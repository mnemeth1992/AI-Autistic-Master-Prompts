"""
Google Drive Synchronization, File Export & Direct Publishing Engine
====================================================================
Manages:
1. Google Drive Cloud API v3 upload via Google Service Account (Streamlit Cloud & Remote).
2. Local Google Drive folder detection and fallback (Laptop / Local filesystem).
3. Structured .txt, .docx, .pdf, .png, and .csv exports.
4. Strict 2026 Etsy CSV generation and sanitization.
5. Gumroad API v2 direct publishing.
"""

import os
import io
import csv
import json
import time
import re
import requests
import logging
from typing import Dict, Any, List, Tuple, Optional
import streamlit as st

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload
    GDRIVE_API_AVAILABLE = True
except ImportError:
    GDRIVE_API_AVAILABLE = False

logger = logging.getLogger("DriveSync")

DEFAULT_DRIVE_ROOT = r"G:\Saját meghajtó\00_VÁLLALKOZÁS_AUDHD_DIGITÁLIS_BIRODALOM"
FALLBACK_LOCAL_ROOT = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "Google_Drive_Mappak")
CONFIG_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "config.json")

DRIVE_FOLDER_MAP = {
    "kdp": "03_📚_AMAZON_KDP",
    "kdp_covers": "03_📚_AMAZON_KDP/02_Borítók_és_Forrásfájlok",
    "kdp_interiors": "03_📚_AMAZON_KDP/01_Belső_PDF_ek",
    "etsy": "04_🖼️_ETSY_DIGITAL",
    "etsy_wallart": "04_🖼️_ETSY_DIGITAL/01_Faliképek_PNG_JPG",
    "etsy_clipart": "04_🖼️_ETSY_DIGITAL/02_Clipart_ZIP_Csomagok",
    "gumroad": "05_📖_GUMROAD_PLR",
    "marketing": "06_📌_MARKETING_ES_SEO"
}


# ─────────────────────────────────────────────────────────
# 1. GOOGLE DRIVE CLOUD API INTEGRATION (SERVICE ACCOUNT)
# ─────────────────────────────────────────────────────────

def get_service_account_info() -> Optional[Dict[str, Any]]:
    """Retrieves Google Service Account credentials from secrets, env, or config.json."""
    # 1. Check Streamlit secrets
    try:
        if "gcp_service_account" in st.secrets:
            return dict(st.secrets["gcp_service_account"])
        if "GOOGLE_SERVICE_ACCOUNT_JSON" in st.secrets:
            raw = st.secrets["GOOGLE_SERVICE_ACCOUNT_JSON"]
            if isinstance(raw, dict):
                return raw
            return json.loads(raw)
    except Exception:
        pass

    # 2. Check environment variable
    env_sa = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()
    if env_sa:
        try:
            return json.loads(env_sa)
        except Exception:
            pass

    # 3. Check config.json
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                raw_cfg = cfg.get("google_service_account_json", "")
                if isinstance(raw_cfg, dict):
                    return raw_cfg
                if isinstance(raw_cfg, str) and raw_cfg.strip():
                    return json.loads(raw_cfg.strip())
        except Exception:
            pass

    return None


def get_gdrive_api_service():
    """Initializes Google Drive v3 API service with Service Account credentials."""
    if not GDRIVE_API_AVAILABLE:
        return None
    sa_info = get_service_account_info()
    if not sa_info:
        return None
    try:
        scopes = ["https://www.googleapis.com/auth/drive"]
        creds = service_account.Credentials.from_service_account_info(sa_info, scopes=scopes)
        service = build("drive", "v3", credentials=creds, cache_discovery=False)
        return service
    except Exception as e:
        logger.warning(f"Failed to create Google Drive API client: {e}")
        return None


def find_or_create_drive_folder(service, folder_name: str, parent_id: Optional[str] = None) -> Optional[str]:
    """Finds an existing folder by name under parent_id or creates a new one."""
    try:
        query = f"mimeType = 'application/vnd.google-apps.folder' and name = '{folder_name}' and trashed = false"
        if parent_id:
            query += f" and '{parent_id}' in parents"

        results = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
        files = results.get('files', [])
        if files:
            return files[0]['id']

        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            file_metadata['parents'] = [parent_id]

        folder = service.files().create(body=file_metadata, fields='id').execute()
        return folder.get('id')
    except Exception as e:
        logger.warning(f"Error in find_or_create_drive_folder ({folder_name}): {e}")
        return None


def upload_to_google_drive_api(
    filename: str,
    file_bytes: bytes,
    folder_category: str = "kdp",
    mime_type: str = "text/plain"
) -> Tuple[bool, str, Optional[str]]:
    """
    Uploads a file directly to Google Drive via Drive API v3.
    Returns (success: bool, web_link_or_error: str, file_id: str).
    """
    service = get_gdrive_api_service()
    if not service:
        return False, "Google Drive API nincs csatlakoztatva.", None

    try:
        # Resolve target subfolder path
        rel_path = DRIVE_FOLDER_MAP.get(folder_category, folder_category)
        path_parts = [p for p in rel_path.replace("\\", "/").split("/") if p]

        # Top root folder in Drive
        root_folder_id = find_or_create_drive_folder(service, "00_VÁLLALKOZÁS_AUDHD_DIGITÁLIS_BIRODALOM")
        curr_parent_id = root_folder_id

        for part in path_parts:
            curr_parent_id = find_or_create_drive_folder(service, part, parent_id=curr_parent_id)
            if not curr_parent_id:
                break

        file_metadata = {
            'name': filename,
            'parents': [curr_parent_id] if curr_parent_id else []
        }
        media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=True)

        uploaded_file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()

        file_id = uploaded_file.get('id')
        web_link = uploaded_file.get('webViewLink') or f"https://drive.google.com/file/d/{file_id}/view"
        logger.info(f"✅ File uploaded to Google Drive Cloud: {filename} -> {web_link}")
        return True, web_link, file_id

    except Exception as e:
        logger.error(f"Google Drive API Upload error: {e}")
        return False, str(e), None


# ─────────────────────────────────────────────────────────
# 2. LOCAL FILE SYSTEM & HYBRID DRIVE MANAGEMENT
# ─────────────────────────────────────────────────────────

def get_drive_root() -> str:
    """Returns detected Google Drive root directory or fallback local folder."""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                saved_path = cfg.get("drive_root_path", "").strip()
                if saved_path and os.path.exists(saved_path):
                    return saved_path
        except Exception:
            pass

    if os.path.exists(DEFAULT_DRIVE_ROOT):
        return DEFAULT_DRIVE_ROOT

    os.makedirs(FALLBACK_LOCAL_ROOT, exist_ok=True)
    return FALLBACK_LOCAL_ROOT


def resolve_drive_folder(folder_key_or_name: str, custom_root: Optional[str] = None) -> str:
    """Resolves and ensures existence of a Drive target folder."""
    root = custom_root if custom_root is not None else get_drive_root()
    rel = DRIVE_FOLDER_MAP.get(folder_key_or_name, folder_key_or_name).replace("/", os.sep)
    target_path = os.path.join(root, rel)
    os.makedirs(target_path, exist_ok=True)
    return target_path


def sanitize_filename(name: str) -> str:
    """Cleans a string to create a safe filename."""
    cleaned = re.sub(r'[\\/*?:"<>|\r\n\t]', " ", name)
    cleaned = re.sub(r'\s+', "_", cleaned).strip("_")
    return cleaned[:40] if cleaned else "Termek"


def save_prompts_file_to_drive(folder_key_or_name: str, theme_title: str, prompt_content: str, header_info: str = "") -> Tuple[bool, str]:
    """
    Hybrid Save:
    1. If Google Drive Cloud API is connected, uploads directly to Google Drive in the cloud.
    2. Also saves to local disk (if available).
    """
    sanitized_title = sanitize_filename(theme_title)
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename = f"Promptek_{sanitized_title}_{timestamp}.txt"

    full_text = ""
    if header_info:
        full_text += f"=== {header_info} ===\n"
        full_text += f"Dátum: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        full_text += f"Téma: {theme_title}\n"
        full_text += "=" * 40 + "\n\n"
    full_text += prompt_content.strip() + "\n"

    file_bytes = full_text.encode("utf-8")

    # 1. Try Google Drive Cloud API upload
    if get_service_account_info():
        ok_cloud, link_or_err, _ = upload_to_google_drive_api(
            filename=filename,
            file_bytes=file_bytes,
            folder_category=folder_key_or_name,
            mime_type="text/plain"
        )
        if ok_cloud:
            return True, f"☁️ Google Drive Felhőbe mentve: {link_or_err}"

    # 2. Local disk save
    try:
        target_dir = resolve_drive_folder(folder_key_or_name)
        file_path = os.path.join(target_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        return True, file_path
    except Exception as e:
        return False, str(e)


def save_binary_file_to_drive(folder_key_or_name: str, filename: str, file_bytes: bytes, mime_type: str = "application/pdf") -> Tuple[bool, str]:
    """Saves PDF, PNG or Word document to Drive (Cloud API + Local disk)."""
    # 1. Cloud API
    if get_service_account_info():
        ok_cloud, link_or_err, _ = upload_to_google_drive_api(
            filename=filename,
            file_bytes=file_bytes,
            folder_category=folder_key_or_name,
            mime_type=mime_type
        )
        if ok_cloud:
            return True, f"☁️ Google Drive Felhőbe mentve: {link_or_err}"

    # 2. Local disk
    try:
        target_dir = resolve_drive_folder(folder_key_or_name)
        file_path = os.path.join(target_dir, filename)
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        return True, file_path
    except Exception as e:
        return False, str(e)


def create_marketing_docx(title: str, content: str, header_info: str = "") -> io.BytesIO:
    """Generates formatted Word (.docx) file in-memory."""
    if not DOCX_AVAILABLE:
        return io.BytesIO()
    doc = docx.Document()
    doc.add_heading(header_info or title, level=0)
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"📅 Készült: {time.strftime('%Y-%m-%d %H:%M:%S')}  |  📌 Termék: {title}").italic = True
    doc.add_paragraph("─" * 40)
    
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        elif block.startswith("- ") or block.startswith("• "):
            for line in block.split("\n"):
                clean_l = line.lstrip("-•* ").strip()
                if clean_l:
                    doc.add_paragraph(clean_l, style="List Bullet")
        else:
            doc.add_paragraph(block)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ─────────────────────────────────────────────────────────
# 3. ETSY CSV ENGINE (2026 STRICT RULES)
# ─────────────────────────────────────────────────────────

def sanitize_etsy_title(title: str, max_chars: int = 140) -> str:
    """Cleans and truncates title to strictly adhere to Etsy's 140 character limit."""
    cleaned = re.sub(r'[\r\n\t]+', ' ', title).strip()
    cleaned = re.sub(r'\s+', ' ', cleaned)
    if len(cleaned) <= max_chars:
        return cleaned
    truncated = cleaned[:max_chars]
    split_pos = max(truncated.rfind(' '), truncated.rfind(','), truncated.rfind('-'))
    if split_pos > 80:
        return truncated[:split_pos].rstrip(' ,-—|')
    return truncated.rstrip(' ,-—|')


def sanitize_etsy_tags(tags_input: Any, max_tags: int = 13, max_tag_len: int = 20) -> List[str]:
    """Cleans, deduplicates, and formats tags into exactly 13 Etsy-compliant tags."""
    raw_list = []
    if isinstance(tags_input, list):
        raw_list = [str(t) for t in tags_input]
    elif isinstance(tags_input, str):
        raw_list = [t.strip() for t in re.split(r'[,;\n\r]+', tags_input) if t.strip()]

    cleaned_tags = []
    seen = set()

    for item in raw_list:
        clean_item = re.sub(r'^[\d\.\-\*\#\s]+', '', item)
        clean_item = re.sub(r'[\"\'\`\$\%\@\!\?\*\(\)\[\]\{\}\<\>\:\;]', '', clean_item)
        clean_item = re.sub(r'\s+', ' ', clean_item).strip().lower()

        if not clean_item:
            continue

        if len(clean_item) > max_tag_len:
            words = clean_item.split()
            sub_tag = ""
            for w in words:
                test_str = (sub_tag + " " + w).strip() if sub_tag else w.strip()
                if len(test_str) <= max_tag_len:
                    sub_tag = test_str
                else:
                    if sub_tag and sub_tag not in seen:
                        valid_tag = sub_tag[:max_tag_len].strip()
                        if valid_tag and valid_tag not in seen:
                            cleaned_tags.append(valid_tag)
                            seen.add(valid_tag)
                    sub_tag = w[:max_tag_len].strip()
            if sub_tag and sub_tag not in seen:
                valid_tag = sub_tag[:max_tag_len].strip()
                if valid_tag and valid_tag not in seen:
                    cleaned_tags.append(valid_tag)
                    seen.add(valid_tag)
        else:
            valid_tag = clean_item[:max_tag_len].strip()
            if valid_tag and valid_tag not in seen:
                cleaned_tags.append(valid_tag)
                seen.add(valid_tag)

    fallback_tags = [
        "christian wall art", "bible verse print", "scripture poster", "christian gift",
        "faith wall decor", "printable wall art", "digital download", "minimalist scripture",
        "kjv bible art", "spiritual wall decor", "christian home decor", "encouraging scripture",
        "christian printable"
    ]

    for fb in fallback_tags:
        if len(cleaned_tags) >= max_tags:
            break
        clean_fb = fb[:max_tag_len].strip().lower()
        if clean_fb not in seen:
            cleaned_tags.append(clean_fb)
            seen.add(clean_fb)

    return [t[:max_tag_len].strip() for t in cleaned_tags[:max_tags]]


def build_etsy_ffc_description(
    product_title: str,
    features_bullets: str,
    emotional_hook: str = "",
    drive_delivery_note: str = "",
    ai_transparency: bool = True
) -> str:
    """Assembles high-conversion FFC product description with Google Drive delivery notice."""
    clean_bullets = features_bullets.strip() if features_bullets else (
        "• 5 High-Resolution 300 DPI File Ratios included (Print over 20+ frame sizes!)\n"
        "• Crystal-clear museum quality vector & watercolor detailing\n"
        "• Instant Access: Download and print at home immediately\n"
        "• Standard ratios: 2:3, 3:4, 4:5, 11:14 and International ISO A1-A4"
    )
    hook_sec = f"{emotional_hook.strip()}\n\n" if emotional_hook.strip() else ""
    delivery_sec = (
        "📥 INSTANT DIGITAL DOWNLOAD DELIVERY (GOOGLE DRIVE):\n"
        "Please note: This is a 100% DIGITAL item. No physical product will be shipped.\n"
        "After purchase, you instantly receive a download guide PDF with your direct Google Drive link "
        "to all full-resolution 300 DPI artwork files."
    )
    if drive_delivery_note.strip():
        delivery_sec += f"\n\n📂 Access Note: {drive_delivery_note.strip()}"

    ai_notice = (
        "\n\n🤖 AI TRANSPARENCY & CRAFTSMANSHIP DISCLOSURE:\n"
        "In compliance with Etsy's Seller Guidelines, this digital artwork was created using advanced generative AI "
        "as a creative partner, then meticulously hand-curated and upscaled to ultra-high 300 DPI print standards."
        if ai_transparency else ""
    )

    return f"""{product_title}

{hook_sec}✨ WHAT IS INCLUDED IN YOUR DOWNLOAD:
{clean_bullets}

{delivery_sec}

🖼️ HOW TO PRINT:
1. Print at home using high-quality matte photo paper or cardstock.
2. Order online through Shutterfly, Printify, Vistaprint, or Walgreens.
3. Bring files to your local print shop (FedEx, Staples, Office Depot).

📜 TERMS OF USE:
For personal use only. Commercial resale of raw digital files is prohibited.
{ai_notice}
"""


def generate_etsy_csv(
    listings: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> Tuple[bool, str, bytes]:
    """Compiles listings into Etsy's official CSV structure."""
    if not listings:
        return False, "Nincsenek feldolgozható termékek a CSV generáláshoz.", b""

    csv_headers = [
        "Title", "Description", "Price", "Quantity", "Tags", "Materials",
        "Section", "Renewal option", "Type", "SKU", "File URL / Delivery Note"
    ]

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=csv_headers, quoting=csv.QUOTE_ALL)
    writer.writeheader()

    for item in listings:
        raw_title = item.get("title", "Christian Digital Wall Art Printable")
        san_title = sanitize_etsy_title(raw_title, 140)
        raw_tags = item.get("tags", [])
        san_tags = sanitize_etsy_tags(raw_tags, max_tags=13, max_tag_len=20)
        tags_str = ", ".join(san_tags)

        raw_desc = item.get("description", "")
        if not raw_desc:
            raw_desc = build_etsy_ffc_description(
                product_title=san_title,
                features_bullets=item.get("features", ""),
                emotional_hook=item.get("emotional_hook", ""),
                drive_delivery_note=item.get("drive_url", "")
            )

        row = {
            "Title": san_title,
            "Description": raw_desc,
            "Price": item.get("price", "6.99"),
            "Quantity": str(item.get("quantity", "999")),
            "Tags": tags_str,
            "Materials": item.get("materials", "Digital Download, 300 DPI PNG, JPG, PDF"),
            "Section": item.get("section", "Christian Wall Art"),
            "Renewal option": "auto",
            "Type": "download",
            "SKU": item.get("sku", f"DIG-{int(time.time())}"),
            "File URL / Delivery Note": item.get("drive_url", "Google Drive PDF Delivery Link")
        }
        writer.writerow(row)

    csv_text = output.getvalue()
    csv_bytes = csv_text.encode("utf-8-sig")

    if output_path:
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        with open(output_path, "wb") as f_csv:
            f_csv.write(csv_bytes)

    return True, output_path or "CSV sikeresen generálva!", csv_bytes


# ─────────────────────────────────────────────────────────
# 4. GUMROAD API V2 PUBLISHER
# ─────────────────────────────────────────────────────────

GUMROAD_API_URL = "https://api.gumroad.com/v2/products"

def publish_to_gumroad(
    product_name: str,
    price_usd: float,
    description: str,
    drive_delivery_url: str = "",
    access_token: Optional[str] = None
) -> Tuple[bool, str, Dict[str, Any]]:
    """Publishes digital product to Gumroad via API v2."""
    token = (access_token or "").strip()
    if not token and os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                token = cfg.get("gumroad_access_token", "").strip()
        except Exception:
            pass

    if not token:
        return False, "⚠️ Nincs megadva Gumroad API Access Token a Beállításokban.", {}

    if not product_name.strip():
        return False, "⚠️ A termék neve nem lehet üres.", {}

    price_in_cents = int(round(max(0.0, float(price_usd)) * 100))
    receipt_message = "Thank you for your purchase!"
    if drive_delivery_url.strip():
        receipt_message += f"\n\n📥 Access your files on Google Drive:\n{drive_delivery_url.strip()}"

    headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
    payload = {
        "name": product_name.strip(),
        "price": price_in_cents,
        "description": description.strip(),
        "custom_receipt": receipt_message,
        "require_shipping": "false",
        "published": "true"
    }

    try:
        res = requests.post(GUMROAD_API_URL, headers=headers, data=payload, timeout=20)
        data = res.json()
        if res.status_code in [200, 201] and data.get("success"):
            p_obj = data.get("product", {})
            url = p_obj.get("short_url") or p_obj.get("url") or f"https://gumroad.com/l/{p_obj.get('id', '')}"
            return True, url, data
        else:
            err = data.get("message") or data.get("error") or res.text
            return False, f"Gumroad Hiba ({res.status_code}): {err}", data
    except Exception as e:
        return False, f"Hálózati hiba Gumroad API hívásakor: {e}", {}
