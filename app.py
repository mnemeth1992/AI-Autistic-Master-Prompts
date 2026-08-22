import os
import io
import json
import time
import datetime
import uuid
import re
import streamlit as st
import streamlit.components.v1 as components
import dotenv
from PIL import Image

# Word (.docx) document support
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load environment variables
dotenv.load_dotenv()

import importlib
import key_manager
importlib.reload(key_manager)
from key_manager import get_key_manager, KeyStatus, generate_image_with_fallback, edit_image_with_fallback
import prompts
importlib.reload(prompts)
from prompts import (
    NICHE_CATEGORIES,
    get_niche_prompt_context,
    get_model_profile,
    IMAGE_MODEL_PROFILES,
    build_ffc_avatar_research_prompt,
    build_ffc_big_domino_hooks_prompt,
    build_ffc_value_stack_prompt,
    build_ffc_sales_letter_prompt,
    build_google_sites_landing_page_prompt,
    build_polsia_landing_page_prompt,
    build_email_funnel_3day_prompt,
    build_email_funnel_30day_prompt,
    build_social_seo_calendar_30day_prompt,
    build_kdp_autopilot_manifest_prompt,
    parse_kdp_autopilot_manifest_json,
    build_strict_etsy_seo_prompt,
    parse_strict_etsy_seo_output,
    build_kdp_dynamic_cover_prompt,
    build_illustrated_book_manifest_prompt,
    parse_illustrated_book_manifest_json
)
import kdp_math
importlib.reload(kdp_math)
from kdp_math import calculate_kdp_cover_dimensions, TRIM_SIZES, PAPER_MULTIPLIERS, parse_trim_size
import kdp_pdf_engine
importlib.reload(kdp_pdf_engine)
from kdp_pdf_engine import build_kdp_book_pdf, build_illustrated_book_pdf, REPORTLAB_AVAILABLE
import etsy_csv_engine
importlib.reload(etsy_csv_engine)
from etsy_csv_engine import sanitize_etsy_title, sanitize_etsy_tags, build_etsy_ffc_description, generate_etsy_csv
import gumroad_publisher
importlib.reload(gumroad_publisher)
from gumroad_publisher import publish_to_gumroad, get_stored_gumroad_token

# FFC & Google Ecosystem Modules
try:
    from modules.ffc_engine import generate_ffc_sales_pack, generate_offline_ffc_pack
    from modules.reels_generator import generate_faceless_reels_batch, generate_reels_broll_image
    from modules.google_hub import get_apps_script_webhook_template, get_stripe_setup_guide, get_google_sites_embed_button
    FFC_MODULES_AVAILABLE = True
except ImportError:
    FFC_MODULES_AVAILABLE = False

# New Research & RAG Modules
try:
    from app.modules.notebooklm_rag import render_notebooklm_rag_module
    from app.modules.tax_calculator_2026 import render_tax_calculator_2026_module
    from app.core.sidecar_dock import render_sidecar_dock
except (ModuleNotFoundError, ImportError):
    try:
        from modules.notebooklm_rag import render_notebooklm_rag_module
        from modules.tax_calculator_2026 import render_tax_calculator_2026_module
        from core.sidecar_dock import render_sidecar_dock
    except Exception:
        render_notebooklm_rag_module = None
        render_tax_calculator_2026_module = None
        render_sidecar_dock = None

# Config & Time Log file persistence
CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
TIME_LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "time_log.json")

# ─────────────────────────────────────────────────────────
# AUDHD 2-HOUR FOCUS & TIMEBOXING CONSTANTS & PERSISTENCE
# ─────────────────────────────────────────────────────────

AUDHD_DAY_PLANS = {
    "Hétfő": {
        "title": "🔍 Hétfő: Tiszta Kutatás & RAG Data Mining (NotebookLM & Gemini)",
        "description": "2 órás mélyfókusz: Amazon/Etsy kulcsszókutatás, 100% pontos KJV igehely- és jelenetkutatás, versenytársi Review Mining.",
        "target_minutes": 120,
        "tasks": [
            "🎯 1. Amazon KDP és Etsy kulcsszókutatás (High Volume, Low Competition) (30 perc)",
            "📖 2. KJV Biblia feltöltése és 30 jelenet/ige kinyerése NotebookLM-ben (30 perc)",
            "⭐ 3. Top versenytársak 1-3 csillagos értékeléseinek Review Mining elemzése (30 perc)",
            "📁 4. Piaci hibák és 5 kiemelkedő termékelőny mentése a projektbe (15 perc)",
            "📋 5. Heti termékspecifikációk véglegesítése a NotebookLM RAG fülön (15 perc)"
        ]
    },
    "Kedd": {
        "title": "📑 Kedd: Teológiai Mátrix & Szöveggenerálás (NotebookLM & Gemini)",
        "description": "2 órás mélyfókusz: 30 napos teológiai mátrix felépítése forrásokból, Gemini Advanced Master Prompt futtatás.",
        "target_minutes": 120,
        "tasks": [
            "🕊️ 1. Teológiai jegyzetek és könyvek feltöltése NotebookLM jegyzetfüzetbe (20 perc)",
            "📑 2. 30 napos táblázatos mátrix legenerálása operatív prompttal (40 perc)",
            "💎 3. Gemini Advanced Master Prompt futtatása a mátrix soraihoz (40 perc)",
            "✍️ 4. 200 szavas reflexiók, imádságok és 3 önreflexiós kérdés finomhangolása (10 perc)",
            "💾 5. Kész kézirat mentése Google Docs-ba és a Drive 05_GUMROAD mappába (10 perc)"
        ]
    },
    "Szerda": {
        "title": "🎨 Szerda: Vizuális Generálás & Képszerkesztés (FLUX / Gemini)",
        "description": "2 órás mélyfókusz: 8.5x11 4K fekete-fehér színezők, 4:5 faliképek, clipart csomagok és többkörös háttéreltávolítás.",
        "target_minutes": 120,
        "tasks": [
            "🎨 1. KDP Színező belső oldalak generálása 4K Master Prompttal (45 perc)",
            "🖼️ 2. Etsy 4:5 Skandináv eukaliptusz igés faliképek generálása (30 perc)",
            "✂️ 3. Clipart illusztrációk generálása tiszta fehér háttérrel (25 perc)",
            "✨ 4. Többkörös beszélgetős háttéreltávolítás (Transparent PNG) (10 perc)",
            "📁 5. Képek mentése és rendszerezése a Google Drive projektmappákba (10 perc)"
        ]
    },
    "Csütörtök": {
        "title": "📐 Csütörtök: Kiadványszerkesztés & PDF Szerkesztés (ReportLab / Canva)",
        "description": "2 órás mélyfókusz: ReportLab nyomdakész KDP belső PDF összeállítása, borító méretezés és Etsy ZIP csomagolás.",
        "target_minutes": 120,
        "tasks": [
            "📖 1. KDP belső PDF összeállítása filcátütés-gátló lapokkal és margókkal (40 perc)",
            "🎨 2. KDP 17.412:11.25 arányú Wrap-Around borító generálása és méretezése (35 perc)",
            "🛍️ 3. Etsy digitális falikép méretcsomagok (4:5, 3:4, 2:3, 1:1) és mockupok készítése (25 perc)",
            "📦 4. Clipart PNG ZIP csomagok és vásárlói útmutatók összeállítása (10 perc)",
            "💾 5. Nyomdakész és letöltendő fájlok mentése a Drive mappákba (10 perc)"
        ]
    },
    "Péntek": {
        "title": "🚀 Péntek: Automata Publikálás & Audio Upsell (KDP, Etsy, Gumroad)",
        "description": "2 órás mélyfókusz: Termékfeltöltések, Pinterest passzív SEO leírások és NotebookLM Audio Devotional ($29->$39) generálás.",
        "target_minutes": 120,
        "tasks": [
            "🛍️ 1. Amazon KDP és Etsy termékek feltöltése 13 SEO taggel és leírással (40 perc)",
            "📌 2. Pinterest passzív SEO címek, leírások és közvetlen linkes pinek készítése (20 perc)",
            "🎙️ 3. NotebookLM Audio Overview (Deep Dive podcast MP3) generálása a kéziratból (35 perc)",
            "💰 4. Gumroad termék publikálás + Audio Companion prémium upsell ($39) beállítása (15 perc)",
            "🏆 5. Heti fókuszblokkok értékelése és felkészülés a képernyőmentes hétvégére (10 perc)"
        ]
    },
    "Szombat": {
        "title": "🌿 Szombat: Pihenés & Regeneráció",
        "description": "Kötelező képernyőmentes idő az idegrendszeri regeneráció és az AudHD túlterhelődés elkerülése érdekében.",
        "target_minutes": 0,
        "tasks": [
            "🌿 Séta a természetben és képernyőmentes offline pihenés",
            "☕ Családi és baráti kapcsolatok ápolása",
            "📖 Csendes olvasás, offline hobbik, idegrendszeri feltöltődés"
        ]
    },
    "Vasárnap": {
        "title": "🕊️ Vasárnap: Lelki Feltöltődés & Csendesség",
        "description": "Lelki megújulás, istentisztelet és felkészülés a következő heti 10 órás aszinkron alkotóciklusra.",
        "target_minutes": 0,
        "tasks": [
            "🕊️ Közösség, istentisztelet és hálaadás",
            "🧘 Csendes elmélkedés és lelki megnyugvás",
            "✨ Motivált, békés ráhangolódás a hétfői kutatási napra"
        ]
    }
}

def load_time_logs() -> list:
    """Loads historical focus sessions from time_log.json."""
    if os.path.exists(TIME_LOG_FILE):
        try:
            with open(TIME_LOG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_time_log_entry(entry: dict) -> bool:
    """Appends a new completed focus session entry into time_log.json."""
    try:
        logs = load_time_logs()
        logs.append(entry)
        with open(TIME_LOG_FILE, "w", encoding="utf-8") as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Error saving time log: {e}")
        return False

def clear_time_logs() -> bool:
    """Clears all historical time logs."""
    try:
        with open(TIME_LOG_FILE, "w", encoding="utf-8") as f:
            json.dump([], f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

def get_current_timer_seconds() -> float:
    """Calculates live elapsed seconds from session state."""
    if st.session_state.get("timer_running", False) and st.session_state.get("timer_start_time"):
        return st.session_state.get("timer_elapsed_seconds", 0) + (time.time() - st.session_state["timer_start_time"])
    return st.session_state.get("timer_elapsed_seconds", 0)

def format_seconds_to_hms(seconds: float) -> str:
    """Formats seconds into HH:MM:SS format."""
    secs = int(max(0, seconds))
    hours = secs // 3600
    minutes = (secs % 3600) // 60
    rem_secs = secs % 60
    return f"{hours:02d}:{minutes:02d}:{rem_secs:02d}"

# ─────────────────────────────────────────────────────────
# GOOGLE DRIVE PATH HELPERS & PROMPT EXPORT (HYBRID CLOUD & LOCAL)
# ─────────────────────────────────────────────────────────

DEFAULT_DRIVE_ROOT = r"G:\Saját meghajtó\00_VÁLLALKOZÁS_AUDHD_DIGITÁLIS_BIRODALOM"
FALLBACK_LOCAL_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Google_Drive_Mappak")

try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload
    GDRIVE_API_AVAILABLE = True
except ImportError:
    GDRIVE_API_AVAILABLE = False


def normalize_service_account_dict(d: Any) -> Optional[Dict[str, Any]]:
    if not d:
        return None
    res = None

    if isinstance(d, dict) or hasattr(d, "items") or hasattr(d, "to_dict"):
        try:
            res = dict(d)
        except Exception:
            try:
                res = {str(k): v for k, v in d.items()}
            except Exception:
                pass

    elif isinstance(d, str):
        s = d.strip()
        # 1. Try standard JSON decode
        for test_s in [s, s.strip("'\""), s.strip("`")]:
            try:
                parsed = json.loads(test_s, strict=False)
                if isinstance(parsed, dict):
                    res = parsed
                    break
            except Exception:
                pass

        # 2. Regex fallback extraction for raw strings/TOML multiline
        if not isinstance(res, dict) and ("client_email" in s or "private_key" in s):
            try:
                email_m = re.search(r'["\']client_email["\']\s*[:=]\s*["\']([^"\']+)["\']', s)
                pk_m = re.search(r'["\']private_key["\']\s*[:=]\s*["\'](-----BEGIN [^"\']+)["\']', s, re.DOTALL)
                proj_m = re.search(r'["\']project_id["\']\s*[:=]\s*["\']([^"\']+)["\']', s)
                if email_m or pk_m:
                    res = {
                        "type": "service_account",
                        "client_email": email_m.group(1).strip() if email_m else "",
                        "private_key": pk_m.group(1).strip() if pk_m else "",
                        "project_id": proj_m.group(1).strip() if proj_m else "",
                        "token_uri": "https://oauth2.googleapis.com/token"
                    }
            except Exception:
                pass

    if isinstance(res, dict) and ("client_email" in res or "private_key" in res):
        norm_res = {}
        for k, v in res.items():
            norm_res[str(k).strip()] = v
        if "private_key" in norm_res and isinstance(norm_res["private_key"], str):
            pk = norm_res["private_key"]
            if "\\n" in pk:
                pk = pk.replace("\\n", "\n")
            norm_res["private_key"] = pk
        if "client_email" in norm_res:
            norm_res["client_email"] = str(norm_res["client_email"]).strip()
        return norm_res

    return None


def get_service_account_info() -> Optional[Dict[str, Any]]:
    """Retrieves Google Service Account credentials from secrets, env, or config.json."""
    # 1. Check Streamlit secrets
    try:
        # Check standard keys
        for k in ["gcp_service_account", "google_service_account", "service_account", "GOOGLE_SERVICE_ACCOUNT_JSON", "google_service_account_json", "SERVICE_ACCOUNT_JSON", "GOOGLE_CREDENTIALS", "google_credentials"]:
            if k in st.secrets:
                norm = normalize_service_account_dict(st.secrets[k])
                if norm:
                    return norm

        # Check if secrets root itself is the service account table
        if "private_key" in st.secrets and "client_email" in st.secrets:
            norm = normalize_service_account_dict(st.secrets)
            if norm:
                return norm

        # Scan all secret keys
        for k in st.secrets:
            v = st.secrets[k]
            norm = normalize_service_account_dict(v)
            if norm:
                return norm
    except Exception:
        pass

    # 2. Check environment variable
    env_sa = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()
    if env_sa:
        norm = normalize_service_account_dict(env_sa)
        if norm:
            return norm

    # 3. Check config.json
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                cfg_f = json.load(f)
                raw_cfg = cfg_f.get("google_service_account_json", "")
                norm = normalize_service_account_dict(raw_cfg)
                if norm:
                    return norm
        except Exception:
            pass

    return None




def get_gdrive_api_service():
    if not GDRIVE_API_AVAILABLE:
        return None
    sa_info = get_service_account_info()
    if not sa_info:
        return None
    try:
        scopes = ["https://www.googleapis.com/auth/drive"]
        creds = service_account.Credentials.from_service_account_info(sa_info, scopes=scopes)
        return build("drive", "v3", credentials=creds, cache_discovery=False)
    except Exception:
        return None


def find_or_create_drive_folder(service, folder_name: str, parent_id: Optional[str] = None) -> Optional[str]:
    try:
        query = f"mimeType = 'application/vnd.google-apps.folder' and name = '{folder_name}' and trashed = false"
        if parent_id:
            query += f" and '{parent_id}' in parents"
        results = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
        files = results.get('files', [])
        if files:
            return files[0]['id']
        file_metadata = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
        if parent_id:
            file_metadata['parents'] = [parent_id]
        folder = service.files().create(body=file_metadata, fields='id').execute()
        return folder.get('id')
    except Exception:
        return None


def upload_to_google_drive_api(
    filename: str,
    file_bytes: bytes,
    folder_category: str = "kdp",
    mime_type: str = "text/plain"
) -> Tuple[bool, str, Optional[str]]:
    service = get_gdrive_api_service()
    if not service:
        return False, "Google Drive API nincs csatlakoztatva.", None
    try:
        rel_path = DRIVE_FOLDER_MAP.get(folder_category, folder_category)
        path_parts = [p for p in rel_path.replace("\\", "/").split("/") if p]
        root_folder_id = find_or_create_drive_folder(service, "00_VÁLLALKOZÁS_AUDHD_DIGITÁLIS_BIRODALOM")
        curr_parent_id = root_folder_id
        for part in path_parts:
            curr_parent_id = find_or_create_drive_folder(service, part, parent_id=curr_parent_id)
            if not curr_parent_id:
                break
        file_metadata = {'name': filename, 'parents': [curr_parent_id] if curr_parent_id else []}
        media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=True)
        uploaded = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        web_link = uploaded.get('webViewLink', f"https://drive.google.com/file/d/{uploaded.get('id')}/view")
        return True, web_link, uploaded.get('id')
    except Exception as e:
        return False, str(e), None


def get_drive_root() -> str:
    cfg = load_config()
    saved_path = cfg.get("drive_root_path", "").strip()
    if saved_path and os.path.exists(saved_path):
        return saved_path
    if os.path.exists(DEFAULT_DRIVE_ROOT):
        return DEFAULT_DRIVE_ROOT
    os.makedirs(FALLBACK_LOCAL_ROOT, exist_ok=True)
    return FALLBACK_LOCAL_ROOT

DRIVE_FOLDER_MAP = {
    "kdp": "03_📚_AMAZON_KDP",
    "kdp_covers": os.path.join("03_📚_AMAZON_KDP", "02_Borítók_és_Forrásfájlok"),
    "kdp_interiors": os.path.join("03_📚_AMAZON_KDP", "01_Belső_PDF_ek"),
    "etsy": "04_🖼️_ETSY_DIGITAL",
    "etsy_wallart": os.path.join("04_🖼️_ETSY_DIGITAL", "01_Faliképek_PNG_JPG"),
    "etsy_clipart": os.path.join("04_🖼️_ETSY_DIGITAL", "02_Clipart_ZIP_Csomagok"),
    "gumroad": "05_📖_GUMROAD_PLR",
    "marketing": "06_📌_MARKETING_ES_SEO"
}

def resolve_drive_folder(folder_key_or_name: str, custom_root: str = None) -> str:
    root = custom_root if custom_root is not None else get_drive_root()
    rel = DRIVE_FOLDER_MAP.get(folder_key_or_name, folder_key_or_name)
    target_path = os.path.join(root, rel)
    os.makedirs(target_path, exist_ok=True)
    return target_path

def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/*?:"<>|\r\n\t]', " ", name)
    cleaned = re.sub(r'\s+', "_", cleaned).strip("_")
    return cleaned[:40] if cleaned else "Termek"

def save_prompts_file_to_drive(folder_key_or_name: str, theme_title: str, prompt_content: str, header_info: str = "") -> Tuple[bool, str]:
    sanitized_title = sanitize_filename(theme_title)
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename = f"Promptek_{sanitized_title}_{timestamp}.txt"
    full_text = ""
    if header_info:
        full_text += f"=== {header_info} ===\n"
        full_text += f"Dátum: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        full_text += f"Téma: {theme_title}\n"
        full_text += "="*40 + "\n\n"
    full_text += prompt_content.strip() + "\n"
    file_bytes = full_text.encode("utf-8")

    # 1. Cloud API
    if get_service_account_info():
        ok_cloud, link_or_err, _ = upload_to_google_drive_api(
            filename=filename,
            file_bytes=file_bytes,
            folder_category=folder_key_or_name,
            mime_type="text/plain"
        )
        if ok_cloud:
            return True, f"☁️ Google Drive Felhőbe mentve: {link_or_err}"

    # 2. Local disk
    try:
        target_dir = resolve_drive_folder(folder_key_or_name)
        file_path = os.path.join(target_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        return True, file_path
    except Exception as e:
        return False, str(e)


def create_marketing_docx(title: str, content: str, header_info: str = "") -> io.BytesIO:
    """Creates a formatted Word (.docx) document in-memory for download or disk saving."""
    if not DOCX_AVAILABLE:
        return io.BytesIO()
    doc = docx.Document()
    doc.add_heading(header_info or title, level=0)
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"📅 Készült: {time.strftime('%Y-%m-%d %H:%M:%S')}  |  📌 Termék: {title}").italic = True
    doc.add_paragraph("─" * 40)
    
    for block in content.split("\n\n"):
        b_clean = block.strip()
        if not b_clean:
            continue
        if b_clean.startswith("# "):
            doc.add_heading(b_clean[2:].strip(), level=1)
        elif b_clean.startswith("## "):
            doc.add_heading(b_clean[3:].strip(), level=2)
        elif b_clean.startswith("### "):
            doc.add_heading(b_clean[4:].strip(), level=3)
        elif b_clean.startswith("1. ") or b_clean.startswith("2. ") or b_clean.startswith("3. ") or b_clean.startswith("- ") or b_clean.startswith("* "):
            for line in b_clean.split("\n"):
                s_line = line.strip()
                if s_line.startswith("- ") or s_line.startswith("* "):
                    doc.add_paragraph(s_line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(s_line)
        else:
            doc.add_paragraph(b_clean)
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def save_marketing_file_to_drive(product_name: str, text_content: str, header_info: str = "Marketing Anyag", content_type_tag: str = "") -> tuple[bool, str, str]:
    """
    Saves generated marketing materials (Sales Letter, Email Funnel, Google Sites Landing Page, SEO Leírás)
    into the Google Drive 06_📌_MARKETING_ES_SEO directory (both Cloud API & local filesystem) as .txt and .docx.
    """
    sanitized_prod = sanitize_filename(product_name)
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    tag = f"_{sanitize_filename(content_type_tag)}" if content_type_tag else ""
    filename_txt = f"Marketing_{sanitized_prod}{tag}_{timestamp}.txt"
    
    full_text = ""
    if header_info:
        full_text += f"=== {header_info} ===\n"
        full_text += f"Dátum: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        full_text += f"Termék: {product_name}\n"
        full_text += "="*40 + "\n\n"
    full_text += text_content.strip() + "\n"

    cloud_link = ""
    if get_service_account_info():
        ok_c, link_c, _ = upload_to_google_drive_api(
            filename=filename_txt,
            file_bytes=full_text.encode("utf-8"),
            folder_category="marketing",
            mime_type="text/plain"
        )
        if ok_c:
            cloud_link = link_c

    try:
        target_dir = resolve_drive_folder("marketing")
        file_path_txt = os.path.join(target_dir, filename_txt)
        
        with open(file_path_txt, "w", encoding="utf-8") as f:
            f.write(full_text)
            
        docx_created = False
        file_path_docx = ""
        if DOCX_AVAILABLE:
            try:
                filename_docx = f"Marketing_{sanitized_prod}{tag}_{timestamp}.docx"
                file_path_docx = os.path.join(target_dir, filename_docx)
                bio = create_marketing_docx(product_name, text_content, header_info=header_info)
                with open(file_path_docx, "wb") as f_docx:
                    f_docx.write(bio.getvalue())
                docx_created = True

                if get_service_account_info():
                    upload_to_google_drive_api(
                        filename=filename_docx,
                        file_bytes=bio.getvalue(),
                        folder_category="marketing",
                        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception:
                pass
                
        details = (f"☁️ Cloud: {cloud_link} | " if cloud_link else "") + file_path_txt + (f" és {file_path_docx}" if docx_created else "")
        return True, details, file_path_txt
    except Exception as e:
        if cloud_link:
            return True, f"☁️ Google Drive Felhő: {cloud_link}", cloud_link
        return False, str(e), ""



def parse_prompts_from_text(text: str, default_category: str = "kdp") -> list[dict]:
    """
    Parses a prompt collection text into a list of item dictionaries:
    [{ 'index': 1, 'title': '...', 'visual_prompt': '...' }]
    Handles both multi-block delimited packages (=== Title ===) and numbered lists (1. Title | Prompt).
    """
    if not text or not text.strip():
        return []

    lines = [l.rstrip() for l in text.strip().splitlines()]
    results = []
    
    # Check if there are multiple === Title === delimiters
    equals_headers = [l for l in lines if l.startswith("===") and l.rstrip().endswith("===") and l.strip("=").strip()]
    
    if len(equals_headers) >= 2:
        current_item = {"title": "", "lines": []}
        for l in lines:
            if l.startswith("===") and l.rstrip().endswith("===") and l.strip("=").strip():
                if current_item["lines"]:
                    prompt_txt = "\n".join(current_item["lines"]).strip()
                    if prompt_txt:
                        results.append({
                            "index": len(results) + 1,
                            "title": current_item["title"] or f"Tétel {len(results) + 1}",
                            "visual_prompt": prompt_txt
                        })
                t_match = re.search(r'===\s*(.*?)\s*===', l)
                current_item = {
                    "title": t_match.group(1).strip() if t_match else f"Tétel {len(results) + 1}",
                    "lines": []
                }
            else:
                if not l.startswith("Dátum:") and not l.startswith("Téma:") and not l.startswith("="):
                    current_item["lines"].append(l)
        if current_item["lines"]:
            prompt_txt = "\n".join(current_item["lines"]).strip()
            if prompt_txt:
                results.append({
                    "index": len(results) + 1,
                    "title": current_item["title"] or f"Tétel {len(results) + 1}",
                    "visual_prompt": prompt_txt
                })
        if results:
            return results

    # Line-by-line / Numbered list parsing (handles single file header or plain list)
    for line in lines:
        s_line = line.strip()
        if not s_line or s_line.startswith("===") or s_line.startswith("#") or s_line.startswith("Dátum:") or s_line.startswith("Téma:"):
            continue

        num_match = re.match(r'^(?:[0-9]{1,3}[\.\)]\s*|\-\s*)(.*)', s_line)
        if num_match:
            content = num_match.group(1).strip()
            if "|" in content:
                parts = [p.strip() for p in content.split("|")]
                title = parts[0]
                vis_prompt = parts[1] if len(parts) > 1 else parts[0]
                results.append({
                    "index": len(results) + 1,
                    "title": title,
                    "visual_prompt": vis_prompt
                })
            else:
                results.append({
                    "index": len(results) + 1,
                    "title": content[:40],
                    "visual_prompt": content
                })
        else:
            if len(s_line) > 10 and not s_line.startswith("FŐ PROMPT") and not s_line.startswith("KAPCSOLÓDÓ"):
                results.append({
                    "index": len(results) + 1,
                    "title": s_line[:40],
                    "visual_prompt": s_line
                })

    return results

STYLE_PRESETS = {
    "🖍️ Kedves Gyerek Vektor (Thick Outlines, Cute Chibi)": (
        "Clean, friendly children's vector illustration style, thick bold black outlines, "
        "simple clean shapes, uniform line weight, no shading, pure white background, cute chibi aesthetic."
    ),
    "🧘 Felnőtt Színező (Intricate Zentangle & Mandala Line Art)": (
        "Highly detailed adult coloring page style, fine black line art, intricate zentangle and mandala patterns, "
        "stained glass window motifs, complex floral borders, no shading, pure white background, stress-relief meditative aesthetic."
    ),
    "🎨 Akvarell & Finom Vonal (Pastel Watercolor & Ink)": (
        "Soft delicate watercolor illustration style with fine ink linework, gentle pastel color palette, "
        "whimsical storybook aesthetic, high detail, clean composition."
    ),
    "🌾 Skandináv Minimál (Scandinavian Boho Line Art)": (
        "Modern Scandinavian minimalist design, clean line art, soft earth tones and eucalyptus green accents, "
        "elegant typography, bohemian aesthetic, 4:5 aspect ratio."
    ),
    "📕 Klasszikus Mesekönyv (Classic Storybook Illustration)": (
        "Classic vintage children's storybook illustration style, warm rich colors, hand-drawn vector art, "
        "nostalgic Christian storybook aesthetic."
    ),
    "✨ 3D Aranyos Modell (Cute 3D Claymation / Soft Studio)": (
        "Cute 3D render style, smooth claymation aesthetic, soft studio lighting, vibrant pastel colors, "
        "friendly characters, Pixar-inspired children's design."
    ),
    "✏️ Egyedi Stílus Megadása...": "CUSTOM"
}

def load_config() -> dict:
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_config(new_data: dict):
    try:
        cfg = load_config()
        cfg.update(new_data)
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.sidebar.error(f"Mentési hiba: {e}")

# Try importing google-genai
GENAI_AVAILABLE = False
try:
    from google import genai
    from google.genai import types
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False

# ─────────────────────────────────────────────────────────
# MASTER PROMPT TEMPLATES (from document)
# ─────────────────────────────────────────────────────────

def template_kdp_coloring(scene: str, aspect_ratio: str = "3:4", trim_size: str = "8.5x11") -> str:
    return (
        f"Create a clean, black-and-white children's coloring page of the following scene: {scene}. "
        f"Style requirements: Use simple, thick black outlines on a pure white background. "
        f"No shading, no gray tones, no colors anywhere. "
        f"Friendly, consistent vector illustration style for children's books, "
        f"with uniform line weight, clean shapes, centered composition, full body framing, and generous safe margins. "
        f"Page aspect ratio: {aspect_ratio}, formatted for Amazon KDP {trim_size} print, 4K resolution."
    )

def template_kdp_adult_coloring(scene: str, aspect_ratio: str = "3:4", trim_size: str = "8.5x11") -> str:
    return (
        f"Create a highly detailed, intricate adult coloring book page depicting the following biblical scene: {scene}. "
        f"Target Audience: Adult Christian coloring book (designed for adult stress relief, meditation, and sophisticated fine line art). "
        f"FIGURE & ANIMAL STYLE (REALISTIC ADULT ENGRAVING): Human figures and animals MUST have realistic anatomical proportions, elegant classic engraving line art, regal dignified features, and flowing realistic draped fabric. "
        f"NO cartoon elements, NO chibi style, NO cute rounded childish eyes, NO simplified toy-like shapes. "
        f"Background & Pattern Requirements: Fine black line art on a pure crisp white background, zero shading, zero grayscale, zero colors. "
        f"Artistic stained glass window composition combined with complex zentangle, botanical, and mandala floral background patterns. "
        f"Composition: Centered, full body subject framing, comfortable safe margins. "
        f"Page aspect ratio: {aspect_ratio}, formatted for Amazon KDP {trim_size} print, 4K resolution."
    )

def template_kdp_cover(title: str = "", theme: str = "", subtitle: str = "", **kwargs) -> str:
    theme_val = kwargs.get("theme", theme)
    title_val = kwargs.get("title", title)
    subtitle_val = kwargs.get("subtitle", subtitle)
    aspect_val = kwargs.get("aspect_ratio", "17.412:11.25")
    sub_clause = f", and subtitle/audience '{subtitle_val}'" if subtitle_val and str(subtitle_val).strip() else ""
    return (
        f"A high-quality colorful book cover depicting {theme_val}, "
        f"consistent color palette, vector art style, "
        f"with bold, clean typography at the top that reads exactly: '{title_val}'{sub_clause}, "
        f"high resolution, professional design, {aspect_val} aspect ratio."
    )

def template_kdp_title_color_tester(title: str, subtitle: str = "This Book Belongs To: ________") -> str:
    return (
        f"A clean, black-and-white title page and color test page (color swatches page) for a coloring book titled '{title}'. "
        f"Features bold, elegant vector typography at the top reading '{title}', followed by a decorative frame reading '{subtitle}'. "
        f"The bottom half contains cute geometric shapes, stars, hearts, and color test swatches labeled 'Color Palette Test' for testing markers and colored pencils. "
        f"Style: Bold clean black outlines, pure white background, no shading, no gray tones, vector line art, 8.5x11 inches, 4K resolution."
    )

def template_kdp_color_guide_scripture(scene_title: str, scripture_verse: str, color_suggestions: str = "Soft sky blue, warm sun yellow, emerald green, and gentle blush pink") -> str:
    return (
        f"A standalone left-hand page for a children's coloring book accompanying the scene '{scene_title}'. "
        f"CRITICAL LAYOUT REQUIREMENT: DO NOT include or draw the main coloring scene, characters, or main illustration on this page — this is a dedicated text, Scripture reference, and color palette reference page placed on the left side opposite the coloring drawing. "
        f"Top section: Clean, friendly typography displaying the page title '{scene_title}' and a warm, simple, kid-friendly paraphrase of the Bible verse (e.g. 'God promised to love and protect Noah and all the animals! - Genesis 6:19') based on: '{scripture_verse}'. "
        f"Center/Bottom section: A decorative color testing swatch palette displaying labeled color boxes and suggested palette combinations: '{color_suggestions}'. "
        f"Border style: Delicate vector leaf/star frame motifs matching the coloring book line art style with thin black outlines on a pure white background. "
        f"8.5x11 inches portrait ratio, 4K resolution."
    )

def template_kdp_adult_calligraphy_scripture(scene_title: str, scripture_verse: str) -> str:
    return (
        f"A standalone left-hand page for an adult Christian coloring book accompanying the scene '{scene_title}'. "
        f"CRITICAL LAYOUT REQUIREMENT: DO NOT include or draw the main coloring scene, characters, or central illustration on this page. "
        f"This page is dedicated purely to elegant typography and decorative line art. "
        f"At the top/center, sophisticated hand-lettered calligraphy displays the full scripture quote and reference: '{scripture_verse}'. "
        f"Surrounded by intricate mandala motifs, fine line art floral borders, and delicate botanical scrollwork matching an adult coloring aesthetic. "
        f"Pure white background, clean fine black outlines, 8.5x11 inches portrait ratio, 4K resolution."
    )

def template_kdp_color_guide_companion(scene_desc: str, scripture_verse: str = "", color_palette: str = "Soft pastel tones") -> str:
    return template_kdp_color_guide_scripture(scene_desc, scripture_verse, color_palette)

def template_kdp_adult_scripture_companion(scene_desc: str, scripture_verse: str = "") -> str:
    return template_kdp_adult_calligraphy_scripture(scene_desc, scripture_verse)

def template_etsy_wall_art(verse: str) -> str:
    return (
        f"An elegant minimalist watercolor design with soft green eucalyptus leaves framing a central text. "
        f"The text reads in beautiful, clean typography: '{verse}', "
        f"high resolution, modern Scandinavian Christian wall art style, 4:5 aspect ratio."
    )

def template_etsy_clipart(subject: str) -> str:
    return (
        f"Watercolor illustration of {subject}, cute chibi style, "
        f"soft pastel color palette, isolated on pure white background, "
        f"high detail, consistent clipart bundle style."
    )

def template_bg_removal() -> str:
    return (
        "Please remove the white background from the graphic above and make it fully transparent. "
        "Output as transparent PNG format."
    )

def template_devotional(audience: str, theme: str, day: str) -> str:
    return (
        f"I want to write a deep, authentic, and spiritually uplifting 30-day Christian devotional journal "
        f"for {audience} titled '{theme}'. "
        f"Please write the complete content for Day {day}. "
        f"Include: one relevant Bible verse (from the KJV translation, in English), "
        f"a 200-word encouraging reflective meditation, a daily prayer, "
        f"and 3 deep, thought-provoking self-reflection questions that believers can write in their journals. "
        f"Style instruction: Avoid typical artificial AI phrases and overly dry writing. "
        f"Write in a warm, deeply encouraging, spiritual, respectful, and human tone — "
        f"as if a believing friend were speaking to another."
    )


# ─────────────────────────────────────────────────────────
# IDEA / TOPIC GENERATOR PROMPT TEMPLATES
# ─────────────────────────────────────────────────────────

def template_idea_kdp(target_audience: str = "gyerekeknek", theme: str = "Ószövetségi és Újszövetségi történetek") -> str:
    return (
        f"Kérlek, gyűjts össze 30 népszerű, kifejező bibliai jelenetet {target_audience} részére szóló színezőkönyvhöz, "
        f"amelynek témája: '{theme}'.\n\n"
        f"FORMÁTUM KÖVETELMÉNYEK:\n"
        f"Minden egyes tételnél adj meg:\n"
        f"1. A jelenet rövid magyar nevét\n"
        f"2. A pontos angol nyelvű visual prompt leírást (pl. 'Noah standing on the deck of the ark with two giraffes')\n"
        f"3. A kapcsolódó Bibliai igehelyet (KJV hivatkozással)\n\n"
        f"Példa kimeneti formátum:\n"
        f"1. Noé a bárka fedélzetén | Noah standing on the deck of the ark with two giraffes | Genesis 6:19\n"
        f"2. Mózes kettéválasztja a Vörös-tengert | Moses parting the Red Sea with his staff held high | Exodus 14:21"
    )

def template_idea_wall_art(theme: str = "Vigasztalás, békesség és remény") -> str:
    return (
        f"Kérlek, gyűjts össze 30 mély, inspiráló és népszerű Bibliai igeidézetet (KJV fordításban, angol nyelven) "
        f"Skandináv stílusú keresztény faliképekhez (Wall Art), melyek fő témája: '{theme}'.\n\n"
        f"FORMÁTUM KÖVETELMÉNYEK:\n"
        f"Adjad meg a sorszámozott listát a pontos angol KJV igeidézettel és a fejezet/vers hivatkozással.\n"
        f"Példa:\n"
        f"1. 'He restores my soul' - Psalm 23:3\n"
        f"2. 'Be still, and know that I am God' - Psalm 46:10\n"
        f"3. 'I can do all things through Christ who strengthens me' - Philippians 4:13"
    )

def template_idea_clipart(bundle_name: str = "Bibliai hősök és angyalok") -> str:
    return (
        f"Kérlek, gyűjts össze 30 egyedi karakter, tárgy vagy szimbólum ötletet egy Etsy Clipart csomaghoz, "
        f"melynek témája: '{bundle_name}'.\n\n"
        f"FORMÁTUM KÖVETELMÉNYEK:\n"
        f"Minden elemnél adj meg egy rövid angol leírást, amit közvetlenül a Clipart Master Promptba lehet illeszteni.\n"
        f"Példa:\n"
        f"1. young biblical Moses holding the stone tablets with golden light\n"
        f"2. baby Jesus resting in a wooden manger with a soft star above\n"
        f"3. Noah's white dove carrying a green olive branch in its beak"
    )

def template_idea_devotional(target_audience: str = "nőknek", main_theme: str = "Reménység a nehéz időkben") -> str:
    return (
        f"Kérlek, tervezz meg egy 30 napos keresztény áhítat napló struktúrát {target_audience} részére "
        f"'{main_theme}' címmel.\n\n"
        f"FORMÁTUM KÖVETELMÉNYEK:\n"
        f"Minden naphoz (1–30. nap) adj meg:\n"
        f"1. A nap rövid alcímét / témáját (magyarul és angolul)\n"
        f"2. A naphoz tartozó kulcsfontosságú Bibliai igét (KJV hivatkozással)\n"
        f"3. Egy 1 mondatos szellemi útravalót\n\n"
        f"Példa:\n"
        f"1. Nap: A remény horgonya (Anchor of Hope) | Jeremiás 29:11 | Isten jövőt és reménységet készített számodra.\n"
        f"2. Nap: Békesség a viharban (Peace in the Storm) | Zsoltárok 46:10 | Állj meg és tudd, hogy Ő az Isten."
    )

def call_gemini_with_retry(client, model_name: str, contents, config=None, max_retries: int = 10, status_widget=None) -> tuple[bool, str]:
    """
    Universal Multi-Provider AI caller with automatic Key Rotation & Fallback Manager.
    Tries Groq (Llama 3.3 70B), Gemini, OpenRouter, and built-in offline template engines.
    Returns (success: bool, text_response_or_error: str).
    """
    km = get_key_manager()
    if isinstance(contents, str):
        sys_inst = ""
        temp = 0.7
        if config:
            if hasattr(config, "system_instruction") and config.system_instruction:
                sys_inst = str(config.system_instruction)
            if hasattr(config, "temperature") and config.temperature is not None:
                temp = float(config.temperature)
        return km.generate_text_with_fallback(
            prompt=contents,
            system_instruction=sys_inst,
            model_name=model_name,
            temperature=temp,
            status_widget=status_widget
        )
    else:
        return km.call_gemini_with_rotation(
            model_name=model_name,
            contents=contents,
            config=config,
            max_retries=max_retries,
            status_widget=status_widget
        )


def generate_ideas_live(client, idea_prompt: str, model_name: str, temp: float = 0.7) -> str:
    km = get_key_manager()
    ok, res = km.generate_text_with_fallback(
        prompt=idea_prompt,
        system_instruction="Te egy profi keresztény kiadványszerkesztő és ötletgazda vagy. Adjad meg a kért 30 ötletet átlátható, sorszámozott, szépen strukturált formában.",
        model_name=model_name,
        temperature=temp
    )
    return res if ok else f"Hiba az ötletek generálásakor: {res}"


def generate_marketing_copy_live(client, prompt: str, model_name: str, temp: float = 0.7, system_instruction: str = "Te egy világklasszis elit keresztény marketing és szövegíró szakértő vagy.") -> str:
    """Generates high-converting copywriting, sales letters, and marketing funnels using Multi-Provider AI."""
    km = get_key_manager()
    ok, res = km.generate_text_with_fallback(
        prompt=prompt,
        system_instruction=system_instruction,
        model_name=model_name,
        temperature=temp
    )
    return res if ok else f"Hiba a marketing szöveg generálásakor: {res}"



# ─────────────────────────────────────────────────────────
# GEMINI VISION MULTIMODAL IMAGE ANALYSIS HELPERS
# ─────────────────────────────────────────────────────────

def analyze_cover_image_for_tester_prompt(client, image_pil: Image.Image, model_name: str, temp: float = 0.7, active_style: str = "") -> str:
    """Uses Gemini Vision to analyze a book cover image and generate a matching Title & Color Tester Page Prompt."""
    if not client or not GENAI_AVAILABLE:
        return "⚠️ Nincs megadva Gemini API kulcs az AI képelemzéshez."

    style_note = f"\n- Enforce strict brand visual art style: {active_style}" if active_style else ""

    prompt = (
        "You are an expert AI prompt engineer and children's book designer. "
        "Analyze this book cover image. Extract the exact main book title text, subtitle (if any), and visual theme. "
        "Based on this cover, craft a professional English image generation prompt for a Title Page and Color Tester Page (color swatches page) for this book. "
        "The prompt MUST instruct: "
        "- Top header displaying the exact book title in bold typography\n"
        "- A decorative frame reading 'This Book Belongs To: ________'\n"
        "- Bottom half displaying cute geometric shapes, stars, and color test swatch boxes labeled 'Color Palette Test' for testing markers\n"
        f"- Style: Simple clean thick black vector outlines, pure white background, no shading, 8.5x11 inches portrait ratio, 4K resolution.{style_note}\n"
        "OUTPUT: Return ONLY the final English image prompt, nothing else."
    )

    config = types.GenerateContentConfig(temperature=temp)
    ok, res = call_gemini_with_retry(client, model_name, [image_pil, prompt], config=config)
    return res if ok else f"Képelemzési hiba: {res}"


def analyze_coloring_image_for_guide_prompt(client, image_pil: Image.Image, model_name: str, temp: float = 0.7, index_num: int = 1, active_style: str = "", status_widget=None, is_adult: bool = False) -> dict:
    """Uses Gemini Vision to analyze a coloring page image and generate a Color Guide & Scripture Page Prompt."""
    if not client or not GENAI_AVAILABLE:
        return {"title": f"Page {index_num}", "prompt": "⚠️ Nincs megadva Gemini API kulcs az AI képelemzéshez."}

    style_note = f"\n- Enforce strict brand visual art style matching: {active_style}" if active_style else ""

    if is_adult:
        prompt = (
            "You are an expert AI prompt engineer and Christian publication designer specializing in adult devotional coloring books. "
            "Analyze this black-and-white coloring page image. "
            "1. Identify the biblical scene depicted in the drawing.\n"
            "2. Give the scene a short, clear 3-5 word title in English.\n"
            "3. Find the exact, full Bible verse reference (book chapter:verse) and full quote (KJV).\n\n"
            "Now, craft a professional English image generation prompt for an accompanying STANDALONE LEFT-HAND Full Scripture Calligraphy Page for adults. "
            "CRITICAL LAYOUT INSTRUCTIONS FOR THE PROMPT:\n"
            "- The prompt MUST explicitly state that this is a STANDALONE LEFT-HAND PAGE (placed opposite the main coloring page).\n"
            "- ABSOLUTELY DO NOT include or draw the main coloring scene, characters, or central illustration on this page.\n"
            "- Main focus: Elegant hand-lettered calligraphy displaying the full scripture quote and reference in the center.\n"
            "- Border & Style: Intricate mandala motifs, fine line art floral borders, delicate botanical scrollwork for adult stress-relief coloring, pure white background, fine black outlines, 8.5x11 inches portrait ratio, 4K resolution.{style_note}\n"
            "OUTPUT: Format your response exactly as:\n"
            "TITLE: [Scene Title]\n"
            "PROMPT: [Full English Image Prompt]"
        )
    else:
        prompt = (
            "You are an expert AI prompt engineer and Christian publication designer specializing in children's Bibles and coloring books. "
            "Analyze this children's black-and-white coloring page image. "
            "1. Identify the biblical scene and key characters/objects depicted in the drawing.\n"
            "2. Give the scene a short, clear 3-5 word title in English.\n"
            "3. Find the relevant Bible verse reference (book chapter:verse) and rewrite it into a short, warm, simple, kid-friendly sentence (NOT literal dry KJV text, but a child-friendly paraphrase suitable for young kids, e.g. 'God always keeps His promises! - Genesis 9:13').\n"
            "4. Suggest a harmonious 4-color palette (e.g. 'Soft sky blue, ocean teal, sun yellow, olive green').\n\n"
            "Now, craft a professional English image generation prompt for an accompanying STANDALONE LEFT-HAND Color Guide and Scripture Reference Page. "
            "CRITICAL LAYOUT INSTRUCTIONS FOR THE PROMPT:\n"
            "- The prompt MUST explicitly state that this is a STANDALONE LEFT-HAND PAGE (placed opposite the main coloring page).\n"
            "- ABSOLUTELY DO NOT include or draw the main coloring scene, characters, or central illustration on this page.\n"
            "- Top header: Clean typography displaying the scene title and the kid-friendly Bible message with exact Scripture reference (e.g. 'God promised to protect Noah! - Genesis 6:19').\n"
            "- Decorative color palette section: Labeled color swatch boxes and suggested color palette names.\n"
            f"- Border & Style: Delicate vector border frame motifs matching the book style, thin black outlines on pure white background, 8.5x11 inches portrait ratio, 4K resolution.{style_note}\n"
            "OUTPUT: Format your response exactly as:\n"
            "TITLE: [Scene Title]\n"
            "PROMPT: [Full English Image Prompt]"
        )

    config = types.GenerateContentConfig(temperature=temp)
    ok, res = call_gemini_with_retry(client, model_name, [image_pil, prompt], config=config, status_widget=status_widget)

    if ok:
        title = f"Page {index_num}"
        prompt_text = res
        if "TITLE:" in res and "PROMPT:" in res:
            parts = res.split("PROMPT:")
            title = parts[0].replace("TITLE:", "").strip()
            prompt_text = parts[1].strip()
        return {"title": title, "prompt": prompt_text}
    else:
        return {"title": f"Page {index_num}", "prompt": f"Képelemzési hiba: {res}"}


def analyze_children_image_for_adult_pair_prompts(client, image_pil: Image.Image, model_name: str, temp: float = 0.7, index_num: int = 1, active_style: str = "", status_widget=None) -> dict:
    """Uses Gemini Vision to analyze a children's coloring page image and generate exact matching Adult Coloring Page & Scripture Calligraphy Prompts."""
    if not client or not GENAI_AVAILABLE:
        return {"title": f"Page {index_num}", "adult_coloring_prompt": "⚠️ Nincs megadva Gemini API kulcs az AI képelemzéshez.", "calligraphy_prompt": ""}

    style_note = f"\n- Enforce strict brand visual art style matching: {active_style}" if active_style else ""

    prompt = (
        "You are a world-class AI prompt engineer and Christian book publisher specializing in adult devotional coloring books. "
        "Analyze this children's coloring page image in detail. "
        "1. Identify the exact biblical scene, main characters, background elements, action, and overall composition.\n"
        "2. Give the scene a short, clear 3-5 word title in English.\n"
        "3. Find the exact Bible verse reference (book chapter:verse) and full quote (KJV).\n\n"
        "Now, craft TWO matching professional English image generation prompts for a paired ADULT EDITION coloring book:\n\n"
        "PROMPT 1 (MAIN ADULT COLORING PAGE PROMPT):\n"
        "Write a highly detailed English image generation prompt that translates this EXACT children's scene into a sophisticated, intricate adult coloring book page.\n"
        "INSTRUCTIONS FOR PROMPT 1:\n"
        "- MUST explicitly begin with: 'Create an intricate adult coloring book page depicting...'\n"
        "- Faithfully recreate the scene composition matching the uploaded image (same characters, action, and setting).\n"
        "- REALISTIC FIGURE & ANIMAL DIRECTIVES: Explicitly state that all human figures (Noah, Daniel, Moses, Jesus, etc.) and animals MUST be rendered with realistic adult anatomical proportions, dignified classic fine-line engraving style, elegant realistic features, and flowing draped fabric. NO cartoon elements, NO chibi style, NO cute rounded childish faces, NO simplified toy-like animal shapes.\n"
        "- BACKGROUND & PATTERNS: Artistic stained glass window composition filled with complex zentangle, botanical, and mandala floral background patterns extending to the edges of the page.\n"
        "- TECHNICAL: Fine black line art on pure white background, zero shading, zero grayscale, zero colors, 8.5x11 inches portrait ratio, 4K resolution.\n"
        "- INTENDED USE: Designed for adult coloring, stress relief, and meditation.{style_note}\n\n"
        "PROMPT 2 (LEFT-HAND CALLIGRAPHY & SCRIPTURE PAGE PROMPT):\n"
        "Write a prompt for a standalone left-hand calligraphy page for adults accompanying this scene.\n"
        "- Standalone left-hand page (DO NOT draw the main scene or characters on this page).\n"
        "- Central feature: Elegant hand-lettered scripture calligraphy displaying the full KJV Bible verse quote and reference.\n"
        "- Border: Intricate mandala motifs, fine line art floral borders, and delicate botanical scrollwork matching adult coloring aesthetic.\n"
        "- Technical: Pure white background, fine black outlines, 8.5x11 inches portrait ratio, 4K resolution.\n\n"
        "OUTPUT FORMAT (STRICT): Format your response exactly as:\n"
        "TITLE: [Scene Title]\n"
        "ADULT_COLORING_PROMPT: [Full English Image Prompt for Adult Coloring Page]\n"
        "CALLIGRAPHY_PROMPT: [Full English Image Prompt for Left-Hand Calligraphy Page]"
    )

    config = types.GenerateContentConfig(temperature=temp)
    ok, res = call_gemini_with_retry(client, model_name, [image_pil, prompt], config=config, status_widget=status_widget)

    if ok:
        title = f"Page {index_num}"
        adult_p = res
        callig_p = ""
        if "TITLE:" in res and "ADULT_COLORING_PROMPT:" in res:
            parts = res.split("ADULT_COLORING_PROMPT:")
            title = parts[0].replace("TITLE:", "").strip()
            rest = parts[1]
            if "CALLIGRAPHY_PROMPT:" in rest:
                sub_parts = rest.split("CALLIGRAPHY_PROMPT:")
                adult_p = sub_parts[0].strip()
                callig_p = sub_parts[1].strip()
            else:
                adult_p = rest.strip()
        return {"title": title, "adult_coloring_prompt": adult_p, "calligraphy_prompt": callig_p}
    else:
        return {"title": f"Page {index_num}", "adult_coloring_prompt": f"Képelemzési hiba: {res}", "calligraphy_prompt": ""}


def generate_gemini_gem_master_instruction(
    is_adult: bool = False,
    active_style: str = "",
    aspect_ratio: str = "3:4",
    trim_size: str = "8.5x11",
    book_title: str = "",
    target_audience: str = "",
    style_name: str = "",
    **kwargs
) -> str:
    """Generates a structured system instruction ready to paste into Gemini Gems ('Instructions' / 'Utasítások' field)."""
    if target_audience:
        is_adult = "adult" in target_audience.lower() or "felnőtt" in target_audience.lower()
    final_style = active_style or style_name or kwargs.get("style", "")
    presets = get_all_gem_presets(final_style, aspect_ratio, trim_size)
    return presets["kdp_adult"]["instruction"] if is_adult else presets["kdp_child"]["instruction"]


def get_all_gem_presets(active_style: str = "", aspect_ratio: str = "3:4", trim_size: str = "8.5x11") -> dict:
    """Returns a dictionary of all 7 pre-configured Gemini Gem instructions adapted to trim size and aspect ratio."""
    style_rule = f"\n- ENFORCED BRAND ART STYLE: {active_style}" if active_style else ""
    clean_ar = aspect_ratio.strip() if aspect_ratio else "3:4"
    return {
        "kdp_adult": {
            "name": "🧘 KDP Felnőtt Bibliai Színező Mester",
            "icon": "🧘",
            "desc": f"Fine line art, valósághű anatómiai figurák, rózsaablak/zentangle mandala hátterek, {clean_ar} arány.",
            "instruction": (
                "You are a world-class AI Image Generation Assistant specializing in KDP Adult Christian Coloring Books.\n\n"
                "SYSTEM DIRECTIVE: STÍLUS RÖGZÍTÉS ÉS KÉPGENERÁLÁSI EGYSÉG (STYLE LOCK FOR ALL IMAGES)\n\n"
                "For every single image request provided by the user, you MUST strictly enforce the following visual requirements:\n\n"
                "1. ART STYLE:\n"
                "   - Fine black line art on a pure crisp white background only.\n"
                "   - ABSOLUTELY ZERO shading, zero grayscale, zero gradients, zero colors.\n\n"
                "2. HUMAN & ANIMAL REALISM DIRECTIVE:\n"
                "   - All human figures (Noah, Daniel, Moses, Jesus, David, etc.) and animals MUST be rendered with realistic adult anatomical proportions, dignified classic fine-line engraving style, elegant detailed facial features, and flowing draped fabric folds.\n"
                "   - STRICTLY NO cartoon elements, NO chibi style, NO cute rounded childish faces, NO simplified toy-like animal shapes.\n\n"
                "3. COMPOSITION & PATTERNS:\n"
                "   - Centered composition with full body framing (head to toe), generous margins around page borders, zero cut-off elements.\n"
                "   - Artistic stained glass window framing combined with complex zentangle, botanical, and mandala floral background patterns extending densely to the page borders.\n\n"
                "4. TARGET AUDIENCE:\n"
                "   - Designed specifically for adult Christian coloring, stress relief, and meditation.\n\n"
                "5. TECHNICAL PARAMETERS:\n"
                f"   - Amazon KDP {trim_size} format, exact aspect ratio: {clean_ar}, 4K resolution, sharp clean line art."
                f"{style_rule}\n\n"
                "GEM INSTRUCTIONS:\n"
                "Whenever the user asks you to generate a scene or page, execute the image generation prompt using these exact style lock rules without deviation."
            )
        },
        "kdp_child": {
            "name": "🧒 KDP Gyermek Bibliai Színező Mester",
            "icon": "🧒",
            "desc": f"Bold vector outlines, letisztult egyszerű háttér, cuki barátságos figura rajzok, {clean_ar} arány.",
            "instruction": (
                "You are a world-class AI Image Generation Assistant specializing in KDP Children's Christian Coloring Books.\n\n"
                "SYSTEM DIRECTIVE: STÍLUS RÖGZÍTÉS ÉS KÉPGENERÁLÁSI EGYSÉG (STYLE LOCK FOR ALL IMAGES)\n\n"
                "For every single image request provided by the user, you MUST strictly enforce the following visual requirements:\n\n"
                "1. ART STYLE:\n"
                "   - Simple, clean, bold black vector outlines on a pure crisp white background.\n"
                "   - ABSOLUTELY ZERO shading, zero grayscale, zero colors.\n\n"
                "2. CHARACTER & ANIMAL DIRECTIVE:\n"
                "   - Cute, friendly, warm illustration style with expressive friendly faces, gentle rounded shapes, and clear easy-to-color sections suitable for young children.\n"
                "   - Medium-wide framing showing the full characters from head to toe, comfortable breathing room around the edges, no cut-off subjects.\n\n"
                "3. COMPOSITION & BACKGROUND:\n"
                "   - Uncluttered minimal white background, large easy-to-color central figures, no overly dense small details.\n\n"
                "4. TARGET AUDIENCE:\n"
                "   - Designed for young children's coloring and Sunday school learning.\n\n"
                "5. TECHNICAL PARAMETERS:\n"
                f"   - Amazon KDP {trim_size} format, exact aspect ratio: {clean_ar}, 4K resolution, bold black outlines."
                f"{style_rule}\n\n"
                "GEM INSTRUCTIONS:\n"
                "Whenever the user asks you to generate a scene or page, execute the image generation prompt using these exact style lock rules without deviation."
            )
        },
        "kdp_cover": {
            "name": "📕 KDP Könyvborító Tervező Mester",
            "icon": "📕",
            "desc": "Profi színes vector KDP borítók 17.412:11.25 méretarányban pontos tipográfiával.",
            "instruction": (
                "You are a world-class AI Image Generation Assistant specializing in KDP Book Cover Designs.\n\n"
                "SYSTEM DIRECTIVE: KDP COVER DESIGN MASTER\n"
                "For every cover requested by the user, generate a stunning, vibrant, print-ready book cover image with:\n"
                "- Vibrant harmonious color palette matching Christian publishing standards.\n"
                "- High-end vector illustration style with clean focal points.\n"
                "- Bold typography at the top reading the exact title provided.\n"
                "- Aspect ratio: Exactly 17.412:11.25 (full bleed KDP wrap-around cover standard) or 8.5x11 inches.\n"
                "- 4K resolution, professional graphic design finish."
                f"{style_rule}"
            )
        },
        "etsy_wallart": {
            "name": "🖼️ Etsy Falikép & Poszter Mester",
            "icon": "🖼️",
            "desc": "Akvarell, Boho, és Aranyfóliás Igés faliképek (300 DPI, 4:5 képarány).",
            "instruction": (
                "You are a world-class AI Image Generation Assistant specializing in Etsy Christian Wall Art & Scripture Prints.\n\n"
                "SYSTEM DIRECTIVE: ETSY WALL ART DESIGN MASTER\n"
                "For every scripture or art scene provided, generate a museum-quality digital print prompt featuring:\n"
                "- Art style options: Soft delicate watercolor, aesthetic boho minimalist, or elegant gold foil typography.\n"
                "- High contrast, beautiful typography for Bible verses.\n"
                "- 4:5 aspect ratio (8x10, 16x20 inches print standard) or 3:4 aspect ratio.\n"
                "- 300 DPI, 4K resolution, ultra-clean aesthetic composition."
            )
        },
        "etsy_clipart": {
            "name": "🎨 Etsy Clipart & Matrica Tervező",
            "icon": "🎨",
            "desc": "Fehér háttéren izolált matrica-szettek és PNG clipart elemek.",
            "instruction": (
                "You are a world-class AI Image Generation Assistant specializing in Etsy Clipart Sets and Sticker Sheets.\n\n"
                "SYSTEM DIRECTIVE: CLIPART & STICKER DESIGN MASTER\n"
                "For every theme provided, generate a multi-item clipart or sticker collection prompt featuring:\n"
                "- Isolated items on a pure solid crisp white background.\n"
                "- Cute sticker borders with subtle white outline framing.\n"
                "- High contrast vector or watercolor illustration style.\n"
                "- 1:1 square ratio, 4K resolution, crisp separation for digital cutting/PNG conversion."
            )
        },
        "gumroad_devotional": {
            "name": "✍️ Gumroad Áhítat & Napló Szövegíró",
            "icon": "✍️",
            "desc": "KJV igealapú áhítatok, imádságok és 30 napos lelki kihívások szövegírója.",
            "instruction": (
                "You are a world-class Christian Author, Devotional Writer, and Curriculum Editor.\n\n"
                "SYSTEM DIRECTIVE: DEVOTIONAL & JOURNAL CONTENT WRITER\n"
                "When given a topic, scripture, or theme, write encouraging, spiritually deep, warm Christian content:\n"
                "- Structure: 1. Relevant Scripture (KJV), 2. Devotional Reflection, 3. Guided Journaling Questions, 4. Closing Prayer.\n"
                "- Tone: Warm, encouraging, grounded in Biblical truth, respectful, inspiring.\n"
                "- Formats: 30-day devotional guides, prayer journals, Sunday school lessons."
            )
        },
        "seo_marketing": {
            "name": "🛍️ Etsy SEO & Gumroad Értékesítési Szövegíró",
            "icon": "🛍️",
            "desc": "Megkapó termékleírások, vásárlási előnyök, és 13 db Etsy SEO keresőcímke (tags).",
            "instruction": (
                "You are an expert E-Commerce Copywriter and Etsy SEO Specialist for Christian digital products.\n\n"
                "SYSTEM DIRECTIVE: ETSY SEO & SALES COPYWRITER\n"
                "When given a product title or theme, generate:\n"
                "1. High-Converting Product Title (optimized for Etsy search).\n"
                "2. Engaging Sales Description featuring bullet points, customer benefits, print instructions, and commercial license terms.\n"
                "3. Exactly 13 Etsy SEO Tags (separated by commas, under 20 characters per tag)."
            )
        }
    }

get_gem_catalog = get_all_gem_presets


# ─────────────────────────────────────────────────────────
# GEMINI PROMPT ENHANCEMENT
# ─────────────────────────────────────────────────────────

ENHANCE_SYSTEM_INSTRUCTION_BASE = """You are a world-class AI prompt engineer specializing in Christian publishing, Etsy digital art products, and KDP self-publishing.

Your task: Take the user's base prompt and enhance it into a highly detailed, professional, and effective prompt for the TARGET IMAGE AI specified below.

{model_enhance_instruction}

CRITICAL MANDATORY RULES:
1. OUTPUT FORMAT: Output ONLY the final enhanced English prompt text — no explanations, no introductions, no markdown formatting, no conversational filler, no bullet points.
2. ADULT COLORING & REALISTIC FIGURES MANDATE: If the base prompt or extra instructions specify an adult coloring page / adult edition / adult audience:
   - You MUST explicitly include the phrase "intricate adult coloring book page" and "designed for adult coloring, stress relief, and meditation".
   - FIGURE & ANIMAL REALISM: You MUST explicitly specify that all human figures (Noah, Daniel, Moses, Jesus, etc.) and animals MUST be rendered with realistic adult anatomical proportions, dignified classic fine-line engraving style, elegant realistic facial features, and flowing fabric folds.
   - NEGATIVE STYLE CONSTRAINTS: Enforce based on the TARGET IMAGE AI rules above.
3. CHILDREN'S COLORING MANDATE: If the base prompt or extra instructions specify a children's coloring page, you MUST explicitly include the phrase "children's coloring book page", "simple thick bold black outlines", and "cute friendly illustration style".
4. COMPOSITION & DETAIL: Be very specific about composition, focal points, line weight (fine line art for adults vs thick vector outlines for kids), background patterns (zentangle/mandala floral patterns for adults vs clean minimal white background for kids), and technical parameters.
5. PRESERVE PARAMETERS: Keep all required dimensions (e.g. 8.5x11 inches, 4K resolution) and KJV scripture quotes from the base prompt intact.
6. QUALITY: Make the result feel professional, publishable, and market-ready."""

# Legacy fallback for any code that references the old constant
ENHANCE_SYSTEM_INSTRUCTION = ENHANCE_SYSTEM_INSTRUCTION_BASE.format(
    model_enhance_instruction="TARGET IMAGE AI: Google Imagen (default)\nUse detailed creative brief style. Explicit negative constraints are OK."
)

def enhance_prompt_with_gemini(client, base_prompt: str, extra_notes: str, model: str, temp: float, active_style: str = "", image_model: str = "") -> str:
    """Uses Gemini to enhance a base prompt with extra user notes, strict brand style enforcement,
    and model-specific prompt engineering rules for the target image AI."""
    # Inject model-specific prompt engineering rules
    profile = get_model_profile(image_model)
    model_enhance_inst = profile.get("enhance_instruction", "TARGET IMAGE AI: Google Imagen (default)\nUse detailed creative brief style.")
    system_instruction_text = ENHANCE_SYSTEM_INSTRUCTION_BASE.format(model_enhance_instruction=model_enhance_inst)
    
    style_instruction = f"\n\nSTRICT BRAND ART STYLE REQUIREMENT (MUST BE ENFORCED IN THE GENERATED PROMPT):\n{active_style}" if active_style else ""
    system_instruction = system_instruction_text + style_instruction

    user_input = base_prompt
    if extra_notes.strip():
        user_input = f"Base prompt:\n{base_prompt}\n\nAdditional instructions from user:\n{extra_notes.strip()}"
    if active_style:
        user_input += f"\n\nStrict visual brand art style requirement to include: {active_style}"

    config = types.GenerateContentConfig(
        system_instruction=system_instruction,
        temperature=temp
    )
    ok, res = call_gemini_with_retry(client, model, user_input, config=config)
    return res if ok else f"[Gemini API hiba: {res}]\n\nAlap prompt (feljavítás nélkül):\n{base_prompt}"


# ─────────────────────────────────────────────────────────
# 📁 GOOGLE DRIVE & LOCAL FOLDER IMAGE GATHERER
# ─────────────────────────────────────────────────────────

def gather_images_from_directory(directory_path: str = "", expected_count: int = 30, target_dir: str = "", **kwargs) -> tuple[dict, list[str]]:
    """
    Scans a folder for images (.png, .jpg, .jpeg, .webp) and automatically maps them
    to page numbers (1..expected_count) and cover. Supports names like '01.png', 'Page_02.png',
    'Noah_03.jpg', 'cover.png', etc.
    Returns: (matched_dict: {page_num: full_path, 'cover': full_path}, unmatched_list: [full_paths])
    """
    matched = {}
    unmatched = []

    dir_path = directory_path or target_dir or kwargs.get("dir", "")
    if not dir_path or not os.path.exists(dir_path):
        return matched, unmatched

    valid_exts = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
    try:
        filenames = sorted([f for f in os.listdir(dir_path) if os.path.splitext(f)[1].lower() in valid_exts])
    except Exception as e:
        logger.error(f"Error scanning directory {dir_path}: {e}")
        return matched, unmatched

    # Pass 1: Match by filename keywords & extracted numbers
    for fn in filenames:
        fp = os.path.join(dir_path, fn)
        base = os.path.splitext(fn)[0].lower()
        if any(k in base for k in ["cover", "borito", "front"]):
            if "cover" not in matched:
                matched["cover"] = fp
                continue

        nums = re.findall(r"\d+", base)
        assigned = False
        if nums:
            p_num = int(nums[0])
            if 1 <= p_num <= expected_count and p_num not in matched:
                matched[p_num] = fp
                assigned = True
            elif len(nums) > 1:
                p_num2 = int(nums[-1])
                if 1 <= p_num2 <= expected_count and p_num2 not in matched:
                    matched[p_num2] = fp
                    assigned = True
        if not assigned:
            unmatched.append(fp)

    # Pass 2: Fallback assign unmatched images sequentially to open page slots
    avail_pages = [p for p in range(1, expected_count + 1) if p not in matched]
    for fp, p in zip(unmatched[:len(avail_pages)], avail_pages):
        matched[p] = fp

    return matched, unmatched


# ─────────────────────────────────────────────────────────
# GEMINI GEM & KDP SCENE CARD RENDERERS
# ─────────────────────────────────────────────────────────

def render_gemini_gem_instruction_card(instruction_text: str, prefix: str = "gem_inst"):
    """
    Renders an elegant, copyable card for the Gemini Gem Master System Directive.
    """
    st.markdown(
        """
        <div style='background: linear-gradient(135deg, rgba(30,58,138,0.25) 0%, rgba(59,130,246,0.15) 100%);
                    border: 1.5px solid #3b82f6; border-radius: 12px; padding: 14px 18px; margin-bottom: 12px;'>
            <div style='display:flex; justify-content:space-between; align-items:center; margin-bottom:6px;'>
                <span style='color:#60a5fa; font-weight:800; font-size:0.92rem;'>💎 GEMINI GEM MESTER UTASÍTÁS (STÍLUSRÖGZÍTÉS)</span>
                <span style='background:#1e3a8a; color:#93c5fd; font-size:0.75rem; font-weight:700; padding:2px 10px; border-radius:12px;'>
                    100% Karakter- & Stílusállandóság
                </span>
            </div>
            <p style='color:#cbd5e1; font-size:0.82rem; margin:0 0 8px 0; line-height:1.4;'>
                Másold be ezt a teljes szöveget a Gemini Web új csevegésébe legelső üzenetként, vagy a Gemini egyéni Gem "Utasítások" mezőjébe!
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.code(instruction_text, language="text")


def render_kdp_scene_cards_grid(scenes_data: list, project_dir: str = "", prefix: str = "kdp_scenes"):
    """
    Renders responsive, structured cards for each scene with:
    - Page number, title (EN & HU), scripture verse
    - Color suggestions
    - Full English image generation prompt (copyable with 1-click)
    - Image status indicator
    """
    if not scenes_data:
        st.info("Nincsenek megjeleníthető jelenetek.")
        return

    num_cols = 2
    cols = st.columns(num_cols, gap="medium")

    for idx, sc in enumerate(scenes_data):
        target_col = cols[idx % num_cols]
        with target_col:
            p_num = sc.get("page_number", idx + 1)
            t_en = sc.get("title", f"Scene {p_num}")
            t_hu = sc.get("title_hu", f"{p_num}. Jelenet")
            s_ref = sc.get("scripture_reference", "")
            s_txt = sc.get("scripture_text", "")
            v_prompt = sc.get("visual_prompt", "")
            colors = sc.get("color_suggestions", [])
            colors_str = ", ".join(colors) if isinstance(colors, list) else str(colors)
            expected_fn = f"{p_num:02d}.png"
            
            # Check if file exists in project_dir
            img_found = False
            if project_dir and os.path.exists(os.path.join(project_dir, expected_fn)):
                img_found = True
            elif sc.get("filepath") and os.path.exists(sc.get("filepath")):
                img_found = True

            badge_html = "<span style='background:#064e3b; color:#34d399; font-size:0.72rem; font-weight:700; padding:2px 8px; border-radius:10px;'>🟢 Kép Letöltve</span>" if img_found else f"<span style='background:#1e293b; color:#94a3b8; font-size:0.72rem; padding:2px 8px; border-radius:10px;'>⚪ Fájlnév: <code>{expected_fn}</code></span>"

            st.markdown(
                f"""
                <div style='background:#1a2338; border:1px solid #2d3748; border-radius:12px; padding:12px 14px; margin-bottom:8px;'>
                    <div style='display:flex; justify-content:space-between; align-items:center; margin-bottom:6px;'>
                        <span style='font-weight:800; color:#38bdf8; font-size:0.9rem;'>#{p_num}. Oldal · {t_en}</span>
                        {badge_html}
                    </div>
                    <div style='color:#e2e8f0; font-size:0.8rem; margin-bottom:4px;'>🇭🇺 <i>{t_hu}</i></div>
                    {f"<div style='color:#facc15; font-size:0.78rem; margin-bottom:4px;'>📖 <b>{s_ref}</b>: \"{s_txt[:90]}...\"</div>" if s_ref else ""}
                    {f"<div style='color:#94a3b8; font-size:0.75rem; margin-bottom:6px;'>🎨 Paletta: {colors_str}</div>" if colors_str else ""}
                </div>
                """,
                unsafe_allow_html=True
            )
            st.code(v_prompt, language="text")


# ─────────────────────────────────────────────────────────
# CANVA / LITPAL-STYLE VISUAL GALLERY & GEMINI ASSISTANT
# ─────────────────────────────────────────────────────────

def render_canva_image_gallery(
    records: list[dict],
    prefix: str = "gal",
    context_type: str = "kdp",
    image_model_name: str = ""
):
    """
    Renders an interactive 3-column visual image gallery optimized for Gemini Web & Drive workflow.
    Displays image previews, prompt copy boxes, manual upload/replacement slots, and file downloads.
    """
    if not records:
        st.info("Nincsenek megjeleníthető képek a galériában.")
        return

    num_cols = 3
    cols = st.columns(num_cols)

    for g_idx, rec in enumerate(records):
        col_target = cols[g_idx % num_cols]
        with col_target:
            item_num = rec.get("index", rec.get("page_number", g_idx + 1))
            item_title = rec.get("title", f"Kép #{item_num}")
            item_filename = rec.get("filename", os.path.basename(rec.get("filepath", f"{item_num:02d}_kep.png")))
            item_prompt = rec.get("prompt", rec.get("visual_prompt", ""))
            item_bytes = rec.get("image_bytes")
            item_pil = rec.get("pil_image")
            item_path = rec.get("filepath", "")

            # If image_bytes exists but pil_image is None
            if item_bytes and item_pil is None:
                try:
                    item_pil = Image.open(io.BytesIO(item_bytes))
                    rec["pil_image"] = item_pil
                except Exception:
                    pass
            elif item_path and os.path.exists(item_path) and item_pil is None:
                try:
                    item_pil = Image.open(item_path)
                    rec["pil_image"] = item_pil
                    with open(item_path, "rb") as f_rd:
                        item_bytes = f_rd.read()
                        rec["image_bytes"] = item_bytes
                except Exception:
                    pass

            dim_info = f"{item_pil.width}×{item_pil.height} px" if item_pil else "Nincs kép"

            # Card Header
            st.markdown(
                f"""
                <div style='background:#1e2536; border:1px solid #2d3748; border-radius:12px 12px 0 0; padding:8px 12px; margin-bottom:-4px;'>
                    <div style='display:flex; justify-content:space-between; align-items:center;'>
                        <span style='font-weight:700; color:#34d399; font-size:0.88rem;'>#{item_num} · {item_title[:28]}</span>
                        <span style='background:#0f172a; color:#94a3b8; font-size:0.72rem; padding:2px 8px; border-radius:10px;'>{dim_info}</span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )

            # Card Image Display
            if item_pil:
                st.image(item_pil, use_container_width=True)
            elif item_bytes:
                st.image(item_bytes, use_container_width=True)
            elif item_path and os.path.exists(item_path):
                st.image(item_path, use_container_width=True)
            else:
                st.markdown(
                    """
                    <div style='height:200px; display:flex; flex-direction:column; align-items:center; justify-content:center; background:#0f172a; color:#64748b; border:1px dashed #334155; border-radius:8px;'>
                        <span style='font-size:2rem; margin-bottom:4px;'>🖼️</span>
                        <span style='font-size:0.82rem;'>Kép még nincs a Drive mappában</span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # Per-Image Card Actions (Prompt copy & replacement)
            with st.expander(f"⚙️ Kép Kezelése & Prompt (#{item_num})", expanded=False):
                st.caption(f"📄 `{item_filename}`")
                
                if item_prompt:
                    st.markdown("**📋 Gemini Prompt ehhez a képhez:**")
                    st.code(item_prompt, language="text")

                uploaded_rep = st.file_uploader(
                    f"📤 Kép feltöltése / cseréje (#{item_num}):",
                    type=["png", "jpg", "jpeg", "webp"],
                    key=f"{prefix}_up_rep_{g_idx}_{item_num}"
                )
                if uploaded_rep:
                    new_bytes = uploaded_rep.getvalue()
                    rec["image_bytes"] = new_bytes
                    rec["pil_image"] = Image.open(io.BytesIO(new_bytes))
                    if item_path:
                        try:
                            with open(item_path, "wb") as f_sv:
                                f_sv.write(new_bytes)
                        except Exception:
                            pass
                    st.toast(f"✅ #{item_num} Kép sikeresen kicserélve!", icon="🖼️")
                    time.sleep(0.4)
                    st.rerun()

                # Download button
                if item_bytes:
                    st.download_button(
                        label="⬇️ PNG Letöltése",
                        data=item_bytes,
                        file_name=item_filename,
                        mime="image/png",
                        key=f"{prefix}_dl_btn_{g_idx}_{item_num}",
                        use_container_width=True
                    )
            
            st.markdown("<div style='margin-bottom:18px;'></div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────
# 📖 ILLUSTRATED BOOK PER-PAGE STORY & PROMPT CARDS
# ─────────────────────────────────────────────────────────

def render_illustrated_book_editor_cards(
    pages_data: list[dict],
    prefix: str = "ib_card",
    image_model_name: str = ""
):
    """
    Renders dual-column interactive editor cards for each illustrated book page:
    - Left column: Page number, Chapter title, editable story text area, and Save Text button.
    - Right column: Illustration preview, Gemini prompt copy box, and Image replacement slot.
    """
    if not pages_data:
        return

    for idx, page in enumerate(pages_data):
        p_num = page.get("page_number", idx + 1)
        c_title = page.get("chapter_title", f"{p_num}. Fejezet")
        s_text = page.get("story_text", "")
        ill_prompt = page.get("illustration_prompt", "")
        img_bytes = page.get("image_bytes")
        img_pil = page.get("pil_image")
        img_path = page.get("filepath", "")

        if img_path and os.path.exists(img_path) and img_pil is None:
            try:
                img_pil = Image.open(img_path)
                page["pil_image"] = img_pil
            except Exception:
                pass

        with st.container():
            st.markdown(
                f"""
                <div style='background: #182030; border: 1.5px solid #2d3c56; border-radius: 14px; padding: 18px 22px; margin-bottom: 22px; box-shadow: 0 4px 14px rgba(0,0,0,0.25);'>
                    <div style='display:flex; justify-content:space-between; align-items:center; border-bottom: 1px solid #28364e; padding-bottom: 10px; margin-bottom: 14px;'>
                        <span style='background:#1e3a8a; color:#93c5fd; font-weight:800; font-size:0.82rem; padding:4px 12px; border-radius:20px; border:1px solid #3b82f6;'>
                            📖 OLDAL #{p_num}
                        </span>
                        <h4 style='margin:0; color:#f8fafc; font-size:1.1rem;'>{c_title}</h4>
                    </div>
                """,
                unsafe_allow_html=True
            )

            col_ib_text, col_ib_img = st.columns([1.1, 1], gap="large")

            # Left column: Story Text Editor
            with col_ib_text:
                st.markdown("<div class='step-label'>📝 Megírt Kézirat & Történet Szöveg</div>", unsafe_allow_html=True)
                new_text = st.text_area(
                    f"#{p_num} Oldal Szövege:",
                    value=s_text,
                    height=240,
                    key=f"{prefix}_txt_area_{idx}",
                    help="Itt közvetlenül átírhatod, kiegészítheted a mese vagy fejezet szövegét."
                )

                if st.button("💾 Szöveg Mentése", key=f"{prefix}_btn_save_txt_{idx}", use_container_width=True):
                    pages_data[idx]["story_text"] = new_text
                    st.toast(f"✅ #{p_num} Oldal szövege sikeresen elmentve!", icon="📝")
                    time.sleep(0.3)
                    st.rerun()

                if page.get("scene_summary"):
                    st.caption(f"**Jelenet összefoglaló:** *{page['scene_summary']}*")

            # Right column: Color Illustration & Gemini Prompt
            with col_ib_img:
                st.markdown("<div class='step-label'>🎨 Illusztráció & Gemini Prompt</div>", unsafe_allow_html=True)

                if img_pil:
                    st.image(img_pil, use_container_width=True)
                elif img_bytes:
                    st.image(img_bytes, use_container_width=True)
                elif img_path and os.path.exists(img_path):
                    st.image(img_path, use_container_width=True)
                else:
                    st.markdown(
                        """
                        <div style='height:200px; display:flex; flex-direction:column; align-items:center; justify-content:center; background:#111827; color:#94a3b8; border:2px dashed #374151; border-radius:10px;'>
                            <span style='font-size:2rem; margin-bottom:4px;'>🖼️</span>
                            <span style='font-size:0.82rem;'>Kép még nincs a Drive mappában</span>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                st.markdown("**📋 Gemini Prompt ehhez az illusztrációhoz:**")
                st.code(ill_prompt, language="text")

                up_ib_img = st.file_uploader(
                    f"📤 Kép feltöltése / cseréje (#{p_num}):",
                    type=["png", "jpg", "jpeg", "webp"],
                    key=f"{prefix}_up_ib_{idx}"
                )
                if up_ib_img:
                    ib_bytes = up_ib_img.getvalue()
                    pages_data[idx]["image_bytes"] = ib_bytes
                    pages_data[idx]["pil_image"] = Image.open(io.BytesIO(ib_bytes))
                    if img_path:
                        try:
                            with open(img_path, "wb") as f_ib_sv:
                                f_ib_sv.write(ib_bytes)
                        except Exception:
                            pass
                    st.toast(f"✅ #{p_num} Kép frissítve!", icon="🎨")
                    time.sleep(0.4)
                    st.rerun()

            st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────
# VIRTUAL FLIPBOOK / LIVE BOOK PREVIEW ENGINE
# ─────────────────────────────────────────────────────────

def render_virtual_book_flipbook_preview(
    scenes_data: list[dict],
    book_title: str = "Bibliai Kifestőkönyv",
    book_subtitle: str = "Békesség & Fókusz Gyermek Kifestő",
    prefix: str = "kdp_flipbook"
):
    """
    Renders a Canva / LitPal-inspired virtual open-book spread preview with dual-page companion
    and coloring views, interactive navigation, and real-time reflection.
    """
    if not scenes_data:
        return

    total_scenes = len(scenes_data)
    
    # Initialize flipbook index in session_state
    fb_key = f"{prefix}_active_idx"
    if fb_key not in st.session_state:
        st.session_state[fb_key] = 0

    curr_idx = max(0, min(st.session_state[fb_key], total_scenes - 1))
    st.session_state[fb_key] = curr_idx
    curr_scene = scenes_data[curr_idx]

    st.markdown("---")
    st.markdown("### 📖 Teljes Könyv Élő Előnézete (Virtuális Kifestőkönyv Lapozó)")
    st.caption("Lapozz végig a készülő kifestőkönyveden nyomdai oldalpáros elrendezésben (Bal oldal: Díszes Igevers & Színpaletta swatchok · Jobb oldal: 300 DPI színező grafika)!")

    # Navigation Controls
    col_nav1, col_nav2, col_nav3 = st.columns([1, 2.2, 1])
    
    with col_nav1:
        if st.button("◀️ Előző Oldalpár", key=f"{prefix}_btn_prev", use_container_width=True, disabled=(curr_idx == 0)):
            st.session_state[fb_key] = max(0, curr_idx - 1)
            st.rerun()

    with col_nav2:
        new_sel = st.slider(
            f"Jelenet kiválasztása ({curr_idx + 1} / {total_scenes}):",
            min_value=1,
            max_value=total_scenes,
            value=curr_idx + 1,
            key=f"{prefix}_slider_page"
        )
        if new_sel - 1 != curr_idx:
            st.session_state[fb_key] = new_sel - 1
            st.rerun()

    with col_nav3:
        if st.button("Következő Oldalpár ▶️", key=f"{prefix}_btn_next", use_container_width=True, disabled=(curr_idx >= total_scenes - 1)):
            st.session_state[fb_key] = min(total_scenes - 1, curr_idx + 1)
            st.rerun()

    # Open Book Container (Split Dual-Page Spread)
    col_left_page, col_spine, col_right_page = st.columns([1.15, 0.04, 1.15])

    # Left Page: Companion Page (Scripture + Palette Swatches + Reflection)
    with col_left_page:
        s_title = curr_scene.get("title", f"{curr_idx+1}. Jelenet")
        s_ref = curr_scene.get("scripture_reference", "Bibliai Igehely")
        s_text = curr_scene.get("scripture_text", "Mert úgy szerette Isten e világot...")
        s_colors = curr_scene.get("color_suggestions", ["Sky Blue", "Olive Green", "Sun Gold", "Earth Brown", "Blush Pink"])
        s_refl = curr_scene.get("reflection_thought", "Gondolkodj el Isten jóságán ezen a napon, miközben kiszínezed a rajzot!")

        palette_chips_html = " ".join([
            f"<span style='display:inline-block; background:#0f172a; color:#38bdf8; border:1.5px solid #0284c7; border-radius:16px; padding:4px 12px; margin:3px; font-size:0.82rem; font-weight:600;'>🎨 {col}</span>"
            for col in s_colors
        ])

        left_page_num = 2 + (curr_idx * 3)
        
        st.markdown(
            f"""
            <div style='background: #182030; border: 2px solid #334155; border-radius: 12px 0 0 12px; padding: 22px 18px; min-height: 480px; box-shadow: -4px 4px 16px rgba(0,0,0,0.35);'>
                <div style='text-align: center; border-bottom: 1.5px dashed #3b82f6; padding-bottom: 10px; margin-bottom: 14px;'>
                    <span style='color: #60a5fa; font-weight: 800; font-size: 0.76rem; text-transform: uppercase; letter-spacing: 0.08em;'>Amazon KDP Kísérő Lap</span>
                    <h4 style='color: #ffffff; margin: 4px 0 0 0; font-size: 1.05rem;'>{s_title}</h4>
                </div>

                <div style='background: #0f172a; border-left: 4px solid #34d399; border-radius: 8px; padding: 12px; margin-bottom: 14px;'>
                    <div style='font-size: 0.8rem; font-weight: 700; color: #34d399; margin-bottom: 4px;'>📖 {s_ref}</div>
                    <div style='font-style: italic; font-size: 0.9rem; color: #e2e8f0; line-height: 1.45;'>"{s_text}"</div>
                </div>

                <div style='margin-bottom: 14px;'>
                    <div style='font-size: 0.8rem; font-weight: 700; color: #cbd5e1; margin-bottom: 4px;'>🎨 Javasolt Színpaletta:</div>
                    <div>{palette_chips_html}</div>
                </div>

                <div style='background: rgba(30, 41, 59, 0.6); border: 1px solid #334155; border-radius: 8px; padding: 12px; margin-bottom: 20px;'>
                    <div style='font-size: 0.78rem; font-weight: 700; color: #94a3b8;'>✨ Áhítatos Gondolat / Reflexió:</div>
                    <div style='font-size: 0.85rem; color: #cbd5e1; margin-top: 4px;'>{s_refl}</div>
                </div>

                <div style='text-align: center; color: #64748b; font-size: 0.75rem; border-top: 1px solid #283347; padding-top: 8px;'>
                    Oldal {left_page_num} (Bal lap) · {book_title}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Center Spine shadow
    with col_spine:
        st.markdown(
            """
            <div style='height: 480px; width: 100%; background: linear-gradient(to right, rgba(0,0,0,0.5), rgba(0,0,0,0.8), rgba(0,0,0,0.5)); border-radius: 2px;'></div>
            """,
            unsafe_allow_html=True
        )

    # Right Page: Coloring Image Page
    with col_right_page:
        s_img_b = curr_scene.get("image_bytes")
        s_img_pil = curr_scene.get("pil_image")
        s_img_path = curr_scene.get("filepath", "")

        right_page_num = 3 + (curr_idx * 3)

        st.markdown(
            f"""
            <div style='background: #ffffff; border: 2px solid #334155; border-radius: 0 12px 12px 0; padding: 14px 18px; min-height: 480px; box-shadow: 4px 4px 16px rgba(0,0,0,0.35); text-align: center; color: #0f172a;'>
                <div style='font-size: 0.75rem; color: #64748b; margin-bottom: 6px; font-weight: 700;'>
                    SZÍNEZŐ LAP #{curr_idx+1} · {s_title}
                </div>
            """,
            unsafe_allow_html=True
        )

        if s_img_pil:
            st.image(s_img_pil, use_container_width=True)
        elif s_img_b:
            st.image(s_img_b, use_container_width=True)
        elif s_img_path and os.path.exists(s_img_path):
            st.image(s_img_path, use_container_width=True)
        else:
            st.markdown(
                """
                <div style='height:340px; display:flex; align-items:center; justify-content:center; background:#f8fafc; color:#94a3b8; border:2px dashed #cbd5e1; border-radius:8px;'>
                    <span>🖼️ Kép generálása folyamatban vagy hiányzik</span>
                </div>
                """,
                unsafe_allow_html=True
            )

        st.markdown(
            f"""
                <div style='text-align: center; color: #64748b; font-size: 0.75rem; border-top: 1px solid #e2e8f0; padding-top: 6px; margin-top: 6px;'>
                    Oldal {right_page_num} (Jobb lap) · {book_title}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )



# ─────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Keresztény Digitális Alkotóműhely",
    page_icon="✝️",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ─────────────────────────────────────────────────────────
# CSS — AUDHD COZY SLATE & VELVET GRAPHITE THEME
# ─────────────────────────────────────────────────────────

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap');

    /* Global Base & Text Styling - Eye-Comfort Calming Slate */
    html, body, [class*="css"], .stApp {
        font-family: 'Plus Jakarta Sans', 'Inter', -apple-system, sans-serif !important;
        background-color: #161b26 !important;
        color: #f1f5f9 !important;
        -webkit-font-smoothing: antialiased;
    }

    /* All text elements - Crisp, Soft Silver-White Readability */
    p, span, label, div, li, a {
        color: #f1f5f9;
    }

    /* Headings */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Plus Jakarta Sans', sans-serif !important;
        color: #ffffff !important;
        font-weight: 700 !important;
        letter-spacing: -0.02em !important;
    }

    h1 {
        color: #34d399 !important;
    }

    h2, h3 {
        color: #f8fafc !important;
    }

    /* Sidebar - Sleek Deep Slate Hub */
    section[data-testid="stSidebar"] {
        background-color: #11141d !important;
        border-right: 1.5px solid #232a3b !important;
        padding-top: 1rem !important;
        box-shadow: 4px 0 16px rgba(0, 0, 0, 0.25) !important;
    }

    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] label {
        color: #e2e8f0 !important;
    }

    /* ══════════════════════════════════════════════════════════
       SIDEBAR NAVIGATION HUB — MODERN WEBSITE SAAS MENU
       ══════════════════════════════════════════════════════════ */
    .nav-title-box {
        padding: 10px 4px 6px 14px !important;
        margin-top: 10px !important;
        margin-bottom: 4px !important;
    }
    .nav-title-text {
        font-size: 0.76rem !important;
        font-weight: 800 !important;
        letter-spacing: 0.09em !important;
        text-transform: uppercase !important;
        color: #34d399 !important;
    }

    /* 1. Hide EVERY radio indicator, dot, circle, checkbox, SVG, bullet */
    section[data-testid="stSidebar"] div[data-testid="stRadio"] input,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] svg,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] [data-baseweb="radio"] > div:first-child,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] [data-baseweb="radio"] span:not([data-testid="stMarkdownContainer"] span),
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label > div:first-child:not([data-testid="stMarkdownContainer"]) {
        display: none !important;
        visibility: hidden !important;
        width: 0 !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        opacity: 0 !important;
    }

    /* 2. Indented vertical list container with equal spacing */
    section[data-testid="stSidebar"] div[data-testid="stRadio"] div[role="radiogroup"],
    section[data-testid="stSidebar"] div[data-testid="stRadio"] > div {
        display: flex !important;
        flex-direction: column !important;
        gap: 8px !important;
        padding-left: 14px !important;
        padding-right: 6px !important;
        margin-top: 4px !important;
    }

    /* 3. Sleek Website-style Nav Button / Pill */
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label {
        background-color: #1a2130 !important;
        border: 1.5px solid #283347 !important;
        border-radius: 12px !important;
        padding: 11px 16px !important;
        margin: 0 !important;
        cursor: pointer !important;
        font-size: 0.93rem !important;
        font-weight: 600 !important;
        color: #cbd5e1 !important;
        display: flex !important;
        align-items: center !important;
        transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 2px 5px rgba(0, 0, 0, 0.25) !important;
        width: 100% !important;
    }

    /* Hover effect */
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label:hover {
        background-color: #242f44 !important;
        border-color: #10b981 !important;
        color: #ffffff !important;
        transform: translateX(4px) !important;
    }

    /* Active Selected Item */
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label:has(input:checked),
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label[data-checked="true"] {
        background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        border-color: #34d399 !important;
        color: #ffffff !important;
        font-weight: 700 !important;
        box-shadow: 0 4px 14px rgba(16, 185, 129, 0.4) !important;
        transform: translateX(4px) !important;
    }

    section[data-testid="stSidebar"] div[data-testid="stRadio"] label:has(input:checked) p,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label:has(input:checked) span,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label[data-checked="true"] p,
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label[data-checked="true"] span {
        color: #ffffff !important;
        font-weight: 700 !important;
    }

    /* Soft Slate Cards */
    .sanctuary-card {
        background-color: #1e2536 !important;
        border: 1.5px solid #2d374d !important;
        border-radius: 16px !important;
        padding: 24px 28px !important;
        margin-bottom: 20px !important;
        box-shadow: 0 4px 14px rgba(0, 0, 0, 0.25) !important;
    }

    /* Step / Section Label */
    .step-label {
        font-size: 0.84rem !important;
        font-weight: 700 !important;
        color: #34d399 !important;
        letter-spacing: 0.05em !important;
        text-transform: uppercase !important;
        margin-bottom: 8px !important;
        display: flex !important;
        align-items: center !important;
        gap: 6px !important;
    }

    /* Path Badge */
    .path-badge {
        display: inline-flex !important;
        align-items: center !important;
        gap: 6px !important;
        background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        color: #ffffff !important;
        font-weight: 700 !important;
        font-size: 0.82rem !important;
        padding: 6px 16px !important;
        border-radius: 20px !important;
        margin-bottom: 12px !important;
        letter-spacing: 0.03em !important;
        text-transform: uppercase !important;
        box-shadow: 0 2px 8px rgba(5, 150, 105, 0.3) !important;
    }

    /* Info Banner */
    .info-banner {
        background-color: #132726 !important;
        border-left: 4px solid #10b981 !important;
        border-radius: 0 12px 12px 0 !important;
        padding: 14px 20px !important;
        font-size: 0.95rem !important;
        color: #a7f3d0 !important;
        margin-bottom: 20px !important;
        line-height: 1.6 !important;
    }

    /* Prompt Output Box - Dark Velvet, Crystal Sharp Readability */
    .prompt-output {
        background-color: #111520 !important;
        border: 1.5px solid #2d374d !important;
        border-radius: 12px !important;
        padding: 18px 22px !important;
        font-family: 'JetBrains Mono', 'Consolas', monospace !important;
        font-size: 0.92rem !important;
        line-height: 1.7 !important;
        color: #f1f5f9 !important;
        white-space: pre-wrap !important;
        word-break: break-word !important;
        margin-top: 12px !important;
        box-shadow: inset 0 2px 6px rgba(0, 0, 0, 0.3) !important;
    }

    /* Primary Emerald Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        color: #ffffff !important;
        font-weight: 700 !important;
        font-size: 0.95rem !important;
        font-family: 'Plus Jakarta Sans', sans-serif !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 12px 26px !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 4px 12px rgba(5, 150, 105, 0.35) !important;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(16, 185, 129, 0.45) !important;
        color: #ffffff !important;
    }

    /* Download Buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #0284c7 0%, #0369a1 100%) !important;
        color: #ffffff !important;
        font-weight: 700 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 12px rgba(2, 132, 199, 0.35) !important;
    }
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #38bdf8 0%, #0284c7 100%) !important;
        transform: translateY(-2px) !important;
    }

    /* Text Inputs, Textareas, Selectboxes */
    .stTextInput input,
    .stTextArea textarea,
    .stSelectbox select,
    div[data-baseweb="select"] > div {
        background-color: #161c28 !important;
        border: 1.5px solid #2d374d !important;
        border-radius: 10px !important;
        color: #f8fafc !important;
        font-size: 0.95rem !important;
        font-weight: 500 !important;
    }

    .stTextInput input:focus,
    .stTextArea textarea:focus,
    div[data-baseweb="select"]:focus-within {
        border-color: #10b981 !important;
        box-shadow: 0 0 0 3px rgba(16, 185, 129, 0.25) !important;
    }

    /* BaseWeb Dropdown Popovers */
    div[data-baseweb="popover"] {
        background-color: #1e2536 !important;
        border: 1.5px solid #2d374d !important;
        border-radius: 12px !important;
    }
    ul[data-testid="stSelectboxVirtualDropdown"] li,
    div[data-baseweb="popover"] li {
        color: #f1f5f9 !important;
        background-color: #1e2536 !important;
    }
    ul[data-testid="stSelectboxVirtualDropdown"] li:hover,
    div[data-baseweb="popover"] li:hover {
        background-color: #283248 !important;
        color: #34d399 !important;
    }

    /* Success / Warning / Info Alert Boxes */
    .stSuccess {
        background-color: #064e3b !important;
        border: 1.5px solid #059669 !important;
        border-radius: 12px !important;
        color: #a7f3d0 !important;
    }

    .stWarning {
        background-color: #451a03 !important;
        border: 1.5px solid #d97706 !important;
        border-radius: 12px !important;
        color: #fde68a !important;
    }

    .stInfo {
        background-color: #0c4a6e !important;
        border: 1.5px solid #0284c7 !important;
        border-radius: 12px !important;
        color: #bae6fd !important;
    }

    /* Expanders */
    .streamlit-expanderHeader {
        background-color: #1e2536 !important;
        border: 1px solid #2d374d !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        color: #f1f5f9 !important;
    }

    /* Tab navigation */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: #11141d;
        border-radius: 14px;
        padding: 6px;
        border-bottom: none;
        margin-bottom: 16px;
    }
    .stTabs [data-baseweb="tab"] {
        background: transparent !important;
        border: none !important;
        border-radius: 10px !important;
        color: #94a3b8 !important;
        font-weight: 600 !important;
        padding: 8px 18px !important;
        transition: all 0.2s ease !important;
    }
    .stTabs [data-baseweb="tab"]:hover {
        color: #ffffff !important;
        background: #1e2536 !important;
    }
    .stTabs [aria-selected="true"] {
        background: #1e2536 !important;
        border: 1px solid #10b981 !important;
        color: #34d399 !important;
        font-weight: 700 !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3) !important;
    }

    /* Divider */
    hr {
        border-color: #232a3b !important;
        margin: 20px 0 !important;
    }

    /* AuDHD 2-Hour Timeboxing & Focus Panel Styles */
    .audhd-box {
        background: linear-gradient(135deg, #131b2c 0%, #1a2338 100%) !important;
        border: 1.5px solid #2d3c56 !important;
        border-radius: 16px !important;
        padding: 18px 22px !important;
        margin-bottom: 20px !important;
        box-shadow: 0 4px 18px rgba(0, 0, 0, 0.3) !important;
    }
    .timer-clock {
        font-family: 'JetBrains Mono', 'Consolas', monospace !important;
        font-size: 2.5rem !important;
        font-weight: 800 !important;
        letter-spacing: 2px !important;
        color: #34d399 !important;
        text-align: center !important;
        background: #0d121f !important;
        padding: 12px 18px !important;
        border-radius: 12px !important;
        border: 1.5px solid #1e293b !important;
        margin: 8px 0 !important;
    }
    .timer-clock-red {
        color: #f87171 !important;
        border-color: #ef444466 !important;
    }
    .badge-green {
        background: #064e3b !important;
        color: #34d399 !important;
        border: 1px solid #10b981 !important;
        font-weight: 700 !important;
        padding: 4px 12px !important;
        border-radius: 20px !important;
        font-size: 0.82rem !important;
        display: inline-block !important;
    }
    .badge-red {
        background: #7f1d1d !important;
        color: #fca5a5 !important;
        border: 1px solid #ef4444 !important;
        font-weight: 700 !important;
        padding: 4px 12px !important;
        border-radius: 20px !important;
        font-size: 0.82rem !important;
        display: inline-block !important;
    }
    .history-card-green {
        background: rgba(6, 78, 59, 0.22) !important;
        border: 1.5px solid #10b981 !important;
        border-radius: 12px !important;
        padding: 14px 18px !important;
        margin-bottom: 12px !important;
    }
    .history-card-red {
        background: rgba(127, 29, 29, 0.22) !important;
        border: 1.5px solid #ef4444 !important;
        border-radius: 12px !important;
        padding: 14px 18px !important;
        margin-bottom: 12px !important;
    }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────
# CONFIGURATION & GLOBAL STATE INITIALIZATION
# ─────────────────────────────────────────────────────────

cfg = load_config()
km = get_key_manager()
api_key = km.get_available_key()[0] or ""
client = True  # Multi-Provider Hub handles Groq, Gemini, OpenRouter, and Offline synthesizers seamlessly!

# Text models
FALLBACK_TEXT_MODELS = [
    "groq-llama-3.3-70b",
    "openrouter-free",
    "gemini-3.7-flash",
    "gemini-2.5-flash",
    "gemini-2.5-pro"
]
saved_text_model = cfg.get("selected_model", "").strip() or "groq-llama-3.3-70b"
available_models = list(dict.fromkeys(
    ([saved_text_model] if saved_text_model else []) +
    (st.session_state.get("live_text_models") or []) +
    (cfg.get("saved_models") or []) +
    FALLBACK_TEXT_MODELS
))
text_model = saved_text_model if saved_text_model in available_models else available_models[0]

temperature = float(cfg.get("temperature", 0.7))

# Drive Root
saved_drive_root = cfg.get("drive_root_path", DEFAULT_DRIVE_ROOT)
drive_root_input = saved_drive_root

# ─────────────────────────────────────────────────────────
# 📂 PERSISTENT BOOK PROJECT MANAGER (AUTO-SAVE & RESUME)
# ─────────────────────────────────────────────────────────

PROJECTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "projects")

def ensure_projects_dir():
    os.makedirs(PROJECTS_DIR, exist_ok=True)

def sanitize_proj_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/*?:"<>|\r\n\t]', " ", name)
    cleaned = re.sub(r'\s+', "_", cleaned).strip("_")
    return cleaned[:40] if cleaned else "Book_Project"

def get_active_project_filepath(project_type: str = "kdp_coloring") -> str:
    ensure_projects_dir()
    return os.path.join(PROJECTS_DIR, f"active_{project_type}.json")

def save_book_project(project_type: str, data: dict, custom_name: str = "") -> str:
    """
    Saves the entire project state (scenes, chapters, prompts, texts, trim size, aspect ratio, paths)
    to a JSON file both as the active session cache and as a timestamped/named project file.
    """
    ensure_projects_dir()
    now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    title = data.get("title", "").strip() or custom_name.strip() or "Untitled_Book"
    
    # Sanitize scenes / pages data to avoid serializing binary image bytes directly
    sanitized_items = []
    items_key = "scenes" if "scenes" in data else ("pages" if "pages" in data else "items")
    raw_items = data.get(items_key, [])
    
    for item in raw_items:
        clean_item = dict(item)
        # Remove raw non-serializable objects (PIL, bytes)
        clean_item.pop("pil_image", None)
        clean_item.pop("image_bytes", None)
        sanitized_items.append(clean_item)
        
    payload = {
        "project_type": project_type,
        "title": title,
        "subtitle": data.get("subtitle", ""),
        "theme": data.get("theme", ""),
        "is_adult": data.get("is_adult", False),
        "trim_size": data.get("trim_size", "8.5x11"),
        "aspect_ratio": data.get("aspect_ratio", "3:4"),
        "layout_mode": data.get("layout_mode", "half_page"),
        "project_dir": data.get("project_dir", ""),
        "saved_at": now_str,
        items_key: sanitized_items
    }
    
    # 1. Save as active project for immediate restore on F5 reload
    active_path = get_active_project_filepath(project_type)
    try:
        with open(active_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error saving active project: {e}")
        
    # 2. Save named project file in projects folder
    sanitized_title = sanitize_proj_filename(title)
    named_filename = f"{project_type}_{sanitized_title}.json"
    named_path = os.path.join(PROJECTS_DIR, named_filename)
    try:
        with open(named_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as e:
        pass
        
    return named_path

def load_book_project(filepath: str) -> Optional[dict]:
    """Loads a project file from disk."""
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading project {filepath}: {e}")
            return None
    return None

def auto_restore_active_project(project_type: str = "kdp_coloring") -> Optional[dict]:
    """Auto-restores the last active project state from disk."""
    active_path = get_active_project_filepath(project_type)
    return load_book_project(active_path)

def list_all_saved_book_projects(project_type_filter: Optional[str] = None) -> List[dict]:
    """Lists all saved project files in the projects/ folder."""
    ensure_projects_dir()
    results = []
    try:
        filenames = sorted([f for f in os.listdir(PROJECTS_DIR) if f.endswith(".json") and not f.startswith("active_")], reverse=True)
        for fn in filenames:
            fp = os.path.join(PROJECTS_DIR, fn)
            try:
                with open(fp, "r", encoding="utf-8") as f:
                    p = json.load(f)
                    p_type = p.get("project_type", "kdp")
                    if project_type_filter and project_type_filter not in p_type:
                        continue
                    results.append({
                        "filename": fn,
                        "filepath": fp,
                        "project_type": p_type,
                        "title": p.get("title", fn.replace(".json", "")),
                        "subtitle": p.get("subtitle", ""),
                        "saved_at": p.get("saved_at", ""),
                        "trim_size": p.get("trim_size", "8.5x11"),
                        "aspect_ratio": p.get("aspect_ratio", "3:4"),
                        "items_count": len(p.get("scenes", p.get("pages", []))),
                        "data": p
                    })
            except Exception:
                pass
    except Exception:
        pass
    return results


# ── Restore active KDP projects on F5 reload if not already in session state ──
if "kdp_autopilot_scenes" not in st.session_state:
    saved_active_kdp = auto_restore_active_project("kdp_coloring")
    if saved_active_kdp and saved_active_kdp.get("data"):
        d_kdp = saved_active_kdp["data"]
        if d_kdp.get("scenes"):
            st.session_state["kdp_autopilot_scenes"] = d_kdp.get("scenes", [])
            st.session_state["kdp_autopilot_book_title"] = d_kdp.get("title", "")
            st.session_state["kdp_autopilot_book_sub"] = d_kdp.get("subtitle", "")
            st.session_state["kdp_autopilot_theme"] = d_kdp.get("theme", "")
            st.session_state["kdp_autopilot_is_adult"] = d_kdp.get("is_adult", False)
            st.session_state["kdp_autopilot_trim_size"] = d_kdp.get("trim_size", "8.5x11")
            st.session_state["kdp_autopilot_aspect_ratio"] = d_kdp.get("aspect_ratio", "3:4")
            st.session_state["kdp_project_dir"] = d_kdp.get("project_dir", "")
            st.session_state["kdp_ap_last_saved_time"] = saved_active_kdp.get("saved_at", "")

if "ib_pages_data" not in st.session_state:
    saved_active_ib = auto_restore_active_project("kdp_illustrated")
    if saved_active_ib and saved_active_ib.get("data"):
        d_ib = saved_active_ib["data"]
        if d_ib.get("pages"):
            st.session_state["ib_pages_data"] = d_ib.get("pages", [])
            st.session_state["ib_book_title"] = d_ib.get("title", "")
            st.session_state["ib_book_sub"] = d_ib.get("subtitle", "")
            st.session_state["ib_book_theme"] = d_ib.get("theme", "")
            st.session_state["ib_trim_size"] = d_ib.get("trim_size", "8.5x8.5")
            st.session_state["ib_aspect_ratio"] = d_ib.get("aspect_ratio", "1:1")
            st.session_state["ib_proj_dir"] = d_ib.get("project_dir", "")
            st.session_state["ib_ed_last_saved_time"] = saved_active_ib.get("saved_at", "")


# ─────────────────────────────────────────────────────────
# 📐 KDP FORMAT & ASPECT RATIO PRESETS & SELECTORS
# ─────────────────────────────────────────────────────────

KDP_FORMAT_PRESETS = {
    "👶 8.5\" × 8.5\" Négyzet (1:1) — Amerikai Gyerekkönyv Bestseller (Ajánlott mesékhez)": {
        "trim_size": "8.5x8.5",
        "aspect_ratio": "1:1",
        "width_in": 8.5,
        "height_in": 8.5,
        "desc": "Az USA #1 legnépszerűbb gyerekkönyv mérete. Tökéletes négyzetes képek, 0 pixel levágás, a Gemini 1:1 natív felbontásban generálja.",
        "badge": "⭐ BESTSELLER GYEREKKÖNYV",
        "badge_color": "#10B981"
    },
    "🎨 8.5\" × 11\" Álló (3:4) — Standard Színező & Munkafüzet (Ajánlott KDP kifestőkhöz)": {
        "trim_size": "8.5x11",
        "aspect_ratio": "3:4",
        "width_in": 8.5,
        "height_in": 11.0,
        "desc": "Az amerikai Letter standard méret. Bőséges hely a színezéshez, elegáns 0.5 hüvelykes biztonsági keretbe illesztve.",
        "badge": "🎨 #1 SZÍNEZŐ STANDARD",
        "badge_color": "#3B82F6"
    },
    "📖 8\" × 10\" Álló (4:5) — Klasszikus Nagyalakú Mesekönyv & Falikép": {
        "trim_size": "8x10",
        "aspect_ratio": "4:5",
        "width_in": 8.0,
        "height_in": 10.0,
        "desc": "Klasszikus nagyalakú képeskönyv arány, prémium illusztrációkhoz és Etsy faliképekhez.",
        "badge": "📖 PRÉMIUM MESEKÖNYV",
        "badge_color": "#8B5CF6"
    },
    "🖼️ 8.5\" × 11\" Féloldalas Fekvő (4:3) — Felső Illusztráció + Alsó Meseszöveg": {
        "trim_size": "8.5x11",
        "aspect_ratio": "4:3",
        "width_in": 8.5,
        "height_in": 11.0,
        "desc": "A lap felső felében fekvő illusztráció, alatta 2-3 bekezdésnyi mese szöveggel.",
        "badge": "🖼️ FÉLOLDALAS MESE",
        "badge_color": "#F59E0B"
    },
    "✝️ 6\" × 9\" Álló (2:3) — Trade Paperback, Áhítat & Regény": {
        "trim_size": "6x9",
        "aspect_ratio": "2:3",
        "width_in": 6.0,
        "height_in": 9.0,
        "desc": "Szöveges könyvek, napi áhítatok, naplók és ifjúsági regények leggyakoribb amerikai mérete.",
        "badge": "✝️ ÁHÍTAT & REGÉNY",
        "badge_color": "#EC4899"
    }
}


def render_kdp_format_and_aspect_selector(widget_prefix: str = "kdp", default_format: str = "8.5x11") -> Tuple[str, str, dict]:
    """
    Renders an informative, interactive KDP Book Format & Aspect Ratio selector directly at the point of generation.
    Returns (trim_size, aspect_ratio, preset_info).
    """
    preset_keys = list(KDP_FORMAT_PRESETS.keys())
    
    # Determine default index
    default_idx = 1 if "11" in default_format else (0 if "8.5" in default_format else 0)
    for idx, k in enumerate(preset_keys):
        if default_format in k:
            default_idx = idx
            break
                
    saved_choice = st.session_state.get(f"{widget_prefix}_format_choice", preset_keys[default_idx])
    if saved_choice not in preset_keys:
        saved_choice = preset_keys[default_idx]
        
    with st.container():
        chosen_format_key = st.selectbox(
            "📐 KDP Könyvformátum & Képarány (Amerikai Piaci Standard):",
            options=preset_keys,
            index=preset_keys.index(saved_choice),
            key=f"{widget_prefix}_format_choice_sel",
            help="Válaszd ki a készülő kiadványod célformátumát! A promptok és a PDF összeállító automatikusan ehhez a mérethez igazodnak."
        )
        st.session_state[f"{widget_prefix}_format_choice"] = chosen_format_key
        
        info = KDP_FORMAT_PRESETS[chosen_format_key]
        trim_size = info["trim_size"]
        aspect_ratio = info["aspect_ratio"]
        
        st.markdown(
            f"""
            <div style='background:#1a2338; border:1px solid #2d3c56; border-radius:10px; padding:10px 14px; margin-top:6px; margin-bottom:12px;'>
                <div style='display:flex; justify-content:space-between; align-items:center; margin-bottom:4px;'>
                    <span style='color:#34d399; font-weight:800; font-size:0.85rem;'>
                        📐 Méret: {info['trim_size']} ({info['width_in']}\" × {info['height_in']}\") · Képarány: <b>{info['aspect_ratio']}</b>
                    </span>
                    <span style='background:#064e3b; color:#a7f3d0; font-size:0.72rem; font-weight:700; padding:2px 8px; border-radius:10px; border:1px solid #10b981;'>
                        {info['badge']}
                    </span>
                </div>
                <div style='color:#cbd5e1; font-size:0.82rem; line-height:1.4;'>
                    {info['desc']}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
        
    return trim_size, aspect_ratio, info


def render_book_project_manager_bar(project_type: str = "kdp_coloring", widget_prefix: str = "kdp_ap"):
    """
    Renders the persistent Project Management Toolbar at the top of the workspace:
    - Lists saved projects for 1-click resume/restore
    - New Project button
    - Save Project Now button
    - Last saved time indicator
    """
    saved_projects = list_all_saved_book_projects(project_type)
    last_saved = st.session_state.get(f"{widget_prefix}_last_saved_time", "")
    
    with st.container():
        st.markdown(
            """
            <div style='background: linear-gradient(135deg, rgba(16,185,129,0.12) 0%, rgba(59,130,246,0.1) 100%);
                        border: 1px solid #2d3c56; border-radius: 12px; padding: 10px 16px; margin-bottom: 16px;'>
                <div style='display:flex; justify-content:space-between; align-items:center;'>
                    <span style='color:#34d399; font-weight:800; font-size:0.88rem;'>
                        📂 KÖNYVPROJEKT MENEDZSER (F5-BIZTOS AUTOMATIKUS MENTÉS)
                    </span>
                    <span style='color:#94a3b8; font-size:0.78rem;'>
                        🔒 Minden változtatás azonnal mentve a lemezre
                    </span>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        col_pm1, col_pm2, col_pm3, col_pm4 = st.columns([1.8, 0.8, 0.8, 0.8])
        
        with col_pm1:
            if saved_projects:
                proj_options = ["-- Válassz egy korábbi könyvprojektet --"] + [f"{p['title']} ({p['saved_at'][:16]})" for p in saved_projects]
                sel_proj = st.selectbox(
                    "Korábbi könyv visszatöltése:",
                    options=proj_options,
                    index=0,
                    key=f"{widget_prefix}_saved_proj_sel",
                    label_visibility="collapsed"
                )
                if sel_proj != "-- Válassz egy korábbi könyvprojektet --":
                    chosen_idx = proj_options.index(sel_proj) - 1
                    target_proj = saved_projects[chosen_idx]
                    if st.button("🔄 Projekt Betöltése & Folytatása", key=f"{widget_prefix}_btn_load_proj", use_container_width=True):
                        # Restore project data into state
                        p_data = target_proj["data"]
                        if project_type == "kdp_coloring":
                            st.session_state["kdp_autopilot_scenes"] = p_data.get("scenes", [])
                            st.session_state["kdp_autopilot_book_title"] = p_data.get("title", "")
                            st.session_state["kdp_autopilot_book_sub"] = p_data.get("subtitle", "")
                            st.session_state["kdp_autopilot_theme"] = p_data.get("theme", "")
                            st.session_state["kdp_autopilot_is_adult"] = p_data.get("is_adult", False)
                            st.session_state["kdp_autopilot_trim_size"] = p_data.get("trim_size", "8.5x11")
                            st.session_state["kdp_autopilot_aspect_ratio"] = p_data.get("aspect_ratio", "3:4")
                            st.session_state["kdp_project_dir"] = p_data.get("project_dir", "")
                            st.session_state["kdp_ap_last_saved_time"] = target_proj.get("saved_at", "")
                        elif project_type == "kdp_illustrated":
                            st.session_state["ib_pages_data"] = p_data.get("pages", [])
                            st.session_state["ib_book_title"] = p_data.get("title", "")
                            st.session_state["ib_book_sub"] = p_data.get("subtitle", "")
                            st.session_state["ib_book_theme"] = p_data.get("theme", "")
                            st.session_state["ib_trim_size"] = p_data.get("trim_size", "8.5x8.5")
                            st.session_state["ib_aspect_ratio"] = p_data.get("aspect_ratio", "1:1")
                            st.session_state["ib_proj_dir"] = p_data.get("project_dir", "")
                            st.session_state["ib_ed_last_saved_time"] = target_proj.get("saved_at", "")
                        st.toast(f"✅ '{target_proj['title']}' sikeresen visszatöltve!", icon="📖")
                        time.sleep(0.4)
                        st.rerun()
            else:
                st.caption("📝 *Még nincsenek elmentett korábbi könyvprojektek.*")
                
        with col_pm2:
            if st.button("💾 Mentés Most", key=f"{widget_prefix}_btn_manual_save", use_container_width=True):
                now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                if project_type == "kdp_coloring" and st.session_state.get("kdp_autopilot_scenes"):
                    save_book_project("kdp_coloring", {
                        "title": st.session_state.get("kdp_autopilot_book_title", "Coloring Book"),
                        "subtitle": st.session_state.get("kdp_autopilot_book_sub", ""),
                        "theme": st.session_state.get("kdp_autopilot_theme", ""),
                        "is_adult": st.session_state.get("kdp_autopilot_is_adult", False),
                        "trim_size": st.session_state.get("kdp_autopilot_trim_size", "8.5x11"),
                        "aspect_ratio": st.session_state.get("kdp_autopilot_aspect_ratio", "3:4"),
                        "project_dir": st.session_state.get("kdp_project_dir", ""),
                        "scenes": st.session_state.get("kdp_autopilot_scenes", [])
                    })
                    st.session_state[f"{widget_prefix}_last_saved_time"] = now_str
                    st.toast("✅ Könyvprojekt sikeresen elmentve!", icon="💾")
                elif project_type == "kdp_illustrated" and st.session_state.get("ib_pages_data"):
                    save_book_project("kdp_illustrated", {
                        "title": st.session_state.get("ib_book_title", "Illustrated Book"),
                        "subtitle": st.session_state.get("ib_book_sub", ""),
                        "theme": st.session_state.get("ib_book_theme", ""),
                        "trim_size": st.session_state.get("ib_trim_size", "8.5x8.5"),
                        "aspect_ratio": st.session_state.get("ib_aspect_ratio", "1:1"),
                        "project_dir": st.session_state.get("ib_proj_dir", ""),
                        "pages": st.session_state.get("ib_pages_data", [])
                    })
                    st.session_state[f"{widget_prefix}_last_saved_time"] = now_str
                    st.toast("✅ Illusztrált könyvprojekt sikeresen elmentve!", icon="💾")
                else:
                    st.toast("⚠️ Még nincs aktív könyvvázlat a mentéshez.", icon="⚠️")
                    
        with col_pm3:
            if st.button("➕ Új Könyv", key=f"{widget_prefix}_btn_new_book", use_container_width=True, help="Tiszta lapot nyit egy új könyv indításához"):
                if project_type == "kdp_coloring":
                    st.session_state.pop("kdp_autopilot_scenes", None)
                    st.session_state.pop("kdp_autopilot_book_title", None)
                    st.session_state.pop("kdp_autopilot_book_sub", None)
                    st.session_state.pop("kdp_autopilot_theme", None)
                    st.session_state.pop("kdp_project_dir", None)
                elif project_type == "kdp_illustrated":
                    st.session_state.pop("ib_pages_data", None)
                    st.session_state.pop("ib_book_title", None)
                    st.session_state.pop("ib_book_sub", None)
                    st.session_state.pop("ib_book_theme", None)
                    st.session_state.pop("ib_proj_dir", None)
                st.toast("✨ Új könyv munkamenet elindítva!", icon="➕")
                time.sleep(0.3)
                st.rerun()
                
        with col_pm4:
            if last_saved:
                st.markdown(f"<div style='font-size:0.75rem; color:#34d399; margin-top:8px;'>🟢 Mentve: {last_saved[11:19]}</div>", unsafe_allow_html=True)
            else:
                st.markdown("<div style='font-size:0.75rem; color:#94a3b8; margin-top:8px;'>⚪ Nincs aktív mentés</div>", unsafe_allow_html=True)


NICHE_DEFAULTS = {
    "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)": {
        "kdp_title": "NOAH'S ARK BIBLE ADVENTURES",
        "kdp_subtitle": "For Kids Ages 4-8 · 30 Inspiring Bible Stories Coloring Book",
        "kdp_theme": "Noah building the ark, animals coming two by two, the great flood, dove with olive branch, the rainbow covenant, and thanksgiving prayer",
        "cov_title": "NOAH'S ARK BIBLE ADVENTURES",
        "cov_sub": "30 Inspiring Bible Stories for Kids Ages 4-8",
        "cov_theme": "Noah's ark on calm water with a bright colorful rainbow and cute smiling animals",
        "storybook_title": "THE BRAVE LITTLE DOVE",
        "storybook_theme": "A gentle white dove sent from Noah's ark carrying an olive branch bringing hope to the world",
        "etsy_verse": "He restores my soul — Psalm 23:3",
        "etsy_subject": "young biblical Moses holding the stone tablets with golden light shining around him",
        "gum_topic": "30 Napos Békesség & Fókusz Keresztény Lelki Napló",
        "gum_title": "30 Napos Békesség & Megújulás Keresztény Digitális Mestercsomag",
        "gum_trans": "Napi 10 perces csendességgel elengedni a szorongást és megtalálni a lelki békét Isten jelenlétében",
        "ffc_prod": "30 Napos Békesség & Fókusz Vezetett Lelki Napló (Printable & Digital)",
        "ffc_aud": "Keresztény édesanyák, alkotók és hívők, akik lelki csendességre vágynak",
        "ffc_trans": "Napi 10 perces vezetett csendességgel elengedni a szorongást, megtalálni a tartós belső békességet és Isten jelenlétében élni",
        "ffc_vehicle": "Mikro-reflexiók, strukturált bibliai igemagyarázatok és letisztult nyomtatható naplólapok",
        "reels_prod": "30 Napos Békesség & Fókusz Vezetett Keresztény Napló",
        "reels_cta_kw": "BÉKESSÉG",
        "gs_prod": "30 Napos Békesség & Fókusz Digitális Napló és Színező",
        "gs_headline": "Találd meg a napi békességet és lelki fókuszt a mindennapok csendjében",
        "gs_tagline": "Egy gyönyörű, nyomtatható 30 napos vezetett áhítat, bibliai igegyűjtemény és művészi színező kollekció a nyugodt, kiegyensúlyozott napokért.",
        "gs_lead_magnet": "3-oldalas ingyenes nyomtatható színező lap és mini áhítat mintacsomag azonnali letöltéssel.",
        "gs_features": "30 napos vezetett áhítat és napló, 30 db 4K felbontású nyomtatható színező oldal, 5 színpaletta ajánló, KJV igehelyek, azonnali digitális PDF letöltés."
    },
    "💼 Online Üzlet & Digitális Termékek (Online Business & PLR)": {
        "kdp_title": "DIGITAL PRODUCT CREATOR BLUEPRINT",
        "kdp_subtitle": "A Step-by-Step Workbook & Planner for Passive Income Creators",
        "kdp_theme": "Modern home office workspace, laptop showing sales dashboard, brainstorming mind maps, product launch checklists, passive income streams",
        "cov_title": "DIGITAL PRODUCT CREATOR BLUEPRINT",
        "cov_sub": "From Zero to First $1,000 in Passive Income with Digital Products",
        "cov_theme": "Sleek modern workspace with golden growth graph, laptop displaying digital shop, and creative design tools",
        "storybook_title": "LEO'S FIRST LEMONADE STAND",
        "storybook_theme": "A curious young boy learning business, saving coins, and building his first successful little shop",
        "etsy_verse": "Build what you love and the growth will follow — Creator Proverb",
        "etsy_subject": "minimalist 3D isometric digital shop icon with gold coins and sparkle stars",
        "gum_topic": "0-ról az Első 1000 Dollárig Digitális Termékekkel",
        "gum_title": "Digitális Termék Indítási Mestercsomag & PLR Értékesítési Tölcsér",
        "gum_trans": "Megalkotni és automatizálni az első digitális termékedet 0 Ft-os eszközökkel, technikai tudás nélkül",
        "ffc_prod": "Faceless Funnel & Digitális Termék Mestercsomag (DFY Sablonok)",
        "ffc_aud": "Kezdő digitális alkotók és szabadúszók, akik passzív jövedelmet akarnak építeni arc nélkül",
        "ffc_trans": "Létrehozni egy teljesen automata digitális értékesítési tölcsért, ami napi 10-20 passzív vásárlást hoz",
        "ffc_vehicle": "Arc nélküli rövid videók (Faceless Reels) + ManyChat kulcsszó automatizáció + 0 Ft-os Stripe tölcsér",
        "reels_prod": "Arc Nélküli Digitális Termék Értékesítési Rendszer",
        "reels_cta_kw": "SZABADSÁG",
        "gs_prod": "Faceless Funnel Digitális Termék Mestercsomag",
        "gs_headline": "Építs 100% Passzív Jövedelmet Arc Nélküli Rövid Videókkal és Digitális Termékekkel",
        "gs_tagline": "Azonnal használható DFY sablonok, értékesítési tölcsérek és automatizált Stripe rendszerek a szabadabb életért.",
        "gs_lead_magnet": "Ingyenes 10-Lépéses Digitális Termék Indítási Ellenőrzőlista és Vázlat Sablon azonnali letöltéssel.",
        "gs_features": "10 db kész Reels forgatókönyv, ManyChat sablon, Google Apps Script webhook, 30 napos e-mail tölcsér és Canva értékesítési csomag."
    },
    "📈 Befektetés & Személyes Pénzügyek (Investing & Personal Finance)": {
        "kdp_title": "FINANCIAL FREEDOM & WEALTH TRACKER",
        "kdp_subtitle": "Budgeting, Sinking Funds & Compound Interest Milestone Planner",
        "kdp_theme": "Financial freedom vision board, compounding wealth tree, piggy bank milestones, debt-free thermometer, investment dividend charts",
        "cov_title": "FINANCIAL FREEDOM & WEALTH TRACKER",
        "cov_sub": "Master Your Money, Eliminate Debt & Build Generational Wealth",
        "cov_theme": "Deep navy blue and emerald green geometric money tree growing under golden stars with clean charts",
        "storybook_title": "PENNY'S MAGIC MONEY JAR",
        "storybook_theme": "A smart girl named Penny learning the power of saving, investing, and watching her seed grow into a big money tree",
        "etsy_verse": "Wealth from get-rich-quick schemes quickly disappears; wealth from hard work grows over time. — Proverbs 13:11",
        "etsy_subject": "golden coin stack sprouting a fresh green plant sprout with radiant glow",
        "gum_topic": "Pénzügyi Függetlenség & Vagyonépítő Naptár",
        "gum_title": "Pénzügyi Szabadság Mestercsomag: Költségvetés, Adósságcsökkentés & Befektetés",
        "gum_trans": "Teljes kontrollt szerezni a pénzügyeid felett, megszabadulni az adósságoktól és felépíteni a vésztartalékot",
        "ffc_prod": "30 Napos Pénzügyi Áttörés & Vagyonépítő Digitális Rendszer",
        "ffc_aud": "Tudatos magánszemélyek és családok, akik stabil anyagi biztonságot akarnak",
        "ffc_trans": "Megszüntetni a hóvégi pénzügyi stresszt és automatizálni a megtakarításaidat",
        "ffc_vehicle": "3-Kasszás költségvetési módszer + Automatikus befektetési kalkulátor + Napi költéskövető",
        "reels_prod": "Pénzügyi Függetlenség és Befektetési Trükkök",
        "reels_cta_kw": "SIKER",
        "gs_prod": "Pénzügyi Szabadság & Vagyonépítő Rendszer",
        "gs_headline": "Vedd Át Az Irányítást A Pénzed Felett És Építs Tartós Vagyont",
        "gs_tagline": "Gyakorlatias, azonnal nyomtatható és digitális pénzügyi tervező, kalkulátorok és lépésről lépésre vezetett útmutatók.",
        "gs_lead_magnet": "Ingyenes Vészhelyzeti Alap & Adósságtörlesztő Tervező Sablon azonnali letöltéssel.",
        "gs_features": "12 havi költségvetés tervező, adósság-hógolyó tracker, kamatos kamat vizualizáció és 30 napos pénzügyi kihívás."
    },
    "⚡ Produktivitás & Notion Rendszerek (Productivity/Notion)": {
        "kdp_title": "MINIMALIST NOTION FOCUS & HABIT PLANNER",
        "kdp_subtitle": "ADHD-Friendly Daily Time-Blocking & Deep Work Journal",
        "kdp_theme": "Clean aesthetic desk setup, Pomodoro timer, habit tracker grid, daily top 3 priorities list, mindful reflection box, clutter-free space",
        "cov_title": "MINIMALIST NOTION FOCUS & HABIT PLANNER",
        "cov_sub": "Defeat Procrastination and Build Laser-Sharp Focus",
        "cov_theme": "Minimalist monochrome aesthetic layout with subtle sage green accents and geometric focus grid",
        "storybook_title": "THE CLOCKWORK OWL",
        "storybook_theme": "A wise little owl who teaches forest animals how to organize their acorns and finish big projects on time",
        "etsy_verse": "Focus on being productive instead of busy. — Tim Ferriss",
        "etsy_subject": "aesthetic minimalist hourglass with glowing golden sand and botanical leaves",
        "gum_topic": "ADHD-Barát Napi Fókusz & Szokásépítő Rendszer",
        "gum_title": "ADHD-Barát Produktivitási Mestercsomag & Notion Életrendező",
        "gum_trans": "Leküzdeni a halogatást, rendszerezni a szétszórt gondolatokat és elérni a napi flow élményt",
        "ffc_prod": "ADHD-Barát 120-Perces Fókusz & Időgazdálkodási Mestercsomag",
        "ffc_aud": "ADHD-s alkotók, túlterhelt szakemberek és vizsgázó egyetemisták",
        "ffc_trans": "Megszüntetni a szétesettséget és napi 2 óra mélyfókuszban elvégezni egy egész napi teendőt",
        "ffc_vehicle": "120 perces Pomodoro ritmus + Döntésmentes sablonok + Napi 3-as prioritás szűrő",
        "reels_prod": "ADHD Fókusz és Produktivitási Trükkök",
        "reels_cta_kw": "FÓKUSZ",
        "gs_prod": "ADHD-Barát Fókusz & Produktivitási Rendszer",
        "gs_headline": "Győzd Le A Halogatást És Találd Meg A Nyugodt, Fókuszált Munka Örömét",
        "gs_tagline": "Súrlódásmentes, vizuális rendszerező eszközök és időblokkoló naplók a túlterheltség ellen.",
        "gs_lead_magnet": "Ingyenes 1-Oldalas Napi Fókusz & Dopamin Tracker azonnali letöltéssel.",
        "gs_features": "30 napos fókusz napló, időblokkoló sablonok, Notion életműszerfal sablon és hangulat-tracker."
    },
    "🧘 Mentális Egészség & Stresszoldás (Mental Health & Stress Relief)": {
        "kdp_title": "CALM MINDFULNESS & ANXIETY RELIEF COLORING",
        "kdp_subtitle": "Gentle Mandala & Calming Nature Scenes for Stress Relief",
        "kdp_theme": "Soothing ocean waves, tranquil zen rock garden, gentle blooming lavender field, warm cup of tea by rainy window, peaceful forest canopy",
        "cov_title": "CALM MINDFULNESS & ANXIETY RELIEF",
        "cov_sub": "A Soothing Coloring Journey to Quiet Your Mind and Relax",
        "cov_theme": "Soft pastel watercolor soothing lavender field under misty mountains and serene gentle lake",
        "storybook_title": "THE BEAR WHO LEARNED TO BREATHE",
        "storybook_theme": "A big cuddly bear who learns to take deep breaths to feel calm when big storms arrive in the forest",
        "etsy_verse": "Peace begins the moment you choose not to allow another person or event to control your emotions.",
        "etsy_subject": "gentle blooming eucalyptus branch in soft pastel watercolor with soft dew drops",
        "gum_topic": "30 Napos Belső Nyugalom & Szorongásoldó Napló",
        "gum_title": "Mentális Egészség & Stresszoldó Megújulás Mestercsomag",
        "gum_trans": "Megnyugtatni a túlterhelt idegrendszert, elcsendesíteni a kavargó gondolatokat és békében aludni",
        "ffc_prod": "30 Napos Idegrendszer-Megnyugtató & Szorongásoldó Digitális Napló",
        "ffc_aud": "Szorongással, kiégéssel és stresszel küzdő emberek, akik belső nyugalomra vágynak",
        "ffc_trans": "Megtanulni 5 perc alatt leföldelni a stresszt és visszanyerni a belső kontrollt",
        "ffc_vehicle": "Szomatikus légzésgyakorlatok + Vezetett érzelemnapló + Művészi relaxációs kártyák",
        "reels_prod": "Szorongásoldás és Idegrendszer Nyugtató Gyakorlatok",
        "reels_cta_kw": "NYUGALOM",
        "gs_prod": "30 Napos Szorongásoldó & Belső Béke Mestercsomag",
        "gs_headline": "Engedd El A Belső Feszültséget És Találd Meg A Tartós Békességet",
        "gs_tagline": "Gyengéd, megnyugtató gyakorlatok, naplókérdések és művészi színezők a kiegyensúlyozott mindennapokért.",
        "gs_lead_magnet": "Ingyenes 5-Perces Vészhelyzeti Nyugtató Légzés & Napló Sablon azonnali letöltéssel.",
        "gs_features": "30 napos reflexió, szomatikus stresszoldó kártyák, 30 db részletgazdag relaxációs színező és hangulat-napló."
    },
    "🥗 Fogyás & Egészséges Táplálkozás (Weight Loss & Nutrition)": {
        "kdp_title": "CLEAN EATING & MEAL PREP MASTER PLANNER",
        "kdp_subtitle": "90-Day Healthy Nutrition, Recipe & Fitness Transformation Log",
        "kdp_theme": "Fresh organic farmer's market vegetables, colorful smoothie bowls, clean kitchen meal prep containers, water tracker drops, healthy balanced plate guide",
        "cov_title": "CLEAN EATING & MEAL PREP MASTER PLANNER",
        "cov_sub": "Transform Your Health with 90 Days of Mindful Nutrition",
        "cov_theme": "Vibrant fresh Mediterranean ingredients on rustic marble table with clean modern typography",
        "storybook_title": "THE LITTLE CHEF'S RAINBOW GARDEN",
        "storybook_theme": "Two cute animal friends discovering colorful fruits and vegetables in a magical sunny vegetable garden",
        "etsy_verse": "Let food be thy medicine and medicine be thy food. — Hippocrates",
        "etsy_subject": "vintage botanical illustration of fresh rosemary, lavender, and lemon with watercolor accents",
        "gum_topic": "90 Napos Tiszta Étkezés & Fenntartható Életmódváltás",
        "gum_title": "Egészséges Táplálkozás & Meal Prep Életmódváltó Mestercsomag",
        "gum_trans": "Önsanyargatás nélkül lefogyni, energikusnak lenni és fenntartható táplálkozási szokásokat kialakítani",
        "ffc_prod": "90 Napos Tiszta Étkezés & Meal Prep Életmódváltó Rendszer",
        "ffc_aud": "Életmódváltásra vágyó nők és férfiak, akik jojó-diéták nélkül akarnak fogyni",
        "ffc_trans": "Könnyedén megtervezni az egész heti egészséges ételeket 1 óra alatt és elérni az álomsúlyt",
        "ffc_vehicle": "Egyszerű heti menütervezési mátrix + Gyors bevásárlólisták + Éhség- és energianapló",
        "reels_prod": "Egészséges Fogyás és Tiszta Étkezési Trükkök",
        "reels_cta_kw": "EGÉSZSÉG",
        "gs_prod": "Tiszta Étkezés & Meal Prep Életmód Mestercsomag",
        "gs_headline": "Fogyj Le Éhezés Nélkül És Nyerj Vissza Határtalan Energiát",
        "gs_tagline": "Strukturált étkezéstervezők, 15 perces receptek és szokásépítő naplók az egészséges életért.",
        "gs_lead_magnet": "Ingyenes 7-Napos Tiszta Étkezési Minta Étrend és Bevásárlólista azonnali letöltéssel.",
        "gs_features": "90 napos étkezési napló, heti meal prep tervező, víz- és kalóriakövető lapok, valamint 50 gyors receptötlet."
    }
}


def sync_niche_preset_to_inputs(niche_key: str, force: bool = False):
    """
    Synchronizes all input fields across all workspaces when a niche is selected
    or when the user clicks 'Alapértelmezések Frissítése'.
    """
    niche_info = get_niche_prompt_context(niche_key)
    preset = NICHE_DEFAULTS.get(niche_key, {})
    
    n_name = niche_info.get("name_en", "General Niche")
    n_aud = niche_info.get("default_audience", "Célközönség az adott rétegpiacon")
    n_keys = niche_info.get("keywords", ["Success", "Growth", "Focus"])

    kdp_t = preset.get("kdp_title", f"{n_name.upper()} MASTER COLORING & JOURNAL")
    kdp_s = preset.get("kdp_subtitle", f"30 Inspiring {n_name} Prompts & Reflections")
    kdp_th = preset.get("kdp_theme", f"Inspiring scenes related to {n_name}, including {', '.join(n_keys[:4])}")
    
    cov_t = preset.get("cov_title", kdp_t)
    cov_s = preset.get("cov_sub", kdp_s)
    cov_th = preset.get("cov_theme", f"Beautiful high resolution artwork depicting {n_name}")

    story_t = preset.get("storybook_title", f"THE WONDERFUL ADVENTURE OF {n_name.upper()}")
    story_th = preset.get("storybook_theme", f"An inspiring illustrated story about courage and discovery in {n_name}")

    ffc_p = preset.get("ffc_prod", f"30 Napos {n_name} Digitális Mestercsomag")
    ffc_a = preset.get("ffc_aud", n_aud)
    ffc_tr = preset.get("ffc_trans", f"Megoldani a legfőbb akadályokat a(z) {n_name} területén és elérni a kívánt célt")
    ffc_v = preset.get("ffc_vehicle", "Strukturált napi gyakorlatok, ellenőrzőlisták és letisztult digitális lapok")
    reels_p = preset.get("reels_prod", ffc_p)
    reels_kw = preset.get("reels_cta_kw", n_keys[0].upper() if n_keys and len(n_keys[0]) <= 8 else "SIKER")

    gs_p = preset.get("gs_prod", ffc_p)
    gs_h = preset.get("gs_headline", f"Érd El A Kiválóságot És Fejlődj A(z) {n_name} Területén Napi 10 Percben")
    gs_tag = preset.get("gs_tagline", f"Prémium nyomtatható és digitális eszközök, sablonok és útmutatók a(z) {n_name} célközönségének.")
    gs_lm = preset.get("gs_lead_magnet", f"Ingyenes mintacsomag és kezdő munkafüzet a(z) {n_name} témájában.")
    gs_ft = preset.get("gs_features", f"30 napos vezetett napló, 30 db prémium nyomtatható sablon a(z) {n_name} témájában, azonnali PDF letöltéssel.")

    keys_to_update = {
        "kdp_ap_title": kdp_t,
        "kdp_ap_subtitle": kdp_s,
        "kdp_ap_theme": kdp_th,
        "ib_title_input": story_t,
        "ib_sub_input": f"For Kids · Inspiring {n_name} Story",
        "ib_theme_input": story_th,
        "cov_title_inp": cov_t,
        "cov_sub_inp": cov_s,
        "cov_theme_inp": cov_th,
        "verse_etsy": preset.get("etsy_verse", f"Inspiring wisdom about {n_name}."),
        "subject_etsy": preset.get("etsy_subject", f"artistic minimalist vector scene of {n_name}"),
        "audience_dev": n_aud,
        "theme_dev": f"30 Napos {n_name} Megújulás",
        "aud_b": n_aud,
        "thm_b": f"{n_name} témájú inspirációk",
        "e4_prod_title": f"{n_name} Printable Digital Pack",
        "ffc_prod_name_v2": ffc_p,
        "ffc_target_aud_v2": ffc_a,
        "ffc_main_trans_v2": ffc_tr,
        "ffc_vehicle_v2": ffc_v,
        "reels_prod_input": reels_p,
        "reels_cta_kw_input": reels_kw,
        "reels_target_input": ffc_a,
        "wh_prod_name_input": ffc_p,
        "gs_prod_name": gs_p,
        "gs_target_aud": ffc_a,
        "gs_headline": gs_h,
        "gs_tagline": gs_tag,
        "gs_lead_magnet": gs_lm,
        "gs_features": gs_ft,
        "em_lead_magnet": gs_lm,
        "em_paid_prod": ffc_p,
        "em_target_aud": ffc_a,
        "cal_prod_name": ffc_p,
        "cal_target_aud": ffc_a
    }

    for k, val in keys_to_update.items():
        if force or k not in st.session_state or not st.session_state[k]:
            st.session_state[k] = val

    st.session_state["last_synced_niche"] = niche_key


def get_niche_slug(niche_key: str = None) -> str:
    if not niche_key:
        niche_key = st.session_state.get("active_niche_choice", "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)")
    return "".join(c if c.isalnum() else "_" for c in niche_key)[:25]


def get_niche_field(field_name: str, niche_key: str = None) -> str:
    """Returns the tailored default value for any field for the active niche."""
    if not niche_key:
        niche_key = st.session_state.get("active_niche_choice", "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)")
        
    preset = NICHE_DEFAULTS.get(niche_key, {})
    if field_name in preset:
        return preset[field_name]
        
    # Fallback generation
    niche_info = get_niche_prompt_context(niche_key)
    n_name = niche_info.get("name_en", "General Niche")
    n_aud = niche_info.get("default_audience", "Célközönség")
    n_keys = niche_info.get("keywords", ["Growth", "Focus", "Success"])
    
    fallbacks = {
        "kdp_title": f"{n_name.upper()} MASTER COLORING & JOURNAL",
        "kdp_subtitle": f"30 Inspiring {n_name} Prompts & Reflections",
        "kdp_theme": f"Inspiring scenes related to {n_name}, including {', '.join(n_keys[:4])}",
        "cov_title": f"{n_name.upper()} MASTER COLORING & JOURNAL",
        "cov_sub": f"30 Inspiring {n_name} Prompts & Reflections",
        "cov_theme": f"Beautiful high resolution artwork depicting {n_name}",
        "storybook_title": f"THE WONDERFUL ADVENTURE OF {n_name.upper()}",
        "storybook_sub": f"For Kids · Inspiring {n_name} Story",
        "storybook_theme": f"An inspiring illustrated story about courage and discovery in {n_name}",
        "verse_etsy": f"Inspiring wisdom about {n_name}.",
        "subject_etsy": f"artistic minimalist vector scene of {n_name}",
        "gum_topic": f"30 Napos {n_name} Megújulás & Vezetett Napló",
        "gum_title": f"30 Napos {n_name} Digitális Mestercsomag",
        "gum_trans": f"Megoldani a legfőbb akadályokat a(z) {n_name} területén és elérni a kívánt célt",
        "ffc_prod": f"30 Napos {n_name} Digitális Mestercsomag",
        "ffc_aud": n_aud,
        "ffc_trans": f"Megoldani a legfőbb akadályokat a(z) {n_name} területén és elérni a kívánt célt",
        "ffc_vehicle": "Strukturált napi gyakorlatok, ellenőrzőlisták és letisztult digitális lapok",
        "reels_prod": f"30 Napos {n_name} Digitális Rendszer",
        "reels_cta_kw": n_keys[0].upper() if n_keys and len(n_keys[0]) <= 8 else "SIKER",
        "gs_prod": f"30 Napos {n_name} Digitális Mestercsomag",
        "gs_headline": f"Érd El A Kiválóságot És Fejlődj A(z) {n_name} Területén Napi 10 Percben",
        "gs_tagline": f"Prémium nyomtatható és digitális eszközök, sablonok és útmutatók a(z) {n_name} célközönségének.",
        "gs_lead_magnet": f"Ingyenes mintacsomag és kezdő munkafüzet a(z) {n_name} témájában.",
        "gs_features": f"30 napos vezetett napló, 30 db prémium nyomtatható sablon a(z) {n_name} témájában, azonnali PDF letöltéssel.",
        "em_lead_magnet": f"Ingyenes mintacsomag és kezdő munkafüzet a(z) {n_name} témájában.",
        "em_paid_prod": f"30 Napos {n_name} Digitális Mestercsomag",
        "cal_prod_name": f"30 Napos {n_name} Digitális Naptár"
    }
    return fallbacks.get(field_name, "")


def get_ffc_preset(niche_key: str, prod_type: str = "", language: str = "Magyar") -> dict:
    """Returns full tailored product details based on Niche + Product Type + Language."""
    is_en = "Angol" in language or "English" in language
    n_info = get_niche_prompt_context(niche_key)
    n_en = n_info.get("name_en", "Christian & Biblical Niche")
    n_aud = n_info.get("default_audience", "Keresztény édesanyák, alkotók és hívők")
    is_christian = "Keresztény" in niche_key or "Bibliai" in niche_key

    if is_en:
        if "Színező" in prod_type or "Coloring" in prod_type:
            return {
                "prod_name": f"30-Day {n_en} Coloring Book & Scripture Reflection Companion (Amazon KDP & Printable)",
                "target_aud": f"Parents, kids, and adults interested in {n_en} looking for creative mindfulness",
                "main_trans": f"Experience deep peace, reduce daily stress, and meditate on uplifting truths of {n_en} through 30 relaxing hand-drawn coloring scenes",
                "vehicle": "Bold clean outlines, inspirational scripture quotes, and structured mindful reflection",
                "extra_notes": f"Target English Amazon KDP {n_en} coloring book audience with 100% money-back guarantee."
            }
        elif "Falikép" in prod_type or "Clipart" in prod_type or "Wall" in prod_type:
            return {
                "prod_name": f"Museum-Quality {n_en} Wall Art & Clipart Master Bundle (300 DPI)",
                "target_aud": f"Etsy shoppers, digital planners, and home decorators interested in {n_en}",
                "main_trans": f"Easily decorate your home and craft best-selling print-on-demand products with timeless {n_en} artwork",
                "vehicle": "High-resolution 300 DPI watercolor vector assets, print-ready PDF formats, and commercial license",
                "extra_notes": "Highlight instant digital download and commercial POD usage rights."
            }
        elif "Tanfolyam" in prod_type or "Masterclass" in prod_type or "Course" in prod_type:
            return {
                "prod_name": f"{n_en} Digital Product & Faceless Funnel Masterclass",
                "target_aud": f"Content creators, freelancers, and entrepreneurs in {n_en}",
                "main_trans": f"Build and launch a 100% automated passive digital product funnel in {n_en} without showing your face on camera",
                "vehicle": "DFY Canva templates, ManyChat keyword automation, and Stripe zero-cost sales funnel",
                "extra_notes": "Include full step-by-step video curriculum and lifetime updates."
            }
        else:  # Devotional / Default
            return {
                "prod_name": f"30-Day {n_en} Guided Devotional & Mindfulness Journal (Printable & Digital)",
                "target_aud": f"Believers, creators, and individuals seeking daily quiet time and focus in {n_en}",
                "main_trans": f"Transform busy mornings into 10 minutes of profound spiritual clarity, release anxiety, and walk in peace through {n_en}",
                "vehicle": "Micro-reflections, timeless scripture anchors, and aesthetic printable daily worksheets",
                "extra_notes": "Include 30-day spiritual breakthrough guarantee."
            }
    else:  # Magyar
        if "Színező" in prod_type:
            return {
                "prod_name": f"30 Napos {n_en} Színezőkönyv & Művészi Lelki Útitárs (Nyomtatható & KDP)",
                "target_aud": f"Szülők, alkotók és felnőttek, akik kreatív elcsendesedésre vágynak a(z) {n_en} témájában",
                "main_trans": f"Megtalálni a napi belső békét, elengedni a stresszt és elmélyülni a(z) {n_en} világában 30 gyönyörű színező lapon keresztül",
                "vehicle": "Kristálytiszta kontúrok, inspiráló idézetek és vezetett relaxációs kérdések",
                "extra_notes": "Nyomdakész Amazon KDP és otthon nyomtatható PDF formátumban."
            }
        elif "Falikép" in prod_type or "Clipart" in prod_type:
            return {
                "prod_name": f"Prémium {n_en} Igés Falikép & Művészi Clipart Mestercsomag (300 DPI)",
                "target_aud": f"Etsy vásárlók, otthonszépítők és alkotók, akik igényes {n_en} grafikákat keresnek",
                "main_trans": f"Otthonodat és digitális termékeidet felemelő, prémium minőségű {n_en} illusztrációkkal díszíteni",
                "vehicle": "300 DPI nyomdakész PDF faliképek és átlátszó hátterű PNG clipartok kereskedelmi joggal",
                "extra_notes": "Azonnali digitális letöltés és Canva szerkeszthetőség."
            }
        elif "Tanfolyam" in prod_type or "Masterclass" in prod_type:
            return {
                "prod_name": f"{n_en} Digitális Termék & Faceless Funnel Mesterkurzus",
                "target_aud": f"Vállalkozók, alkotók és szabadúszók, akik passzív digitális jövedelmet akarnak építeni a(z) {n_en} területén",
                "main_trans": f"Felépíteni egy 100%-ban automatizált digitális termékrendszert arc nélküli videókkal a(z) {n_en} piacán",
                "vehicle": "DFY sablonok, ManyChat kulcsszó automatizáció és 0 Ft-os Stripe értékesítési tölcsér",
                "extra_notes": "Lépésről-lépésre videók és azonnal másolható sablonok."
            }
        else:  # Devotional / Default
            if is_christian:
                return {
                    "prod_name": "30 Napos Békesség & Fókusz Vezetett Lelki Napló (Printable & Digital)",
                    "target_aud": "Keresztény édesanyák, alkotók és hívők, akik lelki csendességre vágynak",
                    "main_trans": "Napi 10 perces vezetett csendességgel elengedni a szorongást, megtalálni a tartós belső békességet és Isten jelenlétében élni",
                    "vehicle": "Mikro-reflexiók, strukturált bibliai igemagyarázatok és letisztult nyomtatható naplólapok",
                    "extra_notes": "30 napos lelki megújulás garanciával."
                }
            else:
                return {
                    "prod_name": f"30 Napos {n_en} Vezetett Napi Munkafüzet & Transzformációs Napló",
                    "target_aud": n_aud,
                    "main_trans": f"Napi 10 perc fókuszált munkával leküzdeni az akadályokat és áttörést elérni a(z) {n_en} területén",
                    "vehicle": "Strukturált napi feladatok, ellenőrzőlisták és letisztult digitális lapok",
                    "extra_notes": "100% pénzvisszafizetési garanciával."
                }


def render_niche_status_bar(workspace_name: str = ""):
    """
    Renders an interactive status and quick-sync bar at the top of every workspace.
    """
    niche_k = st.session_state.get("active_niche_choice", "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)")
    niche_info = get_niche_prompt_context(niche_k)

    col_nb1, col_nb2 = st.columns([2.6, 1])
    with col_nb1:
        st.markdown(f"""
        <div style='background: rgba(16, 185, 129, 0.08); border-left: 4px solid #10b981; border-radius: 8px; padding: 8px 14px; margin-bottom: 14px;'>
            <span style='color:#34d399; font-weight:700;'>🎯 Aktív Célpiac:</span> <b>{niche_k}</b>
            <span style='color:#94a3b8; font-size:0.82rem; margin-left:8px;'>({niche_info.get('group', 'General')})</span><br>
            <span style='color:#cbd5e1; font-size:0.84rem;'>👥 <i>Célközönség: {niche_info.get('default_audience', '')}</i></span>
        </div>
        """, unsafe_allow_html=True)
    with col_nb2:
        if st.button("🔄 Téma Sablonok Újratöltése", key=f"btn_reload_niche_{workspace_name}", use_container_width=True, help="Kattints ide, ha vissza szeretnéd tölteni az ehhez a célpiachoz tartozó gyári sablonokat a mezőkbe!"):
            sync_niche_preset_to_inputs(niche_k, force=True)
            st.toast(f"✅ {niche_k} sablonok sikeresen betöltve!", icon="🎯")
            st.rerun()


def render_style_selector(widget_prefix: str = "ws") -> str:
    """
    Renders an inline visual art style selector component directly on any workspace.
    Synchronizes across session state and displays an immediate live preview of the active style prompt.
    """
    style_keys = list(STYLE_PRESETS.keys())
    saved_style = st.session_state.get("active_style_choice", style_keys[0])
    if saved_style not in style_keys:
        saved_style = style_keys[0]

    with st.container():
        col_st1, col_st2 = st.columns([1.1, 1.25])
        with col_st1:
            chosen_style = st.selectbox(
                "🎨 Vizuális Márka Stílus (Art Style):",
                options=style_keys,
                index=style_keys.index(saved_style),
                key=f"{widget_prefix}_style_select",
                help="Közvetlenül ezen a munkaterületen módosíthatod a generált képek és promptok stílusát!"
            )
            st.session_state["active_style_choice"] = chosen_style

        custom_text = st.session_state.get("custom_style_text", "Custom vector line art style, bold outlines, vibrant colors")
        if "Egyedi Stílus" in chosen_style:
            with col_st2:
                custom_text = st.text_area(
                    "Egyedi stílusleírás angolul:",
                    value=custom_text,
                    height=70,
                    key=f"{widget_prefix}_custom_style_txt"
                )
                st.session_state["custom_style_text"] = custom_text
            res_prompt = custom_text.strip()
        else:
            res_prompt = STYLE_PRESETS.get(chosen_style, "")
            with col_st2:
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.9)); border: 1px solid #3b82f6; border-radius: 10px; padding: 10px 14px; margin-top: 4px;'>
                    <div style='color: #60a5fa; font-size: 0.82rem; font-weight: 700;'>✨ Aktív Vizuális Stílus Prompt:</div>
                    <div style='color: #f1f5f9; font-size: 0.85rem; margin-top: 3px; font-style: italic; line-height: 1.3;'>"{res_prompt}"</div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("---")
    return res_prompt



def render_audhd_focus_dashboard():
    """
    Renders the persistent AuDHD 2-Hour Focus, Timeboxing, Subtask Checklist,
    and Color-Coded Performance Calendar Dashboard.
    """
    # 1. State initializations
    if "timer_running" not in st.session_state:
        st.session_state["timer_running"] = False
    if "timer_start_time" not in st.session_state:
        st.session_state["timer_start_time"] = None
    if "timer_elapsed_seconds" not in st.session_state:
        st.session_state["timer_elapsed_seconds"] = 0
    if "celebrated_days" not in st.session_state:
        st.session_state["celebrated_days"] = []

    weekday_idx = datetime.datetime.now().weekday()
    day_keys = ["Hétfő", "Kedd", "Szerda", "Csütörtök", "Péntek"]
    default_day = day_keys[min(weekday_idx, 4)]
    if "active_focus_day" not in st.session_state:
        st.session_state["active_focus_day"] = default_day

    # Calculate live elapsed seconds
    current_sec = get_current_timer_seconds()
    current_mins = round(current_sec / 60, 1)
    is_overrun = current_mins > 120
    is_running = st.session_state.get("timer_running", False)

    # Top summary badge in header
    timer_hms = format_seconds_to_hms(current_sec)
    status_emoji = "🔥" if is_running else "⏸️"
    status_color_tag = "🔴 TÚLLÉPVE" if is_overrun else "🟢 120p KERETEN BELÜL"

    with st.expander(f"⏱️ AuDHD 2-Órás Fókusz & Naptár Műszerfal ({status_emoji} {timer_hms} · {status_color_tag})", expanded=True):
        tab_focus, tab_history = st.tabs([
            "⏱️ 1. Napi 2-Órás Fókusz & Élő Stopper",
            "📅 2. Teljesítmény Naptár & Történet (Kártya Nézet)"
        ])

        with tab_focus:
            col_f_left, col_f_right = st.columns([1.1, 1], gap="large")

            with col_f_left:
                st.markdown("<div class='step-label'>📅 Napi Fókusz Tervező & Csekklista</div>", unsafe_allow_html=True)
                
                selected_day = st.selectbox(
                    "Válassz Munkanapot (Hétfő – Péntek):",
                    options=day_keys,
                    index=day_keys.index(st.session_state.get("active_focus_day", default_day)),
                    key="sel_focus_day",
                    help="Válaszd ki az aznapi fókuszblokkot! A feladatlista automatikusan betöltődik."
                )
                st.session_state["active_focus_day"] = selected_day
                plan_data = AUDHD_DAY_PLANS.get(selected_day, AUDHD_DAY_PLANS["Hétfő"])

                st.markdown(f"""
                <div style='background:#1a2338; border:1px solid #28364e; border-radius:10px; padding:10px 14px; margin-bottom:12px;'>
                    <strong style='color:#34d399; font-size:0.98rem;'>{plan_data['title']}</strong><br>
                    <span style='color:#cbd5e1; font-size:0.85rem;'>{plan_data['description']}</span>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("**📋 Napi 2-Órás Részfeladatok (Csekklista):**")
                completed_tasks = []
                all_checked = True

                for idx, t_text in enumerate(plan_data["tasks"]):
                    chk_key = f"chk_task_{selected_day}_{idx}"
                    is_chk = st.checkbox(t_text, key=chk_key)
                    if is_chk:
                        completed_tasks.append(t_text)
                    else:
                        all_checked = False

                # Trigger balloons on full completion
                if all_checked and len(plan_data["tasks"]) > 0:
                    if selected_day not in st.session_state["celebrated_days"]:
                        st.balloons()
                        st.session_state["celebrated_days"].append(selected_day)
                    st.success(f"🎉 **Gratulálunk!** A mai ({selected_day}) mind az 5 alfeladatot sikeresen befejezted!")

                focus_note = st.text_input(
                    "📝 Napi megjegyzés / reflexió (opcionális):",
                    placeholder="Pl.: Nagyon jó volt a reggeli flow, a feladatok azonnal sikerültek!",
                    key=f"note_focus_{selected_day}"
                )

            with col_f_right:
                st.markdown("<div class='step-label'>⏱️ Élő 2-Órás Stopper & Időkorlát</div>", unsafe_allow_html=True)

                is_running_js = "true" if is_running else "false"
                base_seconds = current_sec

                live_timer_html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@600;700;800&family=JetBrains+Mono:wght@700;800&display=swap');
    * {{
        box-sizing: border-box;
        margin: 0;
        padding: 0;
        font-family: 'Plus Jakarta Sans', sans-serif;
    }}
    body {{
        background: transparent;
        color: #f1f5f9;
        overflow: hidden;
    }}
    .timer-clock {{
        font-family: 'JetBrains Mono', 'Consolas', monospace;
        font-size: 2.5rem;
        font-weight: 800;
        letter-spacing: 2px;
        color: #34d399;
        text-align: center;
        background: #0d121f;
        padding: 10px 14px;
        border-radius: 12px;
        border: 1.5px solid #1e293b;
        margin: 2px 0 8px 0;
        transition: color 0.3s ease, border-color 0.3s ease;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    }}
    .timer-clock-red {{
        color: #f87171 !important;
        border-color: #ef444466 !important;
        box-shadow: 0 4px 14px rgba(239, 68, 68, 0.25) !important;
    }}
    .badge-center {{
        display: flex;
        justify-content: center;
        margin-bottom: 8px;
    }}
    .badge-green {{
        background: #064e3b;
        color: #34d399;
        border: 1px solid #10b981;
        font-weight: 700;
        padding: 4px 14px;
        border-radius: 20px;
        font-size: 0.82rem;
        display: inline-block;
        letter-spacing: 0.02em;
    }}
    .badge-red {{
        background: #7f1d1d;
        color: #fca5a5;
        border: 1px solid #ef4444;
        font-weight: 700;
        padding: 4px 14px;
        border-radius: 20px;
        font-size: 0.82rem;
        display: inline-block;
        letter-spacing: 0.02em;
    }}
    .progress-track {{
        width: 100%;
        height: 8px;
        background-color: #1e293b;
        border-radius: 6px;
        overflow: hidden;
        margin-top: 2px;
    }}
    .progress-fill {{
        height: 100%;
        width: 0%;
        background-color: #10b981;
        transition: width 0.25s linear, background-color 0.3s ease;
        border-radius: 6px;
    }}
</style>
</head>
<body>
    <div id="clock-display" class="timer-clock">00:00:00</div>
    <div class="badge-center">
        <span id="badge-display" class="badge-green">🟢 SIKERES IDŐKERET (0.0 / 120 perc)</span>
    </div>
    <div class="progress-track">
        <div id="progress-fill" class="progress-fill"></div>
    </div>

<script>
    const isRunning = {is_running_js};
    const baseSeconds = {base_seconds};
    const startEpoch = Date.now();

    function formatHMS(totalSec) {{
        const s = Math.max(0, Math.floor(totalSec));
        const hours = Math.floor(s / 3600);
        const mins = Math.floor((s % 3600) / 60);
        const secs = s % 60;
        return (hours < 10 ? '0' + hours : hours) + ':' +
               (mins < 10 ? '0' + mins : mins) + ':' +
               (secs < 10 ? '0' + secs : secs);
    }}

    function update() {{
        let sec = baseSeconds;
        if (isRunning) {{
            sec += (Date.now() - startEpoch) / 1000;
        }}
        const mins = sec / 60;
        const isOver = mins > 120;

        const clockEl = document.getElementById("clock-display");
        const badgeEl = document.getElementById("badge-display");
        const fillEl = document.getElementById("progress-fill");

        if (clockEl) {{
            clockEl.textContent = formatHMS(sec);
            if (isOver) {{
                clockEl.className = "timer-clock timer-clock-red";
            }} else {{
                clockEl.className = "timer-clock";
            }}
        }}

        if (badgeEl) {{
            if (isOver) {{
                const overM = (mins - 120).toFixed(1);
                badgeEl.className = "badge-red";
                badgeEl.textContent = "🔴 IDŐKERET TÚLLÉPVE (+" + overM + " perc túllépés / " + mins.toFixed(1) + " perc)";
            }} else {{
                const remM = (120 - mins).toFixed(1);
                badgeEl.className = "badge-green";
                badgeEl.textContent = "🟢 SIKERES IDŐKERET (" + mins.toFixed(1) + " / 120 perc · Még " + remM + " perc hátra)";
            }}
        }}

        if (fillEl) {{
            const pct = Math.min(100, (sec / 7200) * 100);
            fillEl.style.width = pct.toFixed(1) + "%";
            fillEl.style.backgroundColor = isOver ? "#ef4444" : "#10b981";
        }}
    }}

    update();
    if (isRunning) {{
        setInterval(update, 250);
    }}
</script>
</body>
</html>
"""
                components.html(live_timer_html, height=135)

                # Timer Action Buttons
                col_b1, col_b2, col_b3 = st.columns(3)
                with col_b1:
                    if not is_running:
                        if st.button("▶️ Start / Folytatás", key="btn_timer_start", use_container_width=True):
                            st.session_state["timer_running"] = True
                            st.session_state["timer_start_time"] = time.time()
                            st.rerun()
                    else:
                        st.button("🟢 Fut az óra...", disabled=True, key="btn_timer_running_dis", use_container_width=True)
                
                with col_b2:
                    if is_running:
                        if st.button("⏸️ Szünet (Pause)", key="btn_timer_pause", use_container_width=True):
                            if st.session_state.get("timer_start_time"):
                                st.session_state["timer_elapsed_seconds"] += (time.time() - st.session_state["timer_start_time"])
                            st.session_state["timer_running"] = False
                            st.session_state["timer_start_time"] = None
                            st.rerun()
                    else:
                        st.button("⏸️ Szünet", disabled=True, key="btn_timer_paused_dis", use_container_width=True)

                with col_b3:
                    if st.button("🔄 Reset", key="btn_timer_reset", use_container_width=True, help="Nullázza az időmérőt."):
                        st.session_state["timer_running"] = False
                        st.session_state["timer_start_time"] = None
                        st.session_state["timer_elapsed_seconds"] = 0
                        st.rerun()

                st.markdown("---")
                
                # Save and complete block
                btn_save_block = st.button("✅ 2-Órás Blokk Lezárása & Mentése", key="btn_save_focus_block", use_container_width=True)
                if btn_save_block:
                    # Finalize current elapsed time
                    final_sec = current_sec
                    final_mins = round(final_sec / 60, 1)
                    final_status = "GREEN" if final_mins <= 120 else "RED"

                    entry = {
                        "id": str(uuid.uuid4())[:8],
                        "date": datetime.date.today().isoformat(),
                        "timestamp": datetime.datetime.now().strftime("%H:%M:%S"),
                        "day_name": selected_day,
                        "task_name": plan_data["title"],
                        "target_minutes": 120,
                        "elapsed_minutes": final_mins,
                        "elapsed_seconds": int(final_sec),
                        "elapsed_formatted": format_seconds_to_hms(final_sec),
                        "status": final_status,
                        "completed_subtasks": completed_tasks,
                        "total_subtasks_count": len(plan_data["tasks"]),
                        "notes": focus_note
                    }

                    ok = save_time_log_entry(entry)
                    if ok:
                        # Reset timer
                        st.session_state["timer_running"] = False
                        st.session_state["timer_start_time"] = None
                        st.session_state["timer_elapsed_seconds"] = 0
                        
                        if final_status == "GREEN":
                            st.balloons()
                            st.success(f"🟢 **Szuper munka!** A(z) {selected_day}i 2-órás fókuszblokkot {final_mins} perc alatt sikeresen teljesítetted és elmentetted!")
                        else:
                            st.warning(f"🔴 **Blokk elmentve!** Időtartam: {final_mins} perc (120 perces keret túllépve).")
                        st.rerun()
                    else:
                        st.error("Hiba történt a blokk mentésekor a time_log.json fájlba.")

        with tab_history:
            st.markdown("<div class='step-label'>📅 Korábbi Fókuszblokkok & Teljesítmény Naptár</div>", unsafe_allow_html=True)
            logs = load_time_logs()

            if logs:
                # Summary metrics
                total_blocks = len(logs)
                green_blocks = sum(1 for l in logs if l.get("status") == "GREEN")
                green_pct = round((green_blocks / total_blocks) * 100, 1) if total_blocks > 0 else 0
                avg_mins = round(sum(l.get("elapsed_minutes", 0) for l in logs) / total_blocks, 1) if total_blocks > 0 else 0
                total_hours = round(sum(l.get("elapsed_minutes", 0) for l in logs) / 60, 1)

                col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                with col_m1:
                    st.metric("📊 Összes Blokk", f"{total_blocks} db")
                with col_m2:
                    st.metric("🟢 Zöld Sikeresség", f"{green_pct}%", delta=f"{green_blocks}/{total_blocks} db")
                with col_m3:
                    st.metric("⏱️ Átlagos Időtartam", f"{avg_mins} perc", delta=f"Cél: 120 perc")
                with col_m4:
                    st.metric("🏆 Fókusz Idő", f"{total_hours} óra")

                st.markdown("---")

                # Render history cards (sorted by date/timestamp descending)
                sorted_logs = sorted(logs, key=lambda x: (x.get("date", ""), x.get("timestamp", "")), reverse=True)

                for item in sorted_logs:
                    i_status = item.get("status", "GREEN")
                    i_mins = item.get("elapsed_minutes", 0)
                    i_day = item.get("day_name", "Ismeretlen nap")
                    i_date = item.get("date", "")
                    i_time = item.get("timestamp", "")
                    i_task = item.get("task_name", "Fókuszblokk")
                    i_subtasks = item.get("completed_subtasks", [])
                    i_total_sub = item.get("total_subtasks_count", 5)
                    i_notes = item.get("notes", "")

                    if i_status == "GREEN":
                        card_html = f"""
                        <div class='history-card-green'>
                            <div style='display:flex; justify-content:space-between; align-items:center;'>
                                <span style='font-weight:700; color:#34d399; font-size:1.02rem;'>📅 {i_date} ({i_day}) · {i_time}</span>
                                <span class='badge-green'>🟢 {i_mins} PERC (SIKERES &lt;= 120p)</span>
                            </div>
                            <h4 style='margin:6px 0 2px 0; color:#f8fafc; font-size:1.05rem;'>{i_task}</h4>
                            <div style='font-size:0.83rem; color:#a7f3d0;'>✅ Elvégzett alfeladatok: {len(i_subtasks)} / {i_total_sub} db</div>
                            {f"<div style='font-size:0.82rem; color:#cbd5e1; margin-top:4px;'><em>📝 {i_notes}</em></div>" if i_notes else ""}
                        </div>
                        """
                    else:
                        over_m = round(i_mins - 120, 1)
                        card_html = f"""
                        <div class='history-card-red'>
                            <div style='display:flex; justify-content:space-between; align-items:center;'>
                                <span style='font-weight:700; color:#f87171; font-size:1.02rem;'>📅 {i_date} ({i_day}) · {i_time}</span>
                                <span class='badge-red'>🔴 {i_mins} PERC (+{over_m}p TÚLLÉPVE)</span>
                            </div>
                            <h4 style='margin:6px 0 2px 0; color:#f8fafc; font-size:1.05rem;'>{i_task}</h4>
                            <div style='font-size:0.83rem; color:#fca5a5;'>⚠️ Elvégzett alfeladatok: {len(i_subtasks)} / {i_total_sub} db</div>
                            {f"<div style='font-size:0.82rem; color:#cbd5e1; margin-top:4px;'><em>📝 {i_notes}</em></div>" if i_notes else ""}
                        </div>
                        """
                    st.markdown(card_html, unsafe_allow_html=True)

                st.markdown("---")
                col_exp1, col_exp2 = st.columns([1, 1])
                with col_exp1:
                    st.download_button(
                        label="⬇️ Időnapló Letöltése (.json)",
                        data=json.dumps(logs, ensure_ascii=False, indent=2).encode("utf-8"),
                        file_name=f"AuDHD_Time_Log_{datetime.date.today().strftime('%Y%m%d')}.json",
                        mime="application/json",
                        key="dl_time_log_json",
                        use_container_width=True
                    )
                with col_exp2:
                    if st.button("🗑️ Időnapló Történet Törlése", key="btn_clear_time_logs", use_container_width=True):
                        clear_time_logs()
                        st.success("Történet törölve!")
                        st.rerun()
            else:
                st.markdown("""
                <div style='text-align:center; padding: 40px 20px; border: 2px dashed #333f56; border-radius: 14px; background: #161f30; color: #94a3b8;'>
                    <div style='font-size: 2.4rem;'>📅</div>
                    <div style='margin-top: 8px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Még nincsenek elmentett fókuszblokkok.<br>Indítsd el a stoppert az 1. fülön, és zárd le a 2-órás blokkot!</div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("---")


# ─────────────────────────────────────────────────────────
# SIDEBAR NAVIGATION HUB
# ─────────────────────────────────────────────────────────

st.sidebar.markdown("""
<div style='text-align: center; padding: 6px 0 16px 0;'>
    <div style='font-size: 2.2rem; margin-bottom: 2px;'>✝️</div>
    <h2 style='margin: 0; font-size: 1.25rem; font-weight: 800; color: #34d399;'>Keresztény Alkotóműhely</h2>
    <div style='font-size: 0.8rem; color: #94a3b8; font-weight: 500;'>AuDHD Digitális Termék Rendszer & 22 Niche Hub</div>
</div>
""", unsafe_allow_html=True)

# ── 22 High-Demand Niche Selector in Sidebar ──
niche_keys = list(NICHE_CATEGORIES.keys())
saved_niche_key = cfg.get("selected_niche", niche_keys[0])
if saved_niche_key not in niche_keys:
    saved_niche_key = niche_keys[0]

chosen_niche = st.sidebar.selectbox(
    "🎯 Cél Niche Kategória (22 Piac):",
    options=niche_keys,
    index=niche_keys.index(saved_niche_key),
    key="sidebar_niche_select",
    help="Válassz a 22 magas keresletű niche közül! Minden munkaterület, prompt és marketing eszköz ehhez a célpiachoz igazodik."
)
st.session_state["active_niche_choice"] = chosen_niche
active_niche_info = get_niche_prompt_context(chosen_niche)

# Automatically sync inputs across all workspaces when niche is selected or changed
if st.session_state.get("last_synced_niche") != chosen_niche:
    sync_niche_preset_to_inputs(chosen_niche, force=True)
    st.session_state["last_synced_niche"] = chosen_niche

st.sidebar.markdown(f"""
<div style='background:#161f30; border:1px solid #28364e; border-radius:10px; padding:8px 12px; margin-bottom:14px; font-size:0.82rem; color:#94a3b8;'>
    <strong style='color:#34d399;'>🏷️ Csoport:</strong> {active_niche_info.get('group', 'General')}<br>
    <strong style='color:#cbd5e1;'>🎯 Piac:</strong> {active_niche_info.get('name_en', '')}
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div class='nav-title-box'>
    <span class='nav-title-text'>🧭 FŐMENÜ & MUNKATERÜLETEK</span>
</div>
""", unsafe_allow_html=True)

# ── Reorganized Hierarchical Sidebar Navigation ──
ai_workspace_choice = st.sidebar.selectbox(
        "AI Generáló Munkaterület Kiválasztása:",
        [
            "🎨 Amazon KDP Színező & PDF Összeállító",
            "📖 Amazon KDP Illusztrált Mesekönyv Műhely",
            "🎨 Amazon KDP Borító & Gerinc Mester",
            "💡 30 Téma & Ötletgeneráló Műhely",
            "🎨 Etsy Faliképek & Clipartok",
            "✍️ Gumroad Áhítatok & SEO",
            "📓 NotebookLM RAG Központ & Review Mining",
            "💰 2026-os Magyar Átalányadó & Pénzügyi Tervező",
            "📷 AI Vision Multimodális Lab",
            "🖼️ Gemini Képbegyűjtő & PDF Központ",
            "🚀 FFC Marketing, Copywriting & Google Sites Stúdió",
            "⚙️ Rendszerbeállítások & API Kulcsok",
            "🎁 DFY Canva Sablonok & Bónuszok"
        ],
        index=0,
        key="ai_workspace_select",
        label_visibility="collapsed"
    )
menu_choice = ai_workspace_choice

st.sidebar.markdown("---")

# Quick status widget in sidebar
st.sidebar.markdown("##### ⚡ Rendszer & AI Állapot")
summary = km.get_summary()

prov_info = []
if summary.get("has_groq"):
    prov_info.append("🚀 Groq (Elsődleges)")
if summary.get("has_openrouter"):
    prov_info.append("🌐 OpenRouter (Másodlagos)")
if summary.get("has_gemini"):
    prov_info.append("💎 Fizetős Gemini (Tartalék)")

if prov_info:
    st.sidebar.markdown(f"📝 **Szöveg AI:** 🟢 {' ➔ '.join(prov_info)}")
else:
    st.sidebar.markdown("📝 **Szöveg AI:** 🛡️ Beépített Sablon Motor")

sa_info = get_service_account_info()
if sa_info:
    st.sidebar.markdown("📁 **Drive:** ☁️ Cloud API Csatlakoztatva")
elif os.path.exists(drive_root_input):
    st.sidebar.markdown("📁 **Drive csatlakoztatva (Helyi)**")
else:
    st.sidebar.markdown("🟡 *Drive mappa helyi fallback-ben*")


if st.sidebar.button("🔄 Kulcsok Resetelése", use_container_width=True, help="Visszaállítja az összes korlátozású kulcsot ACTIVE státuszba."):
    km.reset_all_keys()
    st.sidebar.success("🟢 Kulcsok resetelve!")
    st.rerun()

st.sidebar.caption("Keresztény Digitális Prompt Asszisztens · v3.5")


# ─────────────────────────────────────────────────────────
# MAIN WORKSPACE ROUTING
# ─────────────────────────────────────────────────────────

# Render the persistent AuDHD 2-Hour Focus & Calendar Dashboard on top of main area
render_audhd_focus_dashboard()


# ══════════════════════════════════════════════════════════
# WORKSPACE 1: AMAZON KDP SZÍNEZŐK & NYOMDAKÉSZ PDF
# ══════════════════════════════════════════════════════════

if "Színező" in menu_choice or "KDP Színező" in menu_choice:
    st.markdown("<div class='path-badge'>🎨 Amazon KDP Színező Műhely</div>", unsafe_allow_html=True)
    st.markdown("### 🎨 Bibliai Színezők & Nyomdakész KDP Belsők")
    st.caption("Hozz létre kristálytiszta Gemini színező promptokat, gyűjtsd be a képeket a Google Drive-ról, és fűzd össze nyomdakész PDF könyvbelsővé.")

    active_art_style_prompt = render_style_selector("ws_kdp")
    render_niche_status_bar("kdp")

    tab_kdp_autopilot, tab_kdp_single = st.tabs([
        "⚡ 1. Gemini Képasszisztens & KDP Színező PDF Összeállító",
        "✍️ 2. Egyedi Színező Prompt Készítő (1 Jelenet)"
    ])

    # ── TAB 1: GEMINI KÉPASSZISZTENS & KDP PDF ÖSSZEÁLLÍTÓ ──
    with tab_kdp_autopilot:
        # Project Management & Persistence Bar
        render_book_project_manager_bar(project_type="kdp_coloring", widget_prefix="kdp_ap")

        st.markdown(
            """
            <div style='background: linear-gradient(135deg, rgba(30,58,138,0.2) 0%, rgba(16,185,129,0.15) 100%);
                        border: 1.5px solid #3b82f6; border-radius: 14px; padding: 18px 22px; margin-bottom: 20px;'>
                <div style='display: flex; align-items: center; justify-content: space-between; margin-bottom: 6px;'>
                    <h3 style='margin: 0; color: #60A5FA; font-size: 1.35rem;'>⚡ Kétablakos Gemini Képgeneráló & Drive PDF Összeállító</h3>
                    <span style='background: #1E3A8A; color: #93C5FD; font-size: 0.75rem; font-weight: 700; padding: 3px 12px; border-radius: 20px; border: 1px solid #3B82F6;'>
                        Side-by-Side Mód · KDP Nyomdakész
                    </span>
                </div>
                <p style='margin: 0; color: #CBD5E1; font-size: 0.88rem; line-height: 1.5;'>
                    <b>💡 Kétablakos Munkafolyamat:</b> Nyisd meg a képernyőd bal felén ezt az alkalmazást, a jobb felén pedig a <b>gemini.google.com</b> felületet!<br>
                    1️⃣ <b>Bal ablak:</b> Válassz formátumot/képarányt, add meg a témát és generáld le a vázlatot.<br>
                    2️⃣ <b>Jobb ablak:</b> Másold a Gemini Gem Mester Utasítást és a sorszámozott képpromptokat 1 kattintással a Gemini-be, majd mentsd le a képeket a Drive projektmappába (<code>01.png</code>, <code>02.png</code>...).<br>
                    3️⃣ <b>Összeállítás:</b> Kattints a <b>"Képek Begyűjtése"</b> gombra, majd a <b>"Nyomdakész KDP Belső PDF Összeállítása"</b> gombra!
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        col_ap_left, col_ap_right = st.columns([1.15, 1], gap="large")

        with col_ap_left:
            st.markdown("<div class='step-label'>Lépés 1 — Könyv Alapadatok</div>", unsafe_allow_html=True)
            n_slug_kdp = get_niche_slug(chosen_niche)
            book_title_ap = st.text_input(
                "📖 Könyv Főcíme (Nagybetűkkel jelenik meg a borítón és címoldalon):",
                value=st.session_state.get(f"kdp_ap_title_{n_slug_kdp}", get_niche_field("kdp_title", chosen_niche)),
                key=f"kdp_ap_title_{n_slug_kdp}",
                help="Pl. NOAH'S ARK BIBLE ADVENTURES, HEROES OF FAITH COLORING BOOK"
            )
            book_sub_ap = st.text_input(
                "🏷️ Alcím / Célközönség leírás (opcionális):",
                value=st.session_state.get(f"kdp_ap_subtitle_{n_slug_kdp}", get_niche_field("kdp_subtitle", chosen_niche)),
                key=f"kdp_ap_subtitle_{n_slug_kdp}",
                help="Pl. For Kids Ages 4-8 · 30 Inspiring Bible Stories"
            )
            book_theme_ap = st.text_area(
                "🌟 Könyv Témája & Történetíve (miről szóljanak a jelenetek?):",
                value=st.session_state.get(f"kdp_ap_theme_{n_slug_kdp}", get_niche_field("kdp_theme", chosen_niche)),
                height=110,
                key=f"kdp_ap_theme_{n_slug_kdp}",
                help="Részletezd a történetet, főbb szereplőket, hogy a Gemini változatos és összefüggő jeleneteket alkosson."
            )

        with col_ap_right:
            st.markdown("<div class='step-label'>Lépés 2 — Kiadás, Méretarány & Paraméterek</div>", unsafe_allow_html=True)
            
            # Format & Aspect Ratio Selector
            trim_size_ap, aspect_ratio_ap, info_ap = render_kdp_format_and_aspect_selector(
                widget_prefix="kdp_ap",
                default_format=st.session_state.get("kdp_autopilot_trim_size", "8.5x11")
            )

            kdp_edition_ap = st.radio(
                "👥 Kiadás Módja & Stíluskövetelmény:",
                [
                    "🧒 Gyermek Kiadás (Egyszerű, vastag fekete vonalak, cuki figurák, tiszta fehér háttér)",
                    "🧘 Felnőtt Kiadás (Részletgazdag mandala, zentangle, klasszikus vésnök vonalrajz)"
                ],
                index=1 if st.session_state.get("kdp_autopilot_is_adult", False) else 0,
                key="kdp_ap_edition"
            )
            is_adult_ap = "Felnőtt Kiadás" in kdp_edition_ap

            page_count_ap = st.select_slider(
                "🔢 Színező Oldalak Száma:",
                options=[5, 10, 15, 20, 25, 30],
                value=10,
                key="kdp_ap_page_count",
                help="Gyors teszthez 5-10 oldal, teljes Amazon KDP kötethez 20-30 oldal javasolt."
            )

            st.markdown("<div style='margin-top: 14px;'></div>", unsafe_allow_html=True)
            btn_start_manifest_gen = st.button(
                "✨ 1. LÉPÉS: Könyvvázlat & Promptok Generálása (AI)",
                key="btn_start_kdp_manifest_gen",
                use_container_width=True
            )

        # ── Step 1 Execution: Generate Manifest & Prompts ──
        if btn_start_manifest_gen:
            if not book_title_ap.strip() or not book_theme_ap.strip():
                st.error("⚠️ Kérlek add meg a könyv címét és a történet témáját a vázlat készítéséhez!")
            else:
                with st.spinner(f"AI ({text_model}) megtervezi a(z) {page_count_ap} oldalas könyvvázlatot ({trim_size_ap} · {aspect_ratio_ap}) KJV igékkel és színpalettával..."):
                    p_manifest = build_kdp_autopilot_manifest_prompt(
                        book_title=book_title_ap.strip(),
                        theme=book_theme_ap.strip(),
                        page_count=page_count_ap,
                        target_audience="Adult" if is_adult_ap else "Children",
                        style_name=active_art_style_prompt,
                        image_model="gemini-web",
                        trim_size=trim_size_ap,
                        aspect_ratio=aspect_ratio_ap
                    )

                    ok_manifest, raw_manifest_text = km.generate_text_with_fallback(
                        prompt=p_manifest,
                        model_name=text_model
                    )

                    if not ok_manifest or not raw_manifest_text:
                        st.error(f"❌ Nem sikerült legenerálni a könyvvázlatot: {raw_manifest_text}")
                    else:
                        parsed_scenes = parse_kdp_autopilot_manifest_json(raw_manifest_text)
                        if not parsed_scenes:
                            st.error("❌ Az AI által visszaadott vázlat JSON formátuma nem volt értelmezhető.")
                        else:
                            drive_kdp_dir = resolve_drive_folder("kdp")
                            timestamp_run = time.strftime("%Y%m%d_%H%M%S")
                            sanitized_book_title = sanitize_filename(book_title_ap)
                            book_project_dir = os.path.join(drive_kdp_dir, f"ColoringBook_{sanitized_book_title}_{timestamp_run}")
                            os.makedirs(book_project_dir, exist_ok=True)

                            # Save prompt manifest as text file to Drive folder
                            prompt_file_content = f"AMAZON KDP COLORING BOOK PROMPT MANIFEST\n"
                            prompt_file_content += f"Title: {book_title_ap.strip()}\nSubtitle: {book_sub_ap.strip()}\n"
                            prompt_file_content += f"Format: {trim_size_ap} | Aspect Ratio: {aspect_ratio_ap}\n"
                            prompt_file_content += f"Edition: {'Adult' if is_adult_ap else 'Children'} | Pages: {len(parsed_scenes)}\n\n"
                            prompt_file_content += "="*60 + "\n\n"
                            for s in parsed_scenes:
                                prompt_file_content += f"PAGE {s.get('page_number', 1)}: {s.get('title', '')} ({s.get('title_hu', '')})\n"
                                prompt_file_content += f"Scripture: {s.get('scripture_reference', '')} - {s.get('scripture_text', '')}\n"
                                prompt_file_content += f"Palette: {', '.join(s.get('color_suggestions', []))}\n"
                                prompt_file_content += f"Gemini Prompt:\n{s.get('visual_prompt', '')}\n"
                                prompt_file_content += f"Reflection: {s.get('reflection_thought', '')}\n\n"
                                prompt_file_content += "-"*40 + "\n\n"

                            txt_save_path = os.path.join(book_project_dir, f"Prompts_{sanitized_book_title}.txt")
                            try:
                                with open(txt_save_path, "w", encoding="utf-8") as f_pt:
                                    f_pt.write(prompt_file_content)
                            except Exception:
                                pass

                            st.session_state["kdp_autopilot_scenes"] = parsed_scenes
                            st.session_state["kdp_autopilot_book_title"] = book_title_ap
                            st.session_state["kdp_autopilot_book_sub"] = book_sub_ap
                            st.session_state["kdp_autopilot_theme"] = book_theme_ap
                            st.session_state["kdp_autopilot_is_adult"] = is_adult_ap
                            st.session_state["kdp_autopilot_project_dir"] = book_project_dir
                            st.session_state["kdp_autopilot_manifest_file"] = txt_save_path
                            st.session_state["kdp_autopilot_trim_size"] = trim_size_ap
                            st.session_state["kdp_autopilot_aspect_ratio"] = aspect_ratio_ap

                            # Auto-save immediately to persistent disk storage
                            save_book_project("kdp_coloring", {
                                "title": book_title_ap.strip(),
                                "subtitle": book_sub_ap.strip(),
                                "theme": book_theme_ap.strip(),
                                "is_adult": is_adult_ap,
                                "trim_size": trim_size_ap,
                                "aspect_ratio": aspect_ratio_ap,
                                "project_dir": book_project_dir,
                                "scenes": parsed_scenes
                            })
                            st.session_state["kdp_ap_last_saved_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                            st.balloons()
                            st.success(
                                f"🎉 **{len(parsed_scenes)} Oldalas Könyvvázlat Sikeresen Elkészült!**\n\n"
                                f"📐 **Formátum:** `{trim_size_ap}` ({aspect_ratio_ap}) · 📁 **Projekt mappa:** `{book_project_dir}`"
                            )

        # ── Step 2-5: If Scenes are Generated ──
        if st.session_state.get("kdp_autopilot_scenes"):
            scenes_list = st.session_state["kdp_autopilot_scenes"]
            cur_proj_dir = st.session_state.get("kdp_autopilot_project_dir", resolve_drive_folder("kdp"))
            b_title = st.session_state.get("kdp_autopilot_book_title", "Bible Coloring Book")
            b_sub = st.session_state.get("kdp_autopilot_book_sub", "")
            curr_trim = st.session_state.get("kdp_autopilot_trim_size", trim_size_ap if 'trim_size_ap' in locals() else "8.5x11")
            curr_ar = st.session_state.get("kdp_autopilot_aspect_ratio", aspect_ratio_ap if 'aspect_ratio_ap' in locals() else "3:4")

            st.markdown("---")

            # ── SECTION 2: GEMINI GEM MASTER INSTRUCTION ──
            st.markdown("### 💎 2. LÉPÉS: Gemini Gem Mester Utasítás (Karakter- és Stílusállandóság)")
            st.caption("Másold be ezt a rendszerszintű utasítást a Gemini új csevegésébe (vagy hozz létre egy egyéni Gem-et) a generálás előtt!")

            gem_instruction = generate_gemini_gem_master_instruction(
                is_adult=bool(st.session_state.get("kdp_autopilot_is_adult", False)),
                active_style=active_art_style_prompt,
                aspect_ratio=curr_ar,
                trim_size=curr_trim,
                book_title=b_title,
                target_audience="Adult" if st.session_state.get("kdp_autopilot_is_adult") else "Children",
                style_name=active_art_style_prompt
            )
            render_gemini_gem_instruction_card(gem_instruction, prefix="kdp_ap")

            st.markdown("---")

            # ── SECTION 3: SCENE CARDS & PROMPT COPIERS ──
            st.markdown(f"### 📋 3. LÉPÉS: Sorszámozott Képpromptok ({len(scenes_list)} db) & 1-Kattintásos Másolás")
            st.caption(f"Másold be a promptokat egymás után a jobb oldali Gemini ablakba, majd mentsd le a képeket `{cur_proj_dir}` mappába `01.png`, `02.png` néven!")

            render_kdp_scene_cards_grid(
                scenes_data=scenes_list,
                project_dir=cur_proj_dir,
                prefix="kdp_scenes"
            )

            st.markdown("---")

            # ── SECTION 4: DRIVE IMAGE GATHERING ──
            st.markdown("### 📥 4. LÉPÉS: Képek Begyűjtése a Google Drive Mappából")
            st.caption("Amikor letöltötted a képeket a Drive-ra, kattints a Begyűjtés gombra. A rendszer automatikusan társítja a képeket a könyv oldalaihoz.")

            col_g1, col_g2 = st.columns([2, 1])
            with col_g1:
                cur_proj_dir_input = st.text_input(
                    "📁 Ellenőrzendő Drive Projektmappa:",
                    value=cur_proj_dir,
                    key="kdp_drive_scan_dir_input"
                )
            with col_g2:
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                btn_gather_images = st.button(
                    "🔄 Képek Begyűjtése & Ellenőrzése",
                    key="btn_kdp_gather_images_now",
                    use_container_width=True
                )

            matched_dict, unassigned_list = gather_images_from_directory(
                target_dir=cur_proj_dir_input,
                expected_count=len(scenes_list)
            )

            # Update scenes with matched images
            ready_count = 0
            missing_pages = []
            for s_idx, sc in enumerate(scenes_list):
                p_num = sc.get("page_number", s_idx + 1)
                matched_fp = matched_dict.get(p_num)
                if matched_fp and os.path.exists(matched_fp):
                    sc["filepath"] = matched_fp
                    try:
                        sc["pil_image"] = Image.open(matched_fp)
                        with open(matched_fp, "rb") as f_img_b:
                            sc["image_bytes"] = f_img_b.read()
                    except Exception:
                        pass
                    ready_count += 1
                else:
                    if not sc.get("image_bytes") and not sc.get("filepath"):
                        missing_pages.append(p_num)
                    else:
                        ready_count += 1

            # Status Banner
            total_req = len(scenes_list)
            if ready_count == total_req:
                st.success(f"🎉 **Kiváló! Mind a {ready_count} / {total_req} oldal képe megtalálva a Drive mappában!** Készen áll a nyomdakész PDF összeállítására.")
            else:
                missing_str = ", ".join([f"#{p}" for p in missing_pages]) if missing_pages else "Nincs"
                st.warning(f"🟡 **Állapot: {ready_count} / {total_req} kép elérhető.** Hiányzó oldalak: **{missing_str}**\n\n*Mentsd a képeket a fenti mappába `01.png`, `02.png`... néven, vagy töltsd fel őket egyenként az alábbi kártyákon!*")

            # Visual Gallery with replacement slots
            render_canva_image_gallery(
                records=scenes_list,
                prefix="kdp_gather_gal",
                context_type="kdp"
            )

            st.markdown("---")

            # ── SECTION 5: 1-CLICK KDP PDF COMPILER ──
            st.markdown(f"### 🚀 5. LÉPÉS: Nyomdakész Amazon KDP Belső PDF Összeállítása ({curr_trim})")
            st.caption(f"A ReportLab nyomdai motor összefűzi a könyvet ({curr_trim} méretben). Az alábbi beállításokkal teljesen testreszabhatod a margókat, kereteket és feliratokat!")

            # ⚙️ PDF Layout & Style Customizer Panel
            with st.expander("⚙️ Nyomdai & PDF Elrendezési Beállítások (Margó, Keret, Feliratok)", expanded=True):
                col_set1, col_set2 = st.columns(2, gap="large")
                with col_set1:
                    custom_margin_val = st.slider(
                        "📏 Biztonsági Margó Mérete (hüvelyk / inch):",
                        min_value=0.25,
                        max_value=0.85,
                        value=float(st.session_state.get("kdp_pdf_margin_val", 0.5)),
                        step=0.05,
                        help="KDP ajánlott margó: 0.5 hüvelyk (36 pt). Minimum 0.25 hüvelyk vágásmentes margó.",
                        key="kdp_pdf_margin_slider"
                    )
                    st.session_state["kdp_pdf_margin_val"] = custom_margin_val
                    
                    show_dec_frame = st.checkbox(
                        "🖼️ Dupla díszítő keret a lapok szélén (Decorative Frame)",
                        value=st.session_state.get("kdp_pdf_dec_frame", True),
                        key="kdp_pdf_dec_frame_chk",
                        help="Ha kikapcsolod, a lap szélei teljesen tiszták és keretmentesek maradnak."
                    )
                    st.session_state["kdp_pdf_dec_frame"] = show_dec_frame

                    show_img_border = st.checkbox(
                        "🔲 Fekete kontúrkeret közvetlenül a kép körül (Image Border)",
                        value=st.session_state.get("kdp_pdf_img_border", True),
                        key="kdp_pdf_img_border_chk",
                        help="Kikapcsolva a színező kép körül nem lesz plusz fekete téglalap keret."
                    )
                    st.session_state["kdp_pdf_img_border"] = show_img_border

                    show_swatches_test = st.checkbox(
                        "🧪 Színtesztelő mintaformák a címoldalon (Color Swatches Tester)",
                        value=st.session_state.get("kdp_pdf_swatches_test", True),
                        key="kdp_pdf_swatches_chk"
                    )
                    st.session_state["kdp_pdf_swatches_test"] = show_swatches_test

                with col_set2:
                    show_hdr_txt = st.checkbox(
                        "🏷️ Jelenetcím felirat a színező lap tetején (Header Text)",
                        value=st.session_state.get("kdp_pdf_hdr_txt", True),
                        key="kdp_pdf_hdr_chk",
                        help="Kikapcsolva a színező oldal tetején nem jelenik meg felirat, nagyobb hely marad a rajznak."
                    )
                    st.session_state["kdp_pdf_hdr_txt"] = show_hdr_txt

                    show_ftr_txt = st.checkbox(
                        "🔢 Lábléc & formátum felirat a lap alján (Footer Text)",
                        value=st.session_state.get("kdp_pdf_ftr_txt", True),
                        key="kdp_pdf_ftr_chk"
                    )
                    st.session_state["kdp_pdf_ftr_txt"] = show_ftr_txt

                    inc_companion = st.checkbox(
                        "📖 Bal oldali Kísérő Ige & Színpaletta oldalak (Companion Pages)",
                        value=st.session_state.get("kdp_pdf_inc_companion", True),
                        key="kdp_pdf_companion_chk",
                        help="Kikapcsolva a könyv csak színező oldalakat tartalmaz, kísérő igés lapok nélkül."
                    )
                    st.session_state["kdp_pdf_inc_companion"] = inc_companion

                    inc_bleed = st.checkbox(
                        "🔒 Filcátütés-védelmi üres lapok beszúrása (Bleed Protection)",
                        value=st.session_state.get("kdp_pdf_inc_bleed", True),
                        key="kdp_pdf_bleed_chk",
                        help="Minden rajz után üres hátlapot szúr be, hogy a filctoll ne üssön át a következő képre."
                    )
                    st.session_state["kdp_pdf_inc_bleed"] = inc_bleed

            # Dynamic Page Count Computation
            pages_per_sc = 1 + (1 if inc_companion else 0) + (1 if inc_bleed else 0)
            tot_pages_est = 1 + (len(scenes_list) * pages_per_sc)

            col_pdf_b1, col_pdf_b2 = st.columns([1.2, 1])

            with col_pdf_b1:
                btn_compile_kdp_pdf = st.button(
                    f"🚀 Nyomdakész KDP Belső PDF Összeállítása ({tot_pages_est} oldal · {curr_trim})",
                    key="btn_compile_kdp_interior_pdf",
                    use_container_width=True
                )

            pdf_out_path = os.path.join(cur_proj_dir_input, f"KDP_{sanitize_filename(b_title)}_{curr_trim}.pdf")

            if btn_compile_kdp_pdf:
                with st.spinner(f"Nyomdakész Amazon KDP belső PDF ({curr_trim} · {custom_margin_val}\" margó) összeállítása folyamatban..."):
                    ok_pdf, pdf_bytes, msg_pdf = build_kdp_book_pdf(
                        title=b_title,
                        subtitle=b_sub,
                        pages_data=scenes_list,
                        output_path=pdf_out_path,
                        trim_size=curr_trim,
                        margin_in=custom_margin_val,
                        show_decorative_frame=show_dec_frame,
                        show_image_border=show_img_border,
                        show_header_text=show_hdr_txt,
                        show_footer_text=show_ftr_txt,
                        include_companion_pages=inc_companion,
                        include_bleed_protection=inc_bleed,
                        include_swatches_tester=show_swatches_test
                    )
                    if ok_pdf and pdf_bytes:
                        st.session_state["kdp_autopilot_pdf_bytes"] = pdf_bytes
                        st.session_state["kdp_autopilot_pdf_filename"] = f"KDP_{sanitize_filename(b_title)}_{curr_trim}.pdf"
                        st.session_state["kdp_autopilot_pdf_path"] = pdf_out_path
                        st.balloons()
                        st.success(
                            f"✨ **A Nyomdakész Amazon KDP Belső PDF Sikeresen Elkészült!**\n\n"
                            f"📄 **Összes oldalszám:** `{tot_pages_est} oldal` · **Méret:** `{curr_trim}` · **Margó:** `{custom_margin_val}\"`\n\n"
                            f"📁 **PDF elérési út:** `{pdf_out_path}`"
                        )
                    else:
                        st.error(f"Hiba a PDF összeállításakor: {msg_pdf}")

            with col_pdf_b2:
                saved_pdf_b = st.session_state.get("kdp_autopilot_pdf_bytes")
                saved_pdf_fn = st.session_state.get("kdp_autopilot_pdf_filename", f"KDP_{sanitize_filename(b_title)}_{curr_trim}.pdf")
                if saved_pdf_b:
                    st.download_button(
                        label=f"📥 KÉSZ KDP BELSŐ PDF LETÖLTÉSE ({tot_pages_est} oldal · {curr_trim})",
                        data=saved_pdf_b,
                        file_name=saved_pdf_fn,
                        mime="application/pdf",
                        key="btn_dl_kdp_interior_pdf_final",
                        use_container_width=True
                    )

            # Virtual Flipbook Preview
            render_virtual_book_flipbook_preview(
                scenes_data=scenes_list,
                book_title=b_title,
                book_subtitle=b_sub,
                prefix="kdp_web_flip"
            )

    # ── TAB 2: EGYEDI KDP PROMPT KÉSZÍTŐ (1 JELENET) ──
    with tab_kdp_single:
        col_e_sel1, col_e_sel2 = st.columns([1.1, 1])
        with col_e_sel1:
            kdp_edition = st.radio(
                "👥 Célközönség & Kiadás módja:",
                [
                    "🧒 Gyermek Kiadás (Egyszerű vastag vonalak, színajánló körök, gyerekbarát ige)",
                    "🧘 Felnőtt Kiadás (Vékony vonalak, intrikátus mandala/zentangle minták, teljes ige kalligráfia)"
                ],
                key="kdp_edition_single"
            )
            is_adult_mode = "Felnőtt Kiadás" in kdp_edition

        with col_e_sel2:
            kdp_section = st.radio(
                "Mit szeretnél generálni?",
                [
                    "🖍️ Színező Oldal (B&W Belső Lap)",
                    "🎨 Színajánló & Igehely Oldal (Color Guide / Scripture Calligraphy)",
                    "📕 Könyvborító (Colorful Cover)",
                    "🧪 Címoldal & Színtesztelő Oldal (Title + Color Tester Page)"
                ],
                key="kdp_section_single"
            )

        st.markdown("---")
        col_form, col_out = st.columns([1, 1.05], gap="large")

        gen_companion = False
        companion_base = ""

        with col_form:
            # Aspect ratio and format selector for single prompt
            single_trim, single_ar, single_info = render_kdp_format_and_aspect_selector(
                widget_prefix="kdp_single",
                default_format="8.5x11"
            )

            if "Színező Oldal (" in kdp_section:
                st.markdown(f"<div class='step-label'>Lépés 1 — {'Felnőtt' if is_adult_mode else 'Gyerek'} Színező Jelenet Megadása</div>", unsafe_allow_html=True)
                scene_kdp = st.text_area(
                    "Bibliai jelenet leírása (magyar vagy angol)",
                    value="Daniel in the lions' den surrounded by intricate floral and mandala patterns" if is_adult_mode else "Noah standing on the deck of the ark with two giraffes, rainbow in the background",
                    height=90,
                    key=f"scene_kdp_{'adult' if is_adult_mode else 'child'}",
                    help="Pl.: Moses parting the Red Sea, David with his slingshot facing Goliath"
                )
                verse_kdp_opt = st.text_input(
                    "Bibliai Igehely & Idézet (KJV)",
                    value="Daniel 6:22 - My God hath sent his angel, and hath shut the lions' mouths" if is_adult_mode else "Genesis 6:19 - And of every living thing of all flesh, two of every sort shalt thou bring into the ark",
                    key=f"verse_kdp_opt_{'adult' if is_adult_mode else 'child'}"
                )
                colors_kdp_opt = st.text_input(
                    "Javasolt Színek / Paletta (opcionális)",
                    value="Soft sky blue, ocean teal, warm sun yellow, earthy brown, emerald green",
                    key=f"colors_kdp_opt_{'adult' if is_adult_mode else 'child'}"
                )
                extra_kdp = st.text_area(
                    "➕ Extra megjegyzés / plusz leírás (opcionális)",
                    placeholder="Pl.: Include intricate mandala and stained glass motifs." if is_adult_mode else "Pl.: The animals should look very cute and friendly.",
                    height=70,
                    key=f"extra_kdp_{'adult' if is_adult_mode else 'child'}"
                )

                if is_adult_mode:
                    base = template_kdp_adult_coloring(scene_kdp, aspect_ratio=single_ar, trim_size=single_trim)
                else:
                    base = template_kdp_coloring(scene_kdp, aspect_ratio=single_ar, trim_size=single_trim)

                gen_companion = st.checkbox(
                    f"🎨 Kapcsolódó Bal Oldali {'Kalligráfia Ige' if is_adult_mode else 'Színajánló & Igehely'} Oldal Prompt generálása is",
                    value=True,
                    key=f"gen_companion_{'adult' if is_adult_mode else 'child'}"
                )
                if gen_companion:
                    if is_adult_mode:
                        companion_base = template_kdp_adult_scripture_companion(scene_desc=scene_kdp, scripture_verse=verse_kdp_opt)
                    else:
                        companion_base = template_kdp_color_guide_companion(scene_desc=scene_kdp, scripture_verse=verse_kdp_opt, color_palette=colors_kdp_opt)

            elif "Színajánló" in kdp_section:
                st.markdown(f"<div class='step-label'>Lépés 1 — {'Felnőtt Kalligráfia' if is_adult_mode else 'Gyerek Színajánló'} Paraméterek</div>", unsafe_allow_html=True)
                scene_guide = st.text_input(
                    "Kapcsolódó Jobb Oldali Jelenet Témája",
                    value="Daniel in the lions' den" if is_adult_mode else "Noah's Ark on Mount Ararat",
                    key=f"scene_guide_{'adult' if is_adult_mode else 'child'}"
                )
                verse_guide = st.text_input(
                    "Bibliai Igehely & Teljes Idézet (KJV)",
                    value="Daniel 6:22 - My God hath sent his angel, and hath shut the lions' mouths" if is_adult_mode else "Genesis 8:1 - And God remembered Noah, and every living thing",
                    key=f"verse_guide_{'adult' if is_adult_mode else 'child'}"
                )
                colors_guide = st.text_input(
                    "Színpaletta Leírása (opcionális)",
                    value="Gold, deep blue, warm amber, emerald green",
                    key=f"colors_guide_{'adult' if is_adult_mode else 'child'}"
                )
                extra_kdp = st.text_area(
                    "➕ Extra megjegyzés (opcionális)",
                    placeholder="Pl.: Use elegant botanical border.",
                    height=70,
                    key=f"extra_guide_{'adult' if is_adult_mode else 'child'}"
                )
                if is_adult_mode:
                    base = template_kdp_adult_scripture_companion(scene_guide, verse_guide)
                else:
                    base = template_kdp_color_guide_companion(scene_guide, verse_guide, colors_guide)

            elif "Könyvborító" in kdp_section:
                st.markdown("<div class='step-label'>Lépés 1 — Borító Paraméterek</div>", unsafe_allow_html=True)
                title_cover = st.text_input("Könyv Főcíme (Nagybetűkkel):", value="BIBLE STORIES COLORING BOOK", key="title_cover_single")
                subtitle_cover_opt = st.text_input("Alcím / Célközönség leírás:", value="For Kids Ages 4-8 · 30 Inspiring Stories", key="subtitle_cover_single")
                theme_cover = st.text_area("Borító Illusztráció Leírása:", value="Noah's ark on calm blue water with a vibrant rainbow in the sky, two smiling giraffes and lions on deck", height=80, key="theme_cover_single")
                extra_kdp = st.text_area("➕ Extra megjegyzés:", placeholder="Pl.: Bright sunny palette.", height=70, key="extra_kdp_cover_single")
                base = template_kdp_cover(title_cover, theme_cover, subtitle_cover_opt, aspect_ratio=single_ar)

            else:  # Címoldal & Színtesztelő Oldal
                st.markdown("<div class='step-label'>Lépés 1 — Címoldal & Színtesztelő Adatok</div>", unsafe_allow_html=True)
                title_tester = st.text_input("Könyv Címe", value="BIBLE COLORING BOOK FOR KIDS", key="title_tester_single")
                subtitle_tester = st.text_input("Birtokló / Alcím felirat", value="This Book Belongs To: ________", key="subtitle_tester_single")
                extra_kdp = st.text_area("➕ Extra megjegyzés:", placeholder="Pl.: Include cute heart shapes.", height=70, key="extra_kdp_tester_single")
                base = template_kdp_title_color_tester(title_tester, subtitle_tester)

            generate_kdp = st.button("✨ Profi Prompt Generálása", key="gen_kdp_single", use_container_width=True)

        with col_out:
            st.markdown("<div class='step-label'>Kész Angol Prompt — Másold be a Gemini-be</div>", unsafe_allow_html=True)

            if generate_kdp:
                raw_extra = extra_kdp if "extra_kdp" in locals() else ""
                if is_adult_mode:
                    edition_req = " TARGET EDITION: Intricate adult coloring book page, fine black line art, detailed zentangle and mandala floral background patterns, zero shading, pure white background."
                else:
                    edition_req = " TARGET EDITION: Children's coloring book page, simple clean thick bold black outlines, pure white background, cute friendly illustration."
                extra_val = (raw_extra + edition_req).strip()

                result = enhance_prompt_with_gemini(True, base, extra_val, text_model, temperature, active_art_style_prompt, "")
                comp_res = enhance_prompt_with_gemini(True, companion_base, extra_val, text_model, temperature, active_art_style_prompt, "") if (gen_companion and companion_base) else ""

                st.session_state["kdp_single_result"] = result
                st.session_state["kdp_single_comp"] = comp_res

            if st.session_state.get("kdp_single_result"):
                st.code(st.session_state["kdp_single_result"], language="text")
                if st.session_state.get("kdp_single_comp"):
                    st.markdown("**📄 Kapcsolódó oldal promptja:**")
                    st.code(st.session_state["kdp_single_comp"], language="text")


# ══════════════════════════════════════════════════════════
# WORKSPACE 2: ILLUSZTRÁLT & ÍROTT MESEKÖNYV MŰHELY
# ══════════════════════════════════════════════════════════

elif "Illusztrált" in menu_choice or "Mesekönyv" in menu_choice:
    st.markdown("<div class='path-badge'>📖 Amazon KDP Mesekönyv Műhely</div>", unsafe_allow_html=True)
    st.markdown("### 📖 Illusztrált & Írott Mesekönyv Műhely (Bármilyen Témában)")
    st.caption("Generálj teljes történetet, tanulságos mesét vagy fejezetes könyvet megírt szöveggel és a hozzájuk tartozó színes Gemini illusztrációs promptokkal!")

    active_art_style_prompt = render_style_selector("ws_ib")
    render_niche_status_bar("ib")

    # Project Management & Persistence Bar for Illustrated Books
    render_book_project_manager_bar(project_type="kdp_illustrated", widget_prefix="ib_ed")

    col_ib1, col_ib2 = st.columns([1.1, 1], gap="large")
    with col_ib1:
        n_slug_ib = get_niche_slug(chosen_niche)
        ib_title = st.text_input(
            "📖 Könyv Címe:",
            value=st.session_state.get(f"ib_title_{n_slug_ib}", get_niche_field("storybook_title", chosen_niche)),
            key=f"ib_title_{n_slug_ib}"
        )
        ib_subtitle = st.text_input(
            "🏷️ Alcím / Célközönség:",
            value=st.session_state.get(f"ib_sub_{n_slug_ib}", get_niche_field("storybook_sub", chosen_niche)),
            key=f"ib_sub_{n_slug_ib}"
        )
        ib_theme = st.text_area(
            "🌟 Történet témája és fő tanulsága:",
            value=st.session_state.get(f"ib_theme_{n_slug_ib}", get_niche_field("storybook_theme", chosen_niche)),
            height=100,
            key=f"ib_theme_{n_slug_ib}"
        )
    with col_ib2:
        # Format & Aspect Ratio Selector for Illustrated Books (Default 8.5x8.5 Square)
        ib_trim_size, ib_aspect_ratio, ib_info = render_kdp_format_and_aspect_selector(
            widget_prefix="ib_ed",
            default_format=st.session_state.get("ib_trim_size", "8.5x8.5")
        )

        ib_chapters_cnt = st.select_slider(
            "🔢 Fejezetek / Oldalak Száma:",
            options=[4, 6, 8, 10, 12, 16],
            value=st.session_state.get("ib_chapters_cnt_val", 6),
            key="ib_chapters_slider"
        )
        st.session_state["ib_chapters_cnt_val"] = ib_chapters_cnt

        btn_gen_ib_story = st.button("✨ Történet & Illusztrációs Promptok Generálása (AI)", key="btn_gen_ib_story", use_container_width=True)

    if btn_gen_ib_story:
        if not ib_title.strip() or not ib_theme.strip():
            st.error("⚠️ Kérlek add meg a könyv címét és témáját!")
        else:
            with st.spinner(f"AI ({text_model}) megírja a(z) {ib_chapters_cnt} fejezetes történetet és illusztrációs promptokat ({ib_trim_size} · {ib_aspect_ratio})..."):
                ib_prompt = build_illustrated_book_manifest_prompt(
                    book_title=ib_title.strip(),
                    book_genre="Illustrated Story",
                    target_audience=ib_subtitle.strip(),
                    theme_storyline=ib_theme.strip(),
                    page_count=ib_chapters_cnt,
                    art_style=active_art_style_prompt,
                    trim_size=ib_trim_size,
                    aspect_ratio=ib_aspect_ratio
                )
                ok_ib, res_ib_raw = km.generate_text_with_fallback(prompt=ib_prompt, model_name=text_model)
                if ok_ib and res_ib_raw:
                    parsed_ib = parse_illustrated_book_manifest_json(res_ib_raw)
                    if parsed_ib:
                        drive_kdp_dir = resolve_drive_folder("kdp")
                        ib_proj_dir = os.path.join(drive_kdp_dir, f"Illustrated_{sanitize_filename(ib_title)}_{time.strftime('%Y%m%d_%H%M%S')}")
                        os.makedirs(ib_proj_dir, exist_ok=True)
                        st.session_state["ib_pages_data"] = parsed_ib
                        st.session_state["ib_book_title"] = ib_title
                        st.session_state["ib_book_sub"] = ib_subtitle
                        st.session_state["ib_book_theme"] = ib_theme
                        st.session_state["ib_trim_size"] = ib_trim_size
                        st.session_state["ib_aspect_ratio"] = ib_aspect_ratio
                        st.session_state["ib_proj_dir"] = ib_proj_dir

                        # Auto-save immediately to disk
                        save_book_project("kdp_illustrated", {
                            "title": ib_title.strip(),
                            "subtitle": ib_subtitle.strip(),
                            "theme": ib_theme.strip(),
                            "trim_size": ib_trim_size,
                            "aspect_ratio": ib_aspect_ratio,
                            "project_dir": ib_proj_dir,
                            "pages": parsed_ib
                        })
                        st.session_state["ib_ed_last_saved_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                        st.success(f"🎉 **{len(parsed_ib)} fejezetes történet sikeresen megírva!** ({ib_trim_size} · {ib_aspect_ratio})")
                    else:
                        st.error("Nem sikerült értelmezni a generált történet JSON formátumát.")
                else:
                    st.error(f"Hiba a történet generálásakor: {res_ib_raw}")

    if st.session_state.get("ib_pages_data"):
        ib_pages = st.session_state["ib_pages_data"]
        ib_p_dir = st.session_state.get("ib_proj_dir", resolve_drive_folder("kdp"))
        curr_ib_trim = st.session_state.get("ib_trim_size", ib_trim_size if 'ib_trim_size' in locals() else "8.5x8.5")
        
        st.markdown("---")
        st.markdown("### 📁 Illusztrációk Begyűjtése & Kézirat Szerkesztő")
        
        col_ib_s1, col_ib_s2 = st.columns([2, 1])
        with col_ib_s1:
            ib_dir_in = st.text_input("Drive Illusztrációk Mappája:", value=ib_p_dir, key="input_ib_dir")
        with col_ib_s2:
            st.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
            if st.button("🔄 Képek Begyűjtése", key="btn_gather_ib_imgs", use_container_width=True):
                matched_ib, _ = gather_images_from_directory(ib_dir_in, expected_count=len(ib_pages))
                for idx, pg in enumerate(ib_pages):
                    p_num = pg.get("page_number", idx + 1)
                    if p_num in matched_ib:
                        pg["filepath"] = matched_ib[p_num]
                        try:
                            pg["pil_image"] = Image.open(matched_ib[p_num])
                            with open(matched_ib[p_num], "rb") as f_ib:
                                pg["image_bytes"] = f_ib.read()
                        except Exception:
                            pass
                st.toast("Képek frissítve!", icon="🖼️")
                st.rerun()

        render_illustrated_book_editor_cards(pages_data=ib_pages, prefix="ib_ed", image_model_name="")

        st.markdown("---")
        
        # ⚙️ Illustrated Book PDF Layout Options
        with st.expander("⚙️ Nyomdai & Illusztrált PDF Elrendezési Beállítások", expanded=True):
            col_ib_opt1, col_ib_opt2 = st.columns(2, gap="large")
            with col_ib_opt1:
                ib_margin_val = st.slider(
                    "📏 Biztonsági Margó Mérete (hüvelyk / inch):",
                    min_value=0.25,
                    max_value=0.85,
                    value=float(st.session_state.get("ib_pdf_margin_val", 0.5)),
                    step=0.05,
                    key="ib_pdf_margin_slider"
                )
                st.session_state["ib_pdf_margin_val"] = ib_margin_val

                ib_dec_frame = st.checkbox(
                    "🖼️ Díszítő keret a címoldalon",
                    value=st.session_state.get("ib_pdf_dec_frame", True),
                    key="ib_pdf_dec_frame_chk"
                )
                st.session_state["ib_pdf_dec_frame"] = ib_dec_frame

            with col_ib_opt2:
                ib_img_border = st.checkbox(
                    "🔲 Finom keret az illusztráció körül",
                    value=st.session_state.get("ib_pdf_img_border", True),
                    key="ib_pdf_img_border_chk"
                )
                st.session_state["ib_pdf_img_border"] = ib_img_border

                ib_show_hdr = st.checkbox(
                    "🏷️ Fejezetcím felirat a lap tetején",
                    value=st.session_state.get("ib_pdf_hdr", True),
                    key="ib_pdf_hdr_chk"
                )
                st.session_state["ib_pdf_hdr"] = ib_show_hdr

                ib_show_pgnum = st.checkbox(
                    "🔢 Oldalszámozás a lap alján",
                    value=st.session_state.get("ib_pdf_pgnum", True),
                    key="ib_pdf_pgnum_chk"
                )
                st.session_state["ib_pdf_pgnum"] = ib_show_pgnum

        if st.button(f"🚀 Illusztrált Könyv Nyomdakész PDF Összeállítása ({curr_ib_trim})", key="btn_compile_ib_pdf", use_container_width=True):
            ib_pdf_out = os.path.join(ib_dir_in, f"Illustrated_{sanitize_filename(st.session_state.get('ib_book_title', 'Book'))}_{curr_ib_trim}.pdf")
            ok_ib_pdf, ib_pdf_bytes, ib_pdf_msg = build_illustrated_book_pdf(
                title=st.session_state.get("ib_book_title", "Book"),
                subtitle=st.session_state.get("ib_book_sub", ""),
                pages_data=ib_pages,
                output_path=ib_pdf_out,
                trim_size=curr_ib_trim,
                margin_in=ib_margin_val,
                show_decorative_frame=ib_dec_frame,
                show_image_border=ib_img_border,
                show_page_numbers=ib_show_pgnum,
                show_chapter_header=ib_show_hdr
            )
            if ok_ib_pdf:
                st.success(f"✨ **Illusztrált Könyv PDF Sikeresen Elkészült!**\n\n📁 `{ib_pdf_out}`")
                st.download_button(
                    label=f"📥 KÉSZ ILLUSZTRÁLT KÖNYV PDF LETÖLTÉSE ({curr_ib_trim})",
                    data=ib_pdf_bytes,
                    file_name=f"Illustrated_{sanitize_filename(st.session_state.get('ib_book_title', 'Book'))}_{curr_ib_trim}.pdf",
                    mime="application/pdf",
                    key="btn_dl_ib_pdf_direct"
                )
            else:
                st.error(f"Hiba a PDF fordításakor: {ib_pdf_msg}")


# ══════════════════════════════════════════════════════════
# WORKSPACE 3: KDP BORÍTÓ & DINAMIKUS GERINC MESTER
# ══════════════════════════════════════════════════════════

elif "Borító" in menu_choice or "Gerinc" in menu_choice:
    st.markdown("<div class='path-badge'>🎨 Amazon KDP Borító Stúdió</div>", unsafe_allow_html=True)
    st.markdown("### 🎨 KDP Wrap-Around Borító & Dinamikus Gerinc Mester")
    st.caption("Számítsd ki a pontos gerincvastagságot és teljes borító méretet (Front + Back + Spine), majd generálj hozzá illeszkedő Gemini borító promptot!")

    active_art_style_prompt = render_style_selector("ws_cov")
    render_niche_status_bar("cov")

    n_slug_cov = get_niche_slug(chosen_niche)
    col_cov1, col_cov2 = st.columns(2, gap="large")
    with col_cov1:
        cov_title = st.text_input(
            "📖 Könyv Főcíme a Borítón:",
            value=st.session_state.get(f"cov_title_{n_slug_cov}", get_niche_field("cov_title", chosen_niche)),
            key=f"cov_title_{n_slug_cov}"
        )
        cov_sub = st.text_input(
            "🏷️ Alcím:",
            value=st.session_state.get(f"cov_sub_{n_slug_cov}", get_niche_field("cov_sub", chosen_niche)),
            key=f"cov_sub_{n_slug_cov}"
        )
        cov_theme = st.text_area(
            "🌟 Borító Illusztráció Leírása:",
            value=st.session_state.get(f"cov_theme_{n_slug_cov}", get_niche_field("cov_theme", chosen_niche)),
            height=80,
            key=f"cov_theme_{n_slug_cov}"
        )

    with col_cov2:
        cov_pages = st.number_input("🔢 Teljes Oldalszám a Belső PDF-ben (Spine math):", min_value=24, max_value=800, value=76, step=2, key="cov_pages_inp")
        cov_trim = st.selectbox("📐 Vágási Méret (Trim Size):", list(TRIM_SIZES.keys()), index=0, key="cov_trim_inp")
        cov_paper = st.selectbox("📄 Papírtípus:", ["Fehér papír (White Paper)", "Krém papír (Cream Paper)"], index=0, key="cov_paper_inp")
        is_cream = "Krém" in cov_paper

        cov_dims = calculate_kdp_cover_dimensions(page_count=int(cov_pages), trim_size=cov_trim, paper_type="cream" if is_cream else "white")

        st.markdown(f"""
        <div style='background:#1e2536; border:1px solid #3b82f6; border-radius:10px; padding:10px 14px; font-size:0.85rem; color:#cbd5e1;'>
            <b>📐 Számított KDP Nyomdai Méretek:</b><br>
            • Teljes borítószélesség: <b>{cov_dims['total_cover_width_in']} hüvelyk</b> ({cov_dims['width_px_300dpi']} px)<br>
            • Teljes magasság: <b>{cov_dims['total_cover_height_in']} hüvelyk</b> ({cov_dims['height_px_300dpi']} px)<br>
            • Gerincvastagság (Spine): <b>{cov_dims['spine_width_in']} hüvelyk</b> ({cov_dims['spine_width_px_300dpi']} px)
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    if st.button("✨ KDP Borító Prompt Generálása (Gemini Web)", key="btn_gen_cov_prompt", use_container_width=True):
        cov_prompt = build_kdp_dynamic_cover_prompt(
            book_title=cov_title,
            theme_desc=cov_theme,
            page_count=int(cov_pages),
            trim_size=cov_trim,
            paper_type="cream" if is_cream else "white",
            art_style=active_art_style_prompt,
            subtitle=cov_sub
        )
        st.session_state["kdp_cover_generated_prompt"] = cov_prompt

    if st.session_state.get("kdp_cover_generated_prompt"):
        st.markdown("### 📋 Kész KDP Borító Prompt — Másold be a Gemini-be:")
        st.code(st.session_state["kdp_cover_generated_prompt"], language="text")
        st.info("💡 **Tipp:** A Gemini által generált borítóképet mentsd le `cover.png` néven a projekt mappájába!")


# ==========================================================
# WORKSPACE: ETSY FALIKÉPEK & CLIPARTOK
# ==========================================================

elif "Etsy" in menu_choice or "2." in menu_choice:
    st.markdown("<div class='path-badge'>🎨 Etsy Grafikai Stúdió</div>", unsafe_allow_html=True)
    st.markdown("### 🎨 Keresztény Faliképek & Clipart Csomagok (Etsy)")
    st.caption("Készíts múzeumi minőségű akvarell, boho vagy aranyfóliás falikép és clipart promptokat.")

    active_art_style_prompt = render_style_selector("ws_etsy")
    render_niche_status_bar("etsy")

    tab_etsy_single, tab_etsy_batch = st.tabs([
        "✍️ Egyedi Etsy Prompt Készítő",
        "📦 Tömeges Etsy Prompt Csomag (30 Ötlet alapján)"
    ])

    with tab_etsy_single:
        etsy_sub = st.radio(
            "Mit szeretnél generálni?",
            ["🖼️ Igés Falikép (Wall Art)", "✂️ Clipart Csomag (PNG Képek / Matricák)"],
            horizontal=True,
            key="etsy_sub_choice"
        )

        st.markdown("---")
        col_form2, col_out2 = st.columns([1, 1.05], gap="large")

        with col_form2:
            n_slug_etsy = get_niche_slug(chosen_niche)
            if "Falikép" in etsy_sub:
                st.markdown("<div class='step-label'>Igeidézet / Idézet megadása</div>", unsafe_allow_html=True)
                verse_etsy = st.text_input(
                    "Exakt bibliai ige vagy inspiráló idézet (angolul)",
                    value=st.session_state.get(f"verse_etsy_{n_slug_etsy}", get_niche_field("verse_etsy", chosen_niche)),
                    key=f"verse_etsy_{n_slug_etsy}",
                    help="Pontosan azt írd ide, amit a falikép közepén szeretnél megjeleníteni."
                )
                extra_etsy = st.text_area(
                    "➕ Extra megjegyzés / stílusjegyzet (opcionális)",
                    placeholder="Pl.: Use soft lavender and blush pink tones. Add small watercolor flowers at the corners.",
                    height=90,
                    key=f"extra_etsy_{n_slug_etsy}"
                )
                base2 = template_etsy_wall_art(verse_etsy)

            else:  # Clipart
                st.markdown("<div class='step-label'>Karakter / Tárgy megadása</div>", unsafe_allow_html=True)
                subject_etsy = st.text_area(
                    "Mit ábrázoljon a clipart?",
                    value=st.session_state.get(f"subject_etsy_{n_slug_etsy}", get_niche_field("subject_etsy", chosen_niche)),
                    height=90,
                    key=f"subject_etsy_{n_slug_etsy}",
                    help="Pl.: baby Jesus in a manger, Noah's dove with olive branch, angel with wings"
                )
                extra_etsy = st.text_area(
                    "➕ Extra megjegyzés (opcionális)",
                    placeholder="Pl.: Make 3 variations: one frontal, one side view, one from above. Soft blush pink palette.",
                    height=80,
                    key=f"extra_etsy_clip_{n_slug_etsy}"
                )
                base2 = template_etsy_clipart(subject_etsy)

            with st.expander("📄 Alap sablon (kattints a megtekintéshez)", expanded=False):
                st.code(base2, language="text")

            generate_etsy = st.button("✨ Profi Prompt Generálása", key="gen_etsy", use_container_width=True)

        with col_out2:
            st.markdown("<div class='step-label'>Kész Angol Prompt — Másold be a Gemini-be</div>", unsafe_allow_html=True)

            if "etsy_result" not in st.session_state:
                st.session_state.etsy_result = ""

            if generate_etsy:
                curr_t2 = time.strftime("%H:%M:%S")
                extra_val2 = extra_etsy if "extra_etsy" in locals() else ""
                if client:
                    with st.spinner("Gemini feljavítja a promptot..."):
                        result2 = enhance_prompt_with_gemini(client, base2, extra_val2, text_model, temperature, active_art_style_prompt, current_image_model)
                else:
                    result2 = base2 + ("\n\n[Style: " + active_art_style_prompt + "]\n\n[Extra notes: " + extra_val2 + "]" if extra_val2.strip() else "")
                    if not api_key:
                        st.warning("⚠️ API kulcs nélkül az alap sablont adjuk vissza.")

                st.session_state["etsy_result"] = result2
                st.session_state["etsy_copy_box"] = result2
                st.session_state["etsy_time"] = curr_t2

            if st.session_state.get("etsy_result"):
                st.markdown(f"<div class='prompt-output'>{st.session_state.etsy_result}</div>", unsafe_allow_html=True)
                st.text_area(
                    "📋 Másold innen (Ctrl+A → Ctrl+C):",
                    value=st.session_state.get("etsy_copy_box", st.session_state.etsy_result),
                    height=200,
                    key="etsy_copy_box"
                )
                time_tag2 = f" (Frissítve: {st.session_state.etsy_time})" if "etsy_time" in st.session_state else ""
                st.success(f"✅ Prompt kész{time_tag2}! Másold be a Gemini képgenerálójába.")

                st.markdown("---")
                if st.button("💾 Promptek mentése Google Drive-ra (04_🖼️_ETSY_DIGITAL)", key="btn_save_etsy_drive", use_container_width=True):
                    th_label = verse_etsy if "Falikép" in etsy_sub else (subject_etsy if 'subject_etsy' in locals() else "Etsy_Grafika")
                    ok_s, res_path = save_prompts_file_to_drive("etsy", th_label, st.session_state.etsy_result, header_info=f"Etsy Prompt - {etsy_sub}")
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{res_path}`")
                    else:
                        st.error(f"Hiba a mentéskor: {res_path}")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🎨</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Töltsd ki az adatokat a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">✨ Profi Prompt Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ── TAB 2: TÖMEGES ETSY PROMPT CSOMAG ──
    with tab_etsy_batch:
        st.markdown("#### 📦 Tömeges Etsy Prompt Csomag Generáló (Google Drive & 30 Ötlet)")
        st.caption("Töltsd be a 4. Munkaterületen generált 30 falikép vagy clipart ötletet, vagy válassz ki egy `.txt` fájlt a Google Drive-ról (`04_🖼️_ETSY_DIGITAL`), és generálj le minden tételhez egy kész, professzionális képgeneráló promptot!")

        col_eb_src, col_eb_cfg = st.columns([1, 1], gap="large")

        with col_eb_src:
            st.markdown("<div class='step-label'>📂 1. Lépés: Ötletlista / Forrás Kiválasztása</div>", unsafe_allow_html=True)
            etsy_batch_src_mode = st.radio(
                "Honnan szeretnéd betölteni az ötleteket?",
                [
                    "⚡ 1. Legutóbbi 30 Ötlet betöltése (4. Munkaterületről)",
                    "📁 2. Google Drive .txt fájl betöltése (04_🖼️_ETSY_DIGITAL)",
                    "📤 3. Saját .txt fájl feltöltése",
                    "✍️ 4. Kézi lista beillesztése"
                ],
                key="etsy_batch_src_mode"
            )

            loaded_etsy_batch_text = ""
            batch_etsy_source_name = "Etsy_30_Otlet"

            if "Legutóbbi" in etsy_batch_src_mode:
                last_ideas = st.session_state.get("b_ideas_result", "")
                if last_ideas and not last_ideas.startswith("⚠️"):
                    loaded_etsy_batch_text = last_ideas
                    st.success("⚡ 30 Ötlet sikeresen átvéve a 4. Munkaterületről!")
                else:
                    st.info("ℹ️ A 4. Munkaterületen még nem generáltál ötletlistát a jelenlegi munkamenetben. Generálj 30 ötletet ott, vagy válassz Drive .txt fájlt!")

            elif "Google Drive" in etsy_batch_src_mode:
                etsy_drive_dir = resolve_drive_folder("etsy")
                txt_files_etsy = []
                if os.path.exists(etsy_drive_dir):
                    txt_files_etsy = [f for f in os.listdir(etsy_drive_dir) if f.lower().endswith(".txt")]
                
                if txt_files_etsy:
                    selected_etsy_txt = st.selectbox(
                        f"Válassz .txt fájlt a 04_🖼️_ETSY_DIGITAL mappából ({len(txt_files_etsy)} db):",
                        options=txt_files_etsy,
                        key="sel_etsy_batch_drive_txt"
                    )
                    if selected_etsy_txt:
                        batch_etsy_source_name = selected_etsy_txt.replace(".txt", "")
                        full_p = os.path.join(etsy_drive_dir, selected_etsy_txt)
                        try:
                            with open(full_p, "r", encoding="utf-8") as tf:
                                loaded_etsy_batch_text = tf.read()
                            st.success(f"📄 Betöltve a Drive-ról: `{selected_etsy_txt}`")
                        except Exception as e:
                            st.error(f"Hiba a fájl olvasásakor: {e}")
                else:
                    st.warning(f"⚠️ Ebben a mappában még nincs `.txt` fájl: `{etsy_drive_dir}`")

            elif "feltöltése" in etsy_batch_src_mode:
                uploaded_etsy_batch = st.file_uploader("Tölts fel egy .txt fájlt:", type=["txt"], key="up_etsy_batch_txt")
                if uploaded_etsy_batch:
                    batch_etsy_source_name = uploaded_etsy_batch.name.replace(".txt", "")
                    loaded_etsy_batch_text = uploaded_etsy_batch.getvalue().decode("utf-8", errors="ignore")
                    st.success(f"📄 Feltöltve: `{uploaded_etsy_batch.name}`")

            else:
                loaded_etsy_batch_text = st.text_area(
                    "Illeszd be az Etsy ötleteket (soronként vagy sorszámozva):",
                    value="1. The Lord is my shepherd, I lack nothing — Psalm 23:1\n2. Be strong and courageous, do not be afraid — Joshua 1:9\n3. For I know the plans I have for you — Jeremiah 29:11\n4. I can do all this through him who gives me strength — Philippians 4:13\n5. Trust in the Lord with all your heart — Proverbs 3:5",
                    height=160,
                    key="manual_etsy_batch_txt"
                )
                batch_etsy_source_name = "Kezi_Etsy_Tetelek"

            parsed_etsy_items = parse_prompts_from_text(loaded_etsy_batch_text)
            st.markdown(f"**🎯 Felismert tételek száma:** `{len(parsed_etsy_items)} db`")

            if parsed_etsy_items:
                with st.expander(f"📋 Betöltött Tételek Megtekintése ({len(parsed_etsy_items)} db)", expanded=False):
                    for p_it in parsed_etsy_items:
                        st.markdown(f"**{p_it['index']}. {p_it['title']}**")
                        st.caption(p_it['visual_prompt'])

        with col_eb_cfg:
            st.markdown("<div class='step-label'>⚙️ 2. Lépés: Csomag Paraméterek & Generálás</div>", unsafe_allow_html=True)
            batch_etsy_type = st.radio(
                "Terméktípus (egész csomagra):",
                [
                    "🖼️ Igés Falikép Csomag (Wall Art)",
                    "✂️ Clipart / Matrica Csomag (Átlátszó háttér)"
                ],
                key="etsy_batch_type_choice"
            )
            is_wall_art_batch = "Falikép" in batch_etsy_type

            max_items_etsy = len(parsed_etsy_items) if parsed_etsy_items else 30
            item_limit_etsy = st.slider(
                "Feldolgozandó tételek száma:",
                min_value=1,
                max_value=max(1, max_items_etsy),
                value=max(1, max_items_etsy),
                key="etsy_batch_limit"
            )

            st.markdown("---")
            btn_start_etsy_batch = st.button("✨ Teljes Etsy Profi Prompt Csomag Generálása (AI)", key="btn_start_etsy_batch", use_container_width=True)

        if btn_start_etsy_batch:
            if not parsed_etsy_items:
                st.error("⚠️ Nincsenek feldolgozható tételek! Válassz vagy tölts be egy érvényes listát.")
            else:
                st.markdown("---")
                st.markdown("### 🎬 Etsy Profi Prompt Csomag Generálása Folyamatban...")

                items_to_process = parsed_etsy_items[:item_limit_etsy]
                p_bar = st.progress(0)
                status_etsy_b = st.empty()

                compiled_prompt_blocks = []
                results_records = []

                for b_idx, b_item in enumerate(items_to_process):
                    p_bar.progress((b_idx + 1) / len(items_to_process))
                    status_etsy_b.markdown(f"**⚙️ [{b_idx+1}/{len(items_to_process)}] Prompt készítése:** `{b_item['title']}`...")

                    it_text = b_item["visual_prompt"]
                    if is_wall_art_batch:
                        base_p = template_etsy_wall_art(it_text)
                        req_note = " TARGET: Premium Scandinavian Christian Wall Art Poster, elegant typography, modern boho botanical accents, museum quality, 4:5 aspect ratio."
                    else:
                        base_p = template_etsy_clipart(it_text)
                        req_note = " TARGET: Clean premium transparent clipart illustration, soft watercolor and fine ink outline, cute aesthetic, 1:1 aspect ratio."

                    if client:
                        final_p = enhance_prompt_with_gemini(client, base_p, req_note, text_model, temperature, active_art_style_prompt, current_image_model)
                    else:
                        final_p = base_p + (f"\n\n[Style: {active_art_style_prompt}]" if active_art_style_prompt else "")

                    compiled_prompt_blocks.append(f"=== {b_item['index']}. {b_item['title']} ===")
                    compiled_prompt_blocks.append(final_p.strip())
                    compiled_prompt_blocks.append("")

                    results_records.append({
                        "title": b_item["title"],
                        "prompt": final_p
                    })

                full_etsy_batch_txt = "\n".join(compiled_prompt_blocks).strip()
                st.session_state["etsy_batch_full_result"] = full_etsy_batch_txt
                st.session_state["etsy_batch_records"] = results_records
                st.session_state["etsy_batch_source_name"] = batch_etsy_source_name
                status_etsy_b.empty()
                st.success(f"🎉 **Elkészült!** Mind a {len(items_to_process)} tétel profi képgeneráló promptja legenerálva!")

        if st.session_state.get("etsy_batch_full_result"):
            st.markdown("---")
            st.markdown("### 📋 Generált Etsy Profi Prompt Csomag")
            
            st.text_area(
                "📋 Teljes Prompt Csomag (Ctrl+A → Ctrl+C):",
                value=st.session_state["etsy_batch_full_result"],
                height=300,
                key="area_etsy_batch_full"
            )

            col_eb_save, col_eb_goto = st.columns([1, 1])
            with col_eb_save:
                theme_name = st.session_state.get("etsy_batch_source_name", "Etsy_Csomag")
                if st.button("💾 Kész Etsy Csomag Mentése Google Drive-ra (04_🖼️_ETSY_DIGITAL)", key="btn_save_etsy_batch_drive", use_container_width=True):
                    ok_s, res_path = save_prompts_file_to_drive(
                        "etsy",
                        f"Etsy_Profi_Csomag_{theme_name}",
                        st.session_state["etsy_batch_full_result"],
                        header_info="Etsy Profi Prompt Csomag (Tömeges)"
                    )
                    if ok_s:
                        st.success(f"💾 **Sikeresen elmentve a Google Drive-ra!**\n\n`{res_path}`\n\nA mentett fájlt most már közvetlenül betöltheted és legenerálhatod a **🖼️ Tömeges Képgeneráló Stúdió** fülön!")
                    else:
                        st.error(f"Hiba a mentéskor: {res_path}")

            with col_eb_goto:
                if st.button("🚀 Tovább a Tömeges Képgeneráló Stúdióhoz", key="btn_goto_tab6_etsy", use_container_width=True):
                    st.session_state["main_menu_category_nav"] = "🤖 AI Alkotó Stúdiók & Generálók (Amazon, Etsy, Gumroad...)"
                    st.session_state["ai_workspace_select"] = "🖼️ Tömeges Képgeneráló Stúdió"
                    st.rerun()


# ══════════════════════════════════════════════════════════
# WORKSPACE: GUMROAD ÁHÍTATOK & SEO MARKETING
# ══════════════════════════════════════════════════════════

elif "Gumroad" in menu_choice or "3." in menu_choice:
    st.markdown("<div class='path-badge'>✍️ Gumroad Munkaterület</div>", unsafe_allow_html=True)
    st.markdown("### ✍️ Keresztény Áhítat Naplók & Etsy SEO Marketing")
    st.caption("Írj mély, KJV igealapú áhítatokat és generálj vásárlásösztönző Etsy termékleírásokat.")

    active_art_style_prompt = render_style_selector("ws_gumroad")
    render_niche_status_bar("gumroad")

    gum_sub = st.radio(
        "Mit szeretnél generálni?",
        ["✝️ 30 Napos Áhítat Szöveg (PLR Csomag)", "🛍️ Etsy SEO & Értékesítési Termékleírás"],
        horizontal=True,
        key="gum_sub_choice"
    )

    st.markdown("---")

    if "Áhítat" in gum_sub:
        col_form3, col_out3 = st.columns([1, 1.05], gap="large")

        with col_form3:
            n_slug_gum = get_niche_slug(chosen_niche)
            st.markdown("<div class='step-label'>Áhítat / Napló Paraméterek</div>", unsafe_allow_html=True)
            audience_dev = st.text_input("Célközönség", value=st.session_state.get(f"aud_dev_{n_slug_gum}", get_niche_field("ffc_aud", chosen_niche)), key=f"aud_dev_{n_slug_gum}")
            theme_dev = st.text_input("Áhítat / Napló téma", value=st.session_state.get(f"theme_dev_{n_slug_gum}", get_niche_field("gum_topic", chosen_niche)), key=f"theme_dev_{n_slug_gum}")
            day_dev = st.selectbox("Nap száma", [str(i) for i in range(1, 31)], index=0, key=f"day_dev_{n_slug_gum}")
            extra_dev = st.text_area("➕ Extra utasítás (opcionális)", placeholder="Pl.: Focus on daily reflection.", height=80, key=f"extra_dev_{n_slug_gum}")

            base3 = template_devotional(audience_dev, theme_dev, day_dev)
            generate_dev = st.button("✨ Áhítat Prompt Generálása", key="gen_dev", use_container_width=True)

        with col_out3:
            st.markdown("<div class='step-label'>Kész Prompt — Másold be a Gemini-be</div>", unsafe_allow_html=True)
            if generate_dev:
                if client:
                    with st.spinner("Gemini feljavítja az áhítat promptot..."):
                        res3 = enhance_prompt_with_gemini(client, base3, extra_dev, text_model, temperature, active_art_style_prompt, current_image_model)
                else:
                    res3 = base3
                st.session_state["dev_res"] = res3

            if st.session_state.get("dev_res"):
                st.text_area("📋 Másold innen (Ctrl+A → Ctrl+C):", value=st.session_state["dev_res"], height=240, key="dev_copy_tab4")
                st.markdown("---")
                if st.button("💾 Áhítat mentése Google Drive-ra (05_📖_GUMROAD_PLR)", key="btn_save_dev_drive", use_container_width=True):
                    ok_s, res_path = save_prompts_file_to_drive("gumroad", f"{theme_dev}_Day_{day_dev}", st.session_state["dev_res"], header_info=f"Gumroad Devotional - {theme_dev} - Day {day_dev}")
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{res_path}`")
                    else:
                        st.error(f"Hiba a mentéskor: {res_path}")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>✝️</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Töltsd ki az adatokat a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">✨ Áhítat Prompt Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    else:  # SEO & Marketing
        col_form4, col_out4 = st.columns([1, 1.1], gap="large")
        with col_form4:
            st.markdown("<div class='step-label'>Lépés 1 — Termék & SEO Paraméterek</div>", unsafe_allow_html=True)
            prod_title = st.text_input(
                "Termék Címe / Fő Témája:",
                value="Christian Adult Coloring Book Printable PDF",
                key="seo_prod_title",
                help="Pl.: Psalm 23 Eucalyptus Wall Art, Noah Ark Coloring Book"
            )
            col_tp1, col_tp2 = st.columns(2)
            with col_tp1:
                prod_type = st.selectbox(
                    "Terméktípus:",
                    [
                        "🖼️ Falikép Nyomat (Wall Art)",
                        "✂️ Clipart & Matrica Csomag",
                        "🖍️ Színezőkönyv (Coloring Book)",
                        "📖 Áhítat & Napló (Devotional Journal)",
                        "✨ Egyedi Digitális Termék"
                    ],
                    index=0,
                    key="seo_prod_type"
                )
            with col_tp2:
                prod_price = st.number_input("Ár (USD $):", min_value=0.0, max_value=999.0, value=6.99, step=0.5, key="seo_prod_price")

            prod_drive_url = st.text_input(
                "Google Drive PDF Kézbesítési Link (opcionális):",
                value="https://drive.google.com/drive/folders/EXAMPLE_LINK",
                key="seo_drive_url",
                help="Ezt a linket kapja meg a vásárló a számlán és a termékleírásban a 300 DPI fájlok eléréséhez."
            )
            prod_extra = st.text_area(
                "➕ Extra kulcsszavak / stílusjegyzetek (opcionális):",
                placeholder="Pl.: Scandinavian sage green palette, modern minimalist serif fonts.",
                height=70,
                key="seo_extra_keywords"
            )

            btn_gen_seo = st.button("✨ Szigorú Etsy SEO Csomag Generálása (140-char Title + 13 Tags + FFC)", key="btn_gen_seo", use_container_width=True)

        with col_out4:
            st.markdown("<div class='step-label'>Lépés 2 — Generált Etsy SEO & Export Csomag</div>", unsafe_allow_html=True)
            
            if btn_gen_seo:
                km = get_key_manager()
                with st.spinner("AI generálja a szigorú Etsy SEO csomagot..."):
                    p_seo = build_strict_etsy_seo_prompt(
                        product_title=prod_title,
                        product_type=prod_type,
                        niche_name=st.session_state.get("active_niche_choice", ""),
                        extra_details=prod_extra
                    )
                    ok_seo, res_raw = km.generate_text_with_fallback(prompt=p_seo, model_name=text_model)
                    if ok_seo and res_raw:
                        parsed_seo = parse_strict_etsy_seo_output(res_raw)
                        st.session_state["seo_parsed_data"] = parsed_seo
                        st.session_state["seo_raw_res"] = res_raw
                    else:
                        st.error(f"Hiba a generáláskor: {res_raw}")

            if st.session_state.get("seo_parsed_data"):
                data = st.session_state["seo_parsed_data"]
                curr_t = data.get("title", "")
                curr_tags = data.get("tags", [])
                curr_desc = data.get("description", "")

                t_len = len(curr_t)
                t_color = "#34d399" if t_len <= 140 else "#ef4444"
                st.markdown(f"**🏷️ Front-Loaded Etsy Cím** `<span style='color:{t_color}; font-weight:600;'>({t_len}/140 karakter)</span>`", unsafe_allow_html=True)
                st.code(curr_t, language="text")

                st.markdown(f"**🏷️ 13 db Etsy Keresőcímke (Tags)** `<span style='color:#34d399; font-weight:600;'>({len(curr_tags)} db · egyenként max 20 kar.)</span>`", unsafe_allow_html=True)
                tag_chips_html = " ".join([
                    f"<span style='display:inline-block; background:#1e293b; color:#93c5fd; border:1px solid #3b82f6; border-radius:14px; padding:3px 10px; margin:3px; font-size:0.78rem;'>{t} ({len(t)}/20)</span>"
                    for t in curr_tags
                ])
                st.markdown(f"<div style='margin-bottom:8px;'>{tag_chips_html}</div>", unsafe_allow_html=True)
                st.text_area("📋 Másolható tagek (vesszővel):", value=", ".join(curr_tags), height=65, key="seo_tags_copy_ws3")

                st.markdown("**📝 FFC Termékleírás (Drive Kézbesítés + AI Transzparencia):**")
                st.text_area("📋 Teljes Leírás:", value=curr_desc, height=180, key="seo_desc_copy_ws3")

                st.markdown("---")
                col_exp1, col_exp2 = st.columns(2)
                
                with col_exp1:
                    if st.button("📥 Etsy Import-Kész CSV Exportálása", key="btn_export_etsy_csv_ws3", use_container_width=True):
                        drive_etsy_dir = resolve_drive_folder("etsy")
                        ts_csv = time.strftime("%Y%m%d_%H%M%S")
                        csv_filename = f"Etsy_Listing_{sanitize_filename(curr_t)}_{ts_csv}.csv"
                        csv_filepath = os.path.join(drive_etsy_dir, csv_filename)
                        
                        listing_payload = [{
                            "title": curr_t,
                            "tags": curr_tags,
                            "description": curr_desc,
                            "price": f"{prod_price:.2f}",
                            "quantity": "999",
                            "drive_url": prod_drive_url,
                            "section": prod_type
                        }]
                        ok_csv, msg_csv, csv_b = generate_etsy_csv(listing_payload, output_path=csv_filepath)
                        if ok_csv:
                            st.session_state["last_etsy_csv_bytes"] = csv_b
                            st.session_state["last_etsy_csv_filename"] = csv_filename
                            st.session_state["last_etsy_csv_path"] = csv_filepath
                            st.balloons()
                            st.success(f"🎉 **Etsy CSV sikeresen elmentve:**\n\n`{csv_filepath}`")

                    if st.session_state.get("last_etsy_csv_bytes"):
                        st.download_button(
                            label="⬇️ CSV Fájl Letöltése",
                            data=st.session_state["last_etsy_csv_bytes"],
                            file_name=st.session_state.get("last_etsy_csv_filename", "Etsy_Listing.csv"),
                            mime="text/csv",
                            key="btn_dl_etsy_csv_ws3",
                            use_container_width=True
                        )

                with col_exp2:
                    if st.button("🚀 1-Kattintásos Publikálás Gumroadra", key="btn_pub_gumroad_ws3", use_container_width=True):
                        with st.spinner("Termék feltöltése a Gumroad API-ra..."):
                            ok_gum, res_gum_url, raw_gum = publish_to_gumroad(
                                product_name=curr_t,
                                price_usd=prod_price,
                                description=curr_desc,
                                drive_delivery_url=prod_drive_url
                            )
                            if ok_gum:
                                st.balloons()
                                st.session_state["last_gumroad_url"] = res_gum_url
                                st.success(f"🎉 **Termék sikeresen publikálva a Gumroadon!**\n\n🔗 **Termék URL:** [{res_gum_url}]({res_gum_url})")
                            else:
                                st.error(f"❌ {res_gum_url}")

                    if st.session_state.get("last_gumroad_url"):
                        st.markdown(
                            f"""
                            <div style='background: linear-gradient(135deg, rgba(244,63,94,0.15), rgba(251,113,133,0.10)); border: 1px solid #f43f5e; border-radius: 8px; padding: 8px 12px; margin-top: 5px;'>
                                🔗 <b>Gumroad Élő Link:</b> <a href='{st.session_state["last_gumroad_url"]}' target='_blank' style='color:#fda4af; font-weight:bold;'>{st.session_state["last_gumroad_url"]}</a>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🛍️</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg a termék címét és árát, majd<br>kattints a <strong style="color:#34d399;">✨ Szigorú Etsy SEO Csomag Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# WORKSPACE: NOTEBOOKLM RAG KÖZPONT & REVIEW MINING
# ══════════════════════════════════════════════════════════

elif "NotebookLM" in menu_choice:
    if render_notebooklm_rag_module:
        render_notebooklm_rag_module()
    else:
        st.error("A NotebookLM RAG modul betöltése sikertelen.")

# ══════════════════════════════════════════════════════════
# WORKSPACE: 2026-OS MAGYAR ÁTALÁNYADÓ & PÉNZÜGYI TERVEZŐ
# ══════════════════════════════════════════════════════════

elif "Átalányadó" in menu_choice or "2026" in menu_choice:
    if render_tax_calculator_2026_module:
        render_tax_calculator_2026_module()
    else:
        st.error("A 2026-os Átalányadó modul betöltése sikertelen.")

# ══════════════════════════════════════════════════════════
# WORKSPACE: 30 TÉMA & ÖTLETGENERÁLÓ MŰHELY
# ══════════════════════════════════════════════════════════

elif "Ötletgeneráló" in menu_choice or "30 Téma" in menu_choice or "4." in menu_choice:
    st.markdown("<div class='path-badge'>💡 Ötletbörze & Etsy SEO Stúdió</div>", unsafe_allow_html=True)
    st.markdown("### 💡 30 Téma, Ötletgeneráló & Szigorú Etsy CSV Exportőr")
    st.caption("Készíts 30 tételes ötletlistát KDP színezőkhöz, Etsy faliképekhez vagy exportálj hivatalos, szigorú Etsy SEO CSV listázásokat és publikálj Gumroadra.")

    active_art_style_prompt = render_style_selector("ws_ideas")
    render_niche_status_bar("ideas")

    tab_ideas1, tab_ideas_kdp_batch, tab_ideas2 = st.tabs([
        "💡 1. 30 Téma & Ötletgeneráló Műhely",
        "📦 2. Tömeges KDP Prompt Csomag (30 Ötlet alapján)",
        "🛍️ 3. Szigorú Etsy SEO & Tömeges CSV Exportőr"
    ])

    with tab_ideas1:
        idea_category = st.radio(
            "Milyen termékhez szeretnél 30 ötletet generálni?",
            [
                "🖍️ 30 Bibliai Színező Jelenet (KDP Színezőhöz)",
                "🖼️ 30 Igés Falikép Idézet (Etsy Wall Art-hoz)",
                "✂️ 30 Clipart Karakter / Tárgy Ötlet (Etsy Clipart-hoz)",
                "✝️ 30 Napos Áhítat Tématerv (Gumroad PLR-hez)"
            ],
            key="idea_category"
        )

        st.markdown("---")
        col_i_form, col_i_out = st.columns([1, 1.05], gap="large")

        with col_i_form:
            n_slug_ideas = get_niche_slug(chosen_niche)
            if "Színező" in idea_category:
                st.markdown("<div class='step-label'>Színező Téma Paraméterek</div>", unsafe_allow_html=True)
                aud_b = st.text_input("Célközönség", value=st.session_state.get(f"aud_b_{n_slug_ideas}", get_niche_field("ffc_aud", chosen_niche)), key=f"aud_b_{n_slug_ideas}")
                thm_b = st.text_input("Főtéma", value=st.session_state.get(f"thm_b_{n_slug_ideas}", get_niche_field("kdp_theme", chosen_niche)), key=f"thm_b_{n_slug_ideas}")
                current_idea_prompt = template_idea_kdp(aud_b, thm_b)

            elif "Falikép" in idea_category:
                st.markdown("<div class='step-label'>Falikép Idézet Paraméterek</div>", unsafe_allow_html=True)
                thm_b = st.text_input("Főtéma / Hangulat", value=st.session_state.get(f"thm_wall_b_{n_slug_ideas}", get_niche_field("verse_etsy", chosen_niche)), key=f"thm_wall_b_{n_slug_ideas}")
                current_idea_prompt = template_idea_wall_art(thm_b)

            elif "Clipart" in idea_category:
                st.markdown("<div class='step-label'>Clipart Csomag Paraméterek</div>", unsafe_allow_html=True)
                thm_b = st.text_input("Csomag megnevezése / Témája", value=st.session_state.get(f"thm_clip_b_{n_slug_ideas}", get_niche_field("subject_etsy", chosen_niche)), key=f"thm_clip_b_{n_slug_ideas}")
                current_idea_prompt = template_idea_clipart(thm_b)

            else:  # Devotional
                st.markdown("<div class='step-label'>Áhítat / Napló Tématerv Paraméterek</div>", unsafe_allow_html=True)
                aud_b = st.text_input("Célközönség", value=st.session_state.get(f"aud_dev_b_{n_slug_ideas}", get_niche_field("ffc_aud", chosen_niche)), key=f"aud_dev_b_{n_slug_ideas}")
                thm_b = st.text_input("Áhítat / Napló Főcíme", value=st.session_state.get(f"thm_dev_b_{n_slug_ideas}", get_niche_field("gum_topic", chosen_niche)), key=f"thm_dev_b_{n_slug_ideas}")
                current_idea_prompt = template_idea_devotional(aud_b, thm_b)

            st.markdown("<div class='step-label'>📋 Másolható Ötletelő Prompt</div>", unsafe_allow_html=True)
            st.code(current_idea_prompt, language="text")

            gen_ideas_btn = st.button("🚀 30 Ötlet Generálása Élőben (AI)", key="gen_ideas_btn", use_container_width=True)

        with col_i_out:
            st.markdown("<div class='step-label'>Generált 30 Ötlet Eredménye</div>", unsafe_allow_html=True)

            if "b_ideas_result" not in st.session_state:
                st.session_state.b_ideas_result = ""

            if gen_ideas_btn:
                curr_tb = time.strftime("%H:%M:%S")
                if client:
                    with st.spinner("AI gyűjti a 30 ötletet..."):
                        res_b = generate_ideas_live(client, current_idea_prompt, text_model, temperature)
                else:
                    res_b = (
                        "⚠️ Nincs aktív AI szolgáltató konfigurálva az élő generáláshoz!\n\n"
                        "Állíts be egy API kulcsot a Rendszerbeállításokban!"
                    )

                st.session_state["b_ideas_result"] = res_b
                st.session_state["b_ideas_copy_box"] = res_b
                st.session_state["b_ideas_time"] = curr_tb

            if st.session_state.get("b_ideas_result"):
                time_tag_b = f" (Frissítve: {st.session_state.b_ideas_time})" if "b_ideas_time" in st.session_state else ""
                st.text_area(
                    f"📋 Másold ki az ötleteket{time_tag_b}:",
                    value=st.session_state.get("b_ideas_copy_box", st.session_state.b_ideas_result),
                    height=380,
                    key="b_ideas_copy_box"
                )
                st.markdown("---")

                if "Színező" in idea_category:
                    drive_cat_key = "kdp"
                    target_folder_name = "03_📚_AMAZON_KDP"
                    theme_val = thm_b if 'thm_b' in locals() else "KDP_30_Jelenet"
                elif "Falikép" in idea_category:
                    drive_cat_key = "etsy"
                    target_folder_name = "04_🖼️_ETSY_DIGITAL"
                    theme_val = thm_b if 'thm_b' in locals() else "Etsy_30_Falikepek"
                elif "Clipart" in idea_category:
                    drive_cat_key = "etsy"
                    target_folder_name = "04_🖼️_ETSY_DIGITAL"
                    theme_val = thm_b if 'thm_b' in locals() else "Etsy_30_Clipart"
                else:
                    drive_cat_key = "gumroad"
                    target_folder_name = "05_📖_GUMROAD_PLR"
                    theme_val = thm_b if 'thm_b' in locals() else "Gumroad_30_Ahitatterv"

                if st.button(f"💾 30 Ötlet mentése Google Drive-ra ({target_folder_name})", key="btn_save_ideas_drive", use_container_width=True):
                    ok_s, res_path = save_prompts_file_to_drive(
                        drive_cat_key,
                        f"30_Otlet_{theme_val}",
                        st.session_state["b_ideas_result"],
                        header_info=f"30 Ötlet Gyűjtemény - {idea_category}"
                    )
                    if ok_s:
                        st.success(f"💾 **30 Ötlet sikeresen elmentve:**\n\n`{res_path}`")
                    else:
                        st.error(f"Hiba a mentéskor: {res_path}")

                st.markdown("##### 🚀 Következő Lépés a Munkafolyamatban:")
                col_wf1, col_wf2 = st.columns(2)
                with col_wf1:
                    if "Színező" in idea_category:
                        if st.button("👉 Amazon KDP — 30 Profi Prompt Generálása", key="btn_wf_goto_kdp", use_container_width=True):
                            st.session_state["ai_workspace_select"] = "🎨 Amazon KDP Színező & PDF Összeállító"
                            st.rerun()
                    elif "Falikép" in idea_category or "Clipart" in idea_category:
                        if st.button("👉 Etsy Grafika — 30 Profi Prompt Generálása", key="btn_wf_goto_etsy", use_container_width=True):
                            st.session_state["ai_workspace_select"] = "🎨 Etsy Faliképek & Clipartok"
                            st.rerun()
                    else:
                        if st.button("👉 Gumroad Műhely — Áhítatok Írása", key="btn_wf_goto_gum", use_container_width=True):
                            st.session_state["ai_workspace_select"] = "✍️ Gumroad Áhítatok & SEO"
                            st.rerun()

                with col_wf2:
                    if st.button("🖼️ Tömeges Képgeneráló Stúdió Megnyitása", key="btn_wf_goto_bulk", use_container_width=True):
                        st.session_state["ai_workspace_select"] = "🖼️ Gemini Képbegyűjtő & PDF Központ"
                        st.rerun()

                st.success(f"✅ Ötletek elkészültek{time_tag_b}! A fenti gombokkal azonnal folytathatod a promptok generálását vagy a tömeges képgenerálást!")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #cbd5e1; border-radius: 16px; background: #ffffff; color: #64748b;'>
                    <div style='font-size: 2.8rem;'>💡</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#334155;'>Állítsd be a paramétereket a bal oldalon, majd<br>kattints a <strong style="color:#166534;">🚀 30 Ötlet Generálása Élőben</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ── TAB 2: TÖMEGES KDP PROMPT CSOMAG ──
    with tab_ideas_kdp_batch:
        st.markdown("#### 📦 Tömeges KDP Prompt Csomag Generáló (Google Drive & 30 Ötlet)")
        st.caption("Generálj le egy teljes, 30 tételes promptcsomagot a Gemini Webes munkához, majd mentsd el a Google Drive-ra!")

        col_kb_src, col_kb_cfg = st.columns([1, 1], gap="large")

        with col_kb_src:
            kdp_batch_src_mode = st.radio(
                "Honnan szeretnéd betölteni az ötleteket?",
                [
                    "⚡ 1. Legutóbbi 30 Ötlet betöltése (1. Fülről)",
                    "📁 2. Google Drive .txt fájl betöltése (03_📚_AMAZON_KDP)",
                    "📤 3. Saját .txt fájl feltöltése",
                    "✍️ 4. Kézi lista beillesztése"
                ],
                key="kdp_batch_src_mode_v2"
            )

            loaded_kdp_batch_text = ""
            batch_source_name = "KDP_30_Jelenet"

            if "Legutóbbi" in kdp_batch_src_mode:
                last_ideas = st.session_state.get("b_ideas_result", "")
                if last_ideas and not last_ideas.startswith("⚠️"):
                    loaded_kdp_batch_text = last_ideas
                    st.success("⚡ 30 Ötlet sikeresen átvéve!")
                else:
                    st.info("ℹ️ Még nem generáltál ötletlistát az 1. fülön.")
            elif "Google Drive" in kdp_batch_src_mode:
                kdp_drive_dir = resolve_drive_folder("kdp")
                txt_files_kdp = [f for f in os.listdir(kdp_drive_dir) if f.lower().endswith(".txt")] if os.path.exists(kdp_drive_dir) else []
                if txt_files_kdp:
                    selected_kdp_txt = st.selectbox("Válassz .txt fájlt:", txt_files_kdp, key="sel_kdp_b_txt_v2")
                    if selected_kdp_txt:
                        batch_source_name = selected_kdp_txt.replace(".txt", "")
                        with open(os.path.join(kdp_drive_dir, selected_kdp_txt), "r", encoding="utf-8") as tf:
                            loaded_kdp_batch_text = tf.read()
                        st.success(f"📄 Betöltve: `{selected_kdp_txt}`")
            elif "feltöltése" in kdp_batch_src_mode:
                up_kdp_b = st.file_uploader("Tölts fel egy .txt fájlt:", type=["txt"], key="up_kdp_b_v2")
                if up_kdp_b:
                    loaded_kdp_batch_text = up_kdp_b.getvalue().decode("utf-8", errors="ignore")
                    batch_source_name = up_kdp_b.name.replace(".txt", "")
            else:
                loaded_kdp_batch_text = st.text_area("Illeszd be a jeleneteket:", value="1. Noah building the ark\n2. Animals entering the ark\n3. Rainbow covenant", height=140, key="man_kdp_b_v2")

            parsed_kdp_items = parse_prompts_from_text(loaded_kdp_batch_text)
            st.markdown(f"**🎯 Felismert tételek:** `{len(parsed_kdp_items)} db`")

        with col_kb_cfg:
            batch_edition = st.radio("Célközönség:", ["🧒 Gyermek Kiadás", "🧘 Felnőtt Kiadás"], key="kdp_batch_ed_v2")
            batch_is_adult = "Felnőtt" in batch_edition
            batch_gen_companion = st.checkbox("🎨 Kapcsolódó Bal Oldali Lapok Generálása is", value=True, key="kdp_batch_comp_v2")

            if st.button("✨ Teljes KDP Prompt Csomag Generálása (AI)", key="btn_run_kdp_batch_v2", use_container_width=True):
                if not parsed_kdp_items:
                    st.error("⚠️ Nincsenek feldolgozható jelenetek!")
                else:
                    with st.spinner("Promptok előállítása..."):
                        blocks = []
                        for b_it in parsed_kdp_items:
                            sc_text = b_it["visual_prompt"]
                            base_p = template_kdp_adult_coloring(sc_text) if batch_is_adult else template_kdp_coloring(sc_text)
                            req = " TARGET: Adult zentangle mandala line art" if batch_is_adult else " TARGET: Children bold line art"
                            final_p = enhance_prompt_with_gemini(True, base_p, req, text_model, temperature, active_art_style_prompt, "")
                            blocks.append(f"=== {b_it['index']}. {b_it['title']} ===\n{final_p.strip()}\n")
                        full_txt = "\n".join(blocks)
                        st.session_state["kdp_batch_full_result"] = full_txt
                        st.success("🎉 Teljes csomag sikeresen elkészült!")

        if st.session_state.get("kdp_batch_full_result"):
            st.markdown("### 📋 Generált KDP Prompt Csomag:")
            st.code(st.session_state["kdp_batch_full_result"], language="text")
            if st.button("💾 Mentés Google Drive-ra (.txt)", key="btn_save_b_drive_v2", use_container_width=True):
                ok_s, res_p = save_prompts_file_to_drive("kdp", "KDP_Profi_Csomag", st.session_state["kdp_batch_full_result"])
                if ok_s:
                    st.success(f"💾 Sikeresen mentve: `{res_p}`")

    with tab_ideas2:
        st.markdown("#### 🛍️ Szigorú Etsy SEO & Tömeges CSV Exportőr")
        st.caption("Készíts az Etsy hivatalos szabályzatának (max. 140 kar. cím, 13 db max. 20 kar. keresőcímke, FFC leírás Drive linkkel és AI nyilatkozattal) megfelelő listázást, exportáld CSV-be vagy publikáld 1 kattintással Gumroadra!")

        col_e4_form, col_e4_out = st.columns([1, 1.1], gap="large")

        with col_e4_form:
            st.markdown("<div class='step-label'>Lépés 1 — Termék & SEO Paraméterek</div>", unsafe_allow_html=True)
            e4_prod_title = st.text_input(
                "Termék Címe / Fő Témája:",
                value="Christian Scripture Wall Art Printable Psalm 23 Eucalyptus Botanical",
                key="e4_prod_title",
                help="Pl.: Psalm 23 Eucalyptus Wall Art, Noah Ark Kids Coloring Book"
            )
            col_e4_t1, col_e4_t2 = st.columns(2)
            with col_e4_t1:
                e4_prod_type = st.selectbox(
                    "Terméktípus:",
                    [
                        "🖼️ Falikép Nyomat (Wall Art)",
                        "✂️ Clipart & Matrica Csomag",
                        "🖍️ Színezőkönyv (Coloring Book)",
                        "📖 Áhítat & Napló (Devotional Journal)",
                        "✨ Egyedi Digitális Termék"
                    ],
                    index=0,
                    key="e4_prod_type"
                )
            with col_e4_t2:
                e4_prod_price = st.number_input("Ár (USD $):", min_value=0.0, max_value=999.0, value=6.99, step=0.5, key="e4_prod_price")

            e4_drive_url = st.text_input(
                "Google Drive PDF Kézbesítési Link (opcionális):",
                value="https://drive.google.com/drive/folders/EXAMPLE_LINK",
                key="e4_drive_url",
                help="A vásárló ezt a letöltési linket kapja meg a 300 DPI fájlokhoz."
            )
            e4_extra = st.text_area(
                "➕ Extra kulcsszavak / stílusjegyzetek (opcionális):",
                placeholder="Pl.: Scandinavian sage green palette, modern minimalist serif fonts.",
                height=70,
                key="e4_extra_keywords"
            )

            btn_gen_e4_seo = st.button("✨ Szigorú Etsy SEO Csomag Generálása (140-char Title + 13 Tags + FFC)", key="btn_gen_e4_seo", use_container_width=True)

        with col_e4_out:
            st.markdown("<div class='step-label'>Lépés 2 — Generált Etsy SEO & Export Csomag</div>", unsafe_allow_html=True)
            
            if btn_gen_e4_seo:
                km = get_key_manager()
                with st.spinner("AI generálja a szigorú Etsy SEO csomagot..."):
                    p_seo = build_strict_etsy_seo_prompt(
                        product_title=e4_prod_title,
                        product_type=e4_prod_type,
                        niche_name=st.session_state.get("active_niche_choice", ""),
                        extra_details=e4_extra
                    )
                    ok_seo, res_raw = km.generate_text_with_fallback(prompt=p_seo, model_name=text_model)
                    if ok_seo and res_raw:
                        parsed_seo = parse_strict_etsy_seo_output(res_raw)
                        st.session_state["e4_seo_parsed"] = parsed_seo
                    else:
                        st.error(f"Hiba a generáláskor: {res_raw}")

            if st.session_state.get("e4_seo_parsed"):
                data = st.session_state["e4_seo_parsed"]
                curr_t = data.get("title", "")
                curr_tags = data.get("tags", [])
                curr_desc = data.get("description", "")

                t_len = len(curr_t)
                t_color = "#34d399" if t_len <= 140 else "#ef4444"
                st.markdown(f"**🏷️ Front-Loaded Etsy Cím** `<span style='color:{t_color}; font-weight:600;'>({t_len}/140 karakter)</span>`", unsafe_allow_html=True)
                st.code(curr_t, language="text")

                st.markdown(f"**🏷️ 13 db Etsy Keresőcímke (Tags)** `<span style='color:#34d399; font-weight:600;'>({len(curr_tags)} db · egyenként max 20 kar.)</span>`", unsafe_allow_html=True)
                tag_chips_html = " ".join([
                    f"<span style='display:inline-block; background:#1e293b; color:#93c5fd; border:1px solid #3b82f6; border-radius:14px; padding:3px 10px; margin:3px; font-size:0.78rem;'>{t} ({len(t)}/20)</span>"
                    for t in curr_tags
                ])
                st.markdown(f"<div style='margin-bottom:8px;'>{tag_chips_html}</div>", unsafe_allow_html=True)
                st.text_area("📋 Másolható tagek (vesszővel):", value=", ".join(curr_tags), height=65, key="e4_tags_copy")

                st.markdown("**📝 FFC Termékleírás (Drive Kézbesítés + AI Transzparencia):**")
                st.text_area("📋 Teljes Leírás:", value=curr_desc, height=180, key="e4_desc_copy")

                st.markdown("---")
                col_e4_exp1, col_e4_exp2 = st.columns(2)
                
                with col_e4_exp1:
                    if st.button("📥 Etsy Import-Kész CSV Exportálása", key="btn_export_etsy_csv_tab4", use_container_width=True):
                        drive_etsy_dir = resolve_drive_folder("etsy")
                        ts_csv = time.strftime("%Y%m%d_%H%M%S")
                        csv_filename = f"Etsy_Listing_{sanitize_filename(curr_t)}_{ts_csv}.csv"
                        csv_filepath = os.path.join(drive_etsy_dir, csv_filename)
                        
                        listing_payload = [{
                            "title": curr_t,
                            "tags": curr_tags,
                            "description": curr_desc,
                            "price": f"{e4_prod_price:.2f}",
                            "quantity": "999",
                            "drive_url": e4_drive_url,
                            "section": e4_prod_type
                        }]
                        ok_csv, msg_csv, csv_b = generate_etsy_csv(listing_payload, output_path=csv_filepath)
                        if ok_csv:
                            st.session_state["tab4_etsy_csv_bytes"] = csv_b
                            st.session_state["tab4_etsy_csv_filename"] = csv_filename
                            st.balloons()
                            st.success(f"🎉 **Etsy CSV sikeresen elmentve:**\n\n`{csv_filepath}`")

                    if st.session_state.get("tab4_etsy_csv_bytes"):
                        st.download_button(
                            label="⬇️ CSV Fájl Letöltése",
                            data=st.session_state["tab4_etsy_csv_bytes"],
                            file_name=st.session_state.get("tab4_etsy_csv_filename", "Etsy_Listing.csv"),
                            mime="text/csv",
                            key="btn_dl_etsy_csv_tab4",
                            use_container_width=True
                        )

                with col_e4_exp2:
                    if st.button("🚀 1-Kattintásos Publikálás Gumroadra", key="btn_pub_gumroad_tab4", use_container_width=True):
                        with st.spinner("Termék feltöltése a Gumroad API-ra..."):
                            ok_gum, res_gum_url, raw_gum = publish_to_gumroad(
                                product_name=curr_t,
                                price_usd=e4_prod_price,
                                description=curr_desc,
                                drive_delivery_url=e4_drive_url
                            )
                            if ok_gum:
                                st.balloons()
                                st.session_state["tab4_gumroad_url"] = res_gum_url
                                st.success(f"🎉 **Termék sikeresen publikálva a Gumroadon!**\n\n🔗 **Termék URL:** [{res_gum_url}]({res_gum_url})")
                            else:
                                st.error(f"❌ {res_gum_url}")

                    if st.session_state.get("tab4_gumroad_url"):
                        st.markdown(
                            f"""
                            <div style='background: linear-gradient(135deg, rgba(244,63,94,0.15), rgba(251,113,133,0.10)); border: 1px solid #f43f5e; border-radius: 8px; padding: 8px 12px; margin-top: 5px;'>
                                🔗 <b>Gumroad Élő Link:</b> <a href='{st.session_state["tab4_gumroad_url"]}' target='_blank' style='color:#fda4af; font-weight:bold;'>{st.session_state["tab4_gumroad_url"]}</a>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🛍️</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg a termék címét és árát, majd<br>kattints a <strong style="color:#34d399;">✨ Szigorú Etsy SEO Csomag Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# WORKSPACE: AI VISION MULTIMODÁLIS LAB
# ══════════════════════════════════════════════════════════

elif "AI Vision" in menu_choice or "Vision Lab" in menu_choice or "5." in menu_choice:
    st.markdown("<div class='path-badge'>📷 AI Vision Lab</div>", unsafe_allow_html=True)
    st.markdown("### 📷 Meglévő Képek Feltöltése & AI Képelemzés (Gemini Vision)")
    st.caption("Töltsd fel meglévő könyvborítódat vagy a színező lapjaidat: a Gemini Vision automatikusan kinyeri az igéket és pár-promptokat generál!")

    active_art_style_prompt = render_style_selector("ws_vision")
    render_niche_status_bar("vision")

    st.markdown("---")
    col_up1, col_up2 = st.columns([1, 1], gap="large")

    with col_up1:
        st.markdown("#### 📕 1. Meglévő Borítókép Feltöltése")
        uploaded_cover = st.file_uploader(
            "Töltsd fel a könyvborító képét (PNG / JPG)",
            type=["png", "jpg", "jpeg"],
            key="uploaded_cover"
        )
        if uploaded_cover:
            st.image(uploaded_cover, caption="Feltöltött borítókép", use_container_width=True)
            if st.button("⚡ Borítókép Elemzése & Címoldal/Színtesztelő Prompt Generálása", key="btn_analyze_cover"):
                if client:
                    with st.spinner("Gemini Vision elemzi a borítóképet..."):
                        cover_img_pil = Image.open(uploaded_cover)
                        tester_prompt_gen = analyze_cover_image_for_tester_prompt(client, cover_img_pil, text_model, temperature, active_art_style_prompt)
                        st.session_state["analyzed_cover_prompt"] = tester_prompt_gen
                else:
                    st.warning("⚠️ Az AI képelemzéshez adj meg egy Gemini API kulcsot a Beállításokban!")

        if st.session_state.get("analyzed_cover_prompt"):
            st.markdown("**🧪 Generált Címoldal & Színtesztelő Prompt:**")
            st.markdown(f"<div class='prompt-output'>{st.session_state['analyzed_cover_prompt']}</div>", unsafe_allow_html=True)
            st.text_area(
                "📋 Másold innen (Ctrl+A → Ctrl+C):",
                value=st.session_state['analyzed_cover_prompt'],
                height=180,
                key="copy_analyzed_cover"
            )

    with col_up2:
        st.markdown("#### 🖍️ 2. Meglévő Színező Képek Tömeges Feltöltése")
        uploaded_pages = st.file_uploader(
            "Töltsd fel a színező oldal képeit (akár mind a 32 db-ot)",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="uploaded_pages"
        )
        if uploaded_pages:
            st.info(f"📂 {len(uploaded_pages)} db színező kép feltöltve.")
            btn_col1, btn_col2 = st.columns(2)

            with btn_col1:
                run_guide = st.button(f"⚡ 1. Bal Oldali Színajánló Promptok", key="btn_analyze_pages")
            with btn_col2:
                run_adult_pair = st.button(f"🧘 2. Gyerek ➔ Felnőtt Pár-Promptok", key="btn_analyze_child_to_adult")

            if run_guide:
                if client:
                    results_list = []
                    prog_bar = st.progress(0)
                    status_text = st.empty()

                    for idx, page_file in enumerate(uploaded_pages):
                        status_text.write(f"⚙️ Kép elemzése ({idx+1}/{len(uploaded_pages)}): `{page_file.name}`...")
                        page_pil = Image.open(page_file)
                        res_item = analyze_coloring_image_for_guide_prompt(
                            client, page_pil, text_model, temperature,
                            index_num=idx+1, active_style=active_art_style_prompt,
                            status_widget=status_text, is_adult=False
                        )
                        res_item["filename"] = page_file.name
                        results_list.append(res_item)
                        prog_bar.progress((idx + 1) / len(uploaded_pages))
                        if idx < len(uploaded_pages) - 1:
                            time.sleep(2)

                    status_text.success(f"✅ Mind a {len(uploaded_pages)} színező kép elemzése sikeresen befejeződött!")
                    st.session_state["analyzed_pages_prompts"] = results_list
                else:
                    st.warning("⚠️ Az AI képelemzéshez adj meg egy Gemini API kulcsot a Beállításokban!")

            if run_adult_pair:
                if client:
                    adult_pair_results = []
                    prog_bar = st.progress(0)
                    status_text = st.empty()

                    for idx, page_file in enumerate(uploaded_pages):
                        status_text.write(f"🧘 Gyerek kép átültetése felnőtt kiadásba ({idx+1}/{len(uploaded_pages)}): `{page_file.name}`...")
                        page_pil = Image.open(page_file)
                        res_item = analyze_children_image_for_adult_pair_prompts(
                            client, page_pil, text_model, temperature,
                            index_num=idx+1, active_style=active_art_style_prompt,
                            status_widget=status_text
                        )
                        res_item["filename"] = page_file.name
                        adult_pair_results.append(res_item)
                        prog_bar.progress((idx + 1) / len(uploaded_pages))
                        if idx < len(uploaded_pages) - 1:
                            time.sleep(2)

                    status_text.success(f"✅ Mind a {len(uploaded_pages)} gyerek kép alapján elkészültek a Felnőtt Pár-Promptok!")
                    st.session_state["analyzed_adult_pair_prompts"] = adult_pair_results
                else:
                    st.warning("⚠️ Az AI képelemzéshez adj meg egy Gemini API kulcsot a Beállításokban!")

        if st.session_state.get("analyzed_pages_prompts"):
            pages_res = st.session_state["analyzed_pages_prompts"]
            st.markdown(f"**🎨 Generált Színajánló & Igehely Promptok ({len(pages_res)} db):**")

            all_combined_prompts = []
            for p in pages_res:
                all_combined_prompts.append(f"=== {p.get('filename', 'Page')}: {p.get('title', '')} ===")
                all_combined_prompts.append(p.get('prompt', ''))
                all_combined_prompts.append("")

            combined_text = "\n".join(all_combined_prompts)

            with st.expander("📋 ÖSSZES PROMPT EGYBEN (Ctrl+A → Ctrl+C egyben másoláshoz)", expanded=False):
                st.text_area("Összes prompt egyben:", value=combined_text, height=250, key="all_pages_combined_copy")

            for p_idx, item in enumerate(pages_res):
                with st.expander(f"🖼️ {p_idx+1}. Kép: {item.get('filename', '')} — {item.get('title', '')}", expanded=False):
                    st.code(item.get('prompt', ''), language="text")

        if st.session_state.get("analyzed_adult_pair_prompts"):
            adult_pair_res = st.session_state["analyzed_adult_pair_prompts"]
            st.markdown(f"**🧘 Generált Felnőtt Színező & Kalligráfia Pár-Promptok ({len(adult_pair_res)} db):**")

            all_adult_coloring_prompts = []
            for p in adult_pair_res:
                all_adult_coloring_prompts.append(f"=== {p.get('filename', 'Page')}: {p.get('title', '')} (Felnőtt Színező) ===")
                all_adult_coloring_prompts.append(p.get('adult_coloring_prompt', ''))
                all_adult_coloring_prompts.append("")

            combined_adult_text = "\n".join(all_adult_coloring_prompts)

            with st.expander("📋 ÖSSZES FELNŐTT SZÍNEZŐ PROMPT EGYBEN (Ctrl+A → Ctrl+C egyben másoláshoz)", expanded=True):
                st.text_area("Összes felnőtt színező prompt egyben:", value=combined_adult_text, height=250, key="all_adult_pair_combined_copy")

            for p_idx, item in enumerate(adult_pair_res):
                with st.expander(f"🖼️ {p_idx+1}. Kép: {item.get('filename', '')} — {item.get('title', '')}", expanded=False):
                    st.markdown("**🎨 Felnőtt Színező Oldal Prompt (Stained Glass Zentangle + Realisztikus Figurák):**")
                    st.code(item.get('adult_coloring_prompt', ''), language="text")
                    if item.get('calligraphy_prompt'):
                        st.markdown("**📜 Bal Oldali Felnőtt Bibliai Kalligráfia Prompt:**")
                        st.code(item.get('calligraphy_prompt', ''), language="text")


# ══════════════════════════════════════════════════════════
# WORKSPACE: TÖMEGES KÉPGENERÁLÓ STÚDIÓ
# ══════════════════════════════════════════════════════════

elif "Tömeges Képgeneráló" in menu_choice or "Gemini Képbegyűjtő" in menu_choice or "6." in menu_choice:
    st.markdown("<div class='path-badge'>🖼️ Gemini Képbegyűjtő & PDF Központ</div>", unsafe_allow_html=True)
    st.markdown("### 🖼️ Gemini Képbegyűjtő & Nyomdakész PDF Összeállító (Google Drive Sync)")
    st.caption("Olvasd be a mentett `.txt` promptfájlt, másold a promptokat a Gemini-be, majd gyűjtsd be az elkészült képeket a Google Drive-ról és fűzd össze nyomdakész PDF-fé!")

    active_art_style_prompt = render_style_selector("ws_bulk")
    render_niche_status_bar("bulk")

    col_src, col_cfg = st.columns([1, 1], gap="large")

    with col_src:
        st.markdown("#### 📂 1. Lépés: Prompt Forrás Kiválasztása")
        prompt_source_mode = st.radio(
            "Honnan szeretnéd betölteni a promptokat?",
            [
                "⚡ 1. Legutóbbi KDP / Etsy Generált Csomag Betöltése",
                "📁 2. Google Drive Mappából Tallózás",
                "📤 3. Saját .txt fájl feltöltése",
                "✍️ 4. Kézi prompt-lista beillesztése"
            ],
            key="imagen_source_mode"
        )

        loaded_prompt_text = ""
        source_name = "Kézi_bevitel"

        if "Legutóbbi" in prompt_source_mode:
            recent_kdp = st.session_state.get("kdp_batch_full_result", "")
            recent_etsy = st.session_state.get("etsy_batch_full_result", "")
            recent_ideas = st.session_state.get("b_ideas_result", "")

            avail_batches = []
            if recent_kdp:
                avail_batches.append("📖 Legutóbbi KDP Színező Csomag")
            if recent_etsy:
                avail_batches.append("🎨 Legutóbbi Etsy Grafika Csomag")
            if recent_ideas and not recent_ideas.startswith("⚠️"):
                avail_batches.append("💡 Legutóbbi 30 Ötlet Gyűjtemény")

            if avail_batches:
                chosen_batch_src = st.selectbox("Válassz a generált csomagok közül:", avail_batches, key="sel_avail_session_batch")
                if "KDP" in chosen_batch_src:
                    loaded_prompt_text = recent_kdp
                    source_name = st.session_state.get("kdp_batch_source_name", "KDP_Csomag")
                elif "Etsy" in chosen_batch_src:
                    loaded_prompt_text = recent_etsy
                    source_name = st.session_state.get("etsy_batch_source_name", "Etsy_Csomag")
                else:
                    loaded_prompt_text = recent_ideas
                    source_name = "30_Otlet_Gyujtemeny"
                st.success(f"⚡ Sikeresen betöltve: `{chosen_batch_src}`")
            else:
                st.info("ℹ️ Még nem generáltál csomagot ebben a munkamenetben. Válassz **Google Drive Mappából Tallózás** opciót vagy generálj csomagot az 1., 2. vagy 4. Munkaterületen!")

        elif "Google Drive" in prompt_source_mode:
            drive_folder_choice = st.selectbox(
                "Válassz Google Drive mappát:",
                [
                    "03_📚_AMAZON_KDP",
                    "04_🖼️_ETSY_DIGITAL",
                    "05_📖_GUMROAD_PLR",
                    "06_📌_MARKETING_ES_SEO"
                ],
                key="imagen_drive_folder_sel"
            )

            resolved_src_dir = resolve_drive_folder(drive_folder_choice)
            txt_files = []
            if os.path.exists(resolved_src_dir):
                txt_files = [f for f in os.listdir(resolved_src_dir) if f.lower().endswith(".txt")]

            if txt_files:
                selected_txt_file = st.selectbox(
                    f"Válassz .txt fájlt ({len(txt_files)} db található):",
                    options=txt_files,
                    key="imagen_selected_txt"
                )
                if selected_txt_file:
                    source_name = selected_txt_file.replace(".txt", "")
                    full_txt_path = os.path.join(resolved_src_dir, selected_txt_file)
                    try:
                        with open(full_txt_path, "r", encoding="utf-8") as tf:
                            loaded_prompt_text = tf.read()
                        st.success(f"📄 Betöltve: `{selected_txt_file}`")
                    except Exception as fe:
                        st.error(f"Hiba a fájl olvasásakor: {fe}")
            else:
                st.warning(f"⚠️ Ebben a mappában még nincs `.txt` promptfájl: `{resolved_src_dir}`")

        elif "feltöltése" in prompt_source_mode:
            uploaded_txt = st.file_uploader("Tölts fel egy .txt fájlt:", type=["txt"], key="imagen_file_uploader")
            if uploaded_txt:
                source_name = uploaded_txt.name.replace(".txt", "")
                loaded_prompt_text = uploaded_txt.getvalue().decode("utf-8", errors="ignore")
                st.success(f"📄 Feltöltve: `{uploaded_txt.name}`")

        else:
            loaded_prompt_text = st.text_area(
                "Illeszd be a promptokat (soronként vagy sorszámozva):",
                value="1. Noah's ark floating peacefully on calm waters with rainbow in the sky\n2. Moses with the burning bush on mount Sinai\n3. Daniel in the lions' den praying in peace",
                height=160,
                key="imagen_manual_prompts"
            )
            source_name = "Kézi_Promptok"

        # Parse prompts
        parsed_items = parse_prompts_from_text(loaded_prompt_text)
        st.markdown(f"**🎯 Felismert tételek száma:** `{len(parsed_items)} db`")

    with col_cfg:
        st.markdown("#### ⚙️ 2. Lépés: Google Drive Célmappa & Képbegyűjtés")

        target_dir_choice = st.selectbox(
            "📁 Google Drive Mappa (ahova a Gemini képeket mented):",
            [
                "03_📚_AMAZON_KDP",
                "04_🖼️_ETSY_DIGITAL",
                "05_📖_GUMROAD_PLR",
                "Egyedi mappa megadása..."
            ],
            key="imagen_target_dir_choice"
        )

        if "Egyedi mappa" in target_dir_choice:
            custom_out_dir = st.text_input("Egyedi mappa teljes elérési útja:", value=os.path.join(get_drive_root(), "03_📚_AMAZON_KDP"))
            resolved_out_dir = custom_out_dir
        elif "ETSY" in target_dir_choice:
            resolved_out_dir = resolve_drive_folder("etsy_wallart")
        elif "GUMROAD" in target_dir_choice:
            resolved_out_dir = resolve_drive_folder("gumroad")
        else:
            resolved_out_dir = resolve_drive_folder("kdp")

        st.caption(f"📁 Figyelt mappa: `{resolved_out_dir}`")

        col_b_act1, col_b_act2 = st.columns(2)
        with col_b_act1:
            btn_scan_bulk_imgs = st.button("🔄 Képek Begyűjtése a Mappából", key="btn_scan_bulk_imgs_act", use_container_width=True)
        with col_b_act2:
            if st.button("📁 Mappa Létrehozása", key="btn_mk_bulk_dir", use_container_width=True):
                os.makedirs(resolved_out_dir, exist_ok=True)
                st.toast("Mappa készen áll!", icon="📁")

    # Scanned items and PDF compiler
    if parsed_items:
        st.markdown("---")
        st.markdown("### 📋 3. Lépés: Gemini Promptok & Begyűjtött Képek Állapota")

        matched_bulk, unmatched_bulk = gather_images_from_directory(resolved_out_dir, expected_count=len(parsed_items))

        # Attach images to parsed items
        ready_cnt = 0
        for idx, item in enumerate(parsed_items):
            item_num = item["index"]
            if item_num in matched_bulk:
                item["filepath"] = matched_bulk[item_num]
                try:
                    item["pil_image"] = Image.open(matched_bulk[item_num])
                    with open(matched_bulk[item_num], "rb") as f_b_img:
                        item["image_bytes"] = f_b_img.read()
                except Exception:
                    pass
                ready_cnt += 1

        if ready_cnt == len(parsed_items):
            st.success(f"🎉 **Mind a {ready_cnt} / {len(parsed_items)} kép megtalálva a Drive mappában!**")
        else:
            st.info(f"ℹ️ **{ready_cnt} / {len(parsed_items)} kép elérhető a mappában.** Mentsd le a hiányzó képeket a Gemini-ből a fenti mappába `01.png`, `02.png` néven!")

        render_canva_image_gallery(
            records=parsed_items,
            prefix="bulk_ws6_gal",
            context_type="kdp"
        )

        st.markdown("---")
        st.markdown("### 🚀 4. Lépés: Nyomdakész PDF Összeállítása")
        col_cp1, col_cp2 = st.columns([1.2, 1])
        with col_cp1:
            if st.button(f"🚀 Teljes KDP Belső PDF Összeállítása ({1 + len(parsed_items)*3} oldal)", key="btn_compile_bulk_pdf", use_container_width=True):
                pdf_bulk_path = os.path.join(resolved_out_dir, f"KDP_Book_{sanitize_filename(source_name)}.pdf")
                ok_bp, bp_bytes, bp_msg = build_kdp_book_pdf(
                    title=source_name,
                    subtitle="Christian Digital Coloring Book",
                    pages_data=parsed_items,
                    output_path=pdf_bulk_path
                )
                if ok_bp and bp_bytes:
                    st.session_state["bulk_compiled_pdf_bytes"] = bp_bytes
                    st.session_state["bulk_compiled_pdf_fn"] = f"KDP_{sanitize_filename(source_name)}.pdf"
                    st.success(f"✨ **PDF Sikeresen Elkészült és Mentve a Drive-ra!**\n\n📁 `{pdf_bulk_path}`")
                else:
                    st.error(f"Hiba a PDF fordításakor: {bp_msg}")
        with col_cp2:
            if st.session_state.get("bulk_compiled_pdf_bytes"):
                st.download_button(
                    label="📥 KÉSZ PDF LETÖLTÉSE",
                    data=st.session_state["bulk_compiled_pdf_bytes"],
                    file_name=st.session_state.get("bulk_compiled_pdf_fn", "KDP_Book.pdf"),
                    mime="application/pdf",
                    key="btn_dl_bulk_pdf_done",
                    use_container_width=True
                )



# ==========================================================
# WORKSPACE: RENDSZERBEÁLLÍTÁSOK & KULCSOK
# ==========================================================

elif "Rendszerbeállítások" in menu_choice or "7." in menu_choice:
    st.markdown("<div class='path-badge'>⚙️ Rendszerközpont & AI Beállítások</div>", unsafe_allow_html=True)
    st.markdown("### ⚙️ Rendszerbeállítások, Szöveges AI Modellek & Google Drive Tárhely")
    st.caption("A rendszer a villámgyors Groq Cloud és OpenRouter ingyenes kvótáit használja elsődlegesen, és csak ezek kimerülése esetén vált át a Fizetős Gemini tartalék kulcsra.")

    st.markdown("---")

    col_k1, col_k2 = st.columns([1.15, 1], gap="large")

    with col_k1:
        st.markdown("#### 🔑 1. API Kulcsok & Intelligens Fallback Sorrend")
        st.caption("1️⃣ **Groq Cloud** (Elsődleges) ➔ 2️⃣ **OpenRouter** (Másodlagos) ➔ 3️⃣ **Fizetős Gemini** (Tartalék)")

        # 1. Groq API Key (Primary)
        curr_groq = km.groq_key or cfg.get("groq_api_key", "")
        groq_stat = "🟢" if curr_groq.strip() else "⚪"
        new_groq_key = st.text_input(
            f"{groq_stat} 🚀 1. Groq Cloud API Kulcs (Elsődleges · Llama 3.3 70B · 300 szó/mp · Ingyenes):",
            value=curr_groq,
            type="password",
            key="cfg_groq_key_input",
            help="Ingyenes és villámgyors (300 szó/mp). Regisztráció és ingyenes kulcs: console.groq.com/keys"
        )

        # 2. OpenRouter API Key (Secondary)
        curr_or = km.openrouter_key or cfg.get("openrouter_api_key", "")
        or_stat = "🟢" if curr_or.strip() else "⚪"
        new_openrouter_key = st.text_input(
            f"{or_stat} 🌐 2. OpenRouter API Kulcs (Másodlagos Tartalék · :free nyílt modellek · Ingyenes):",
            value=curr_or,
            type="password",
            key="cfg_openrouter_key_input",
            help="Ingyenes nyílt forráskódú modellek (:free hálózat). Regisztráció: openrouter.ai/keys"
        )

        st.markdown("---")

        # 3. Gemini Paid Key Fallback (Tertiary)
        curr_paid_key = km.paid_key or cfg.get("paid_key", "")
        gem_stat = "🟢" if curr_paid_key.strip() else "⚪"
        new_paid_key = st.text_input(
            f"{gem_stat} 💎 3. Google Gemini Fizetős API Kulcs (Paid Key · Csak kimerüléskor lép életbe):",
            value=curr_paid_key,
            type="password",
            key="cfg_paid_key_input",
            help="Ezt a kulcsot a rendszer CSAK AKKOR használja, ha a Groq és OpenRouter kvóták kimerültek. Így 0 Ft felesleges költség keletkezik normál működéskor."
        )
        st.caption("🛡️ *A fizetős Gemini kulcs 100%-os biztonsági háló: normál esetben az ingyenes Groq és OpenRouter futnak.*")

        col_t1, col_t2 = st.columns(2)
        with col_t1:
            if st.button("🔍 Szöveges AI Tesztelése", key="btn_test_text_providers_v2", use_container_width=True):
                with st.spinner("Szöveggenerálás tesztelése (Groq ➔ OpenRouter ➔ Paid Gemini)..."):
                    ok_txt, res_txt = km.generate_text_with_fallback("Írj 3 rövid bátorító mondatot!")
                    if ok_txt:
                        st.success(f"✅ Sikeres szöveggenerálás!\n\n*{res_txt[:160]}...*")
                    else:
                        st.error(f"Hiba: {res_txt}")
        with col_t2:
            if st.button("🔄 Rendszer Állapot Frissítése", key="btn_reset_keys_page_v2", use_container_width=True):
                km.reset_all_keys()
                st.success("🟢 AI szolgáltatók állapota frissítve!")
                st.rerun()

    with col_k2:
        st.markdown("#### 🤖 2. Alapértelmezett Szöveges Modell & Kreativitás")

        text_models_list = [
            "groq-llama-3.3-70b",
            "openrouter-free",
            "gemini-3.7-flash",
            "gemini-2.5-flash",
            "gemini-2.5-pro"
        ]

        sel_model = st.selectbox(
            "📝 Alapértelmezett Szöveges AI Modell:",
            options=text_models_list,
            index=text_models_list.index(text_model) if text_model in text_models_list else 0,
            key="cfg_model_sel_v2",
            help="A Groq Cloud Llama 3.3 70B az elsődleges, ingyenes és szupergyors modell. A Gemini modellek fizetős kulccsal érhetők el tartalékként."
        )

        sel_temp = st.slider(
            "🌡️ Hőmérséklet (Kreativitás):",
            min_value=0.0,
            max_value=1.0,
            value=temperature,
            step=0.05,
            key="cfg_temp_slider_v2",
            help="0.2 = Precíz | 0.7 = Kiegyensúlyozott | 1.0 = Nagyon kreatív"
        )

        st.markdown("---")
        st.markdown("#### 📁 3. Google Drive Fő Mappa Elérési Útja")
        drive_path_val = st.text_input(
            "Google Drive Fő Mappa Teljes Elérési Útja:",
            value=drive_root_input,
            key="cfg_drive_path_input_v2"
        )

        sa_info_fresh = get_service_account_info()
        if sa_info_fresh:
            email = sa_info_fresh.get("client_email", "Aktív")
            st.success(f"🟢 Google Drive Cloud API csatlakoztatva és aktív! (Robot fiók: `{email}`)")
        elif os.path.exists(drive_path_val):
            st.success("🟢 Google Drive helyi mappa elérhető és csatlakoztatva!")
        else:
            st.info("💡 Helyi módban a laptop meghajtóját használja. Streamlit Cloudhoz add meg a Service Account JSON-t a Secrets-ben vagy illeszd be ide alább!")

        sa_json_val = st.text_area(
            "☁️ Google Service Account JSON Beillesztése (ha nem Secrets-ben adtad meg):",
            value="",
            height=70,
            placeholder='{"type": "service_account", "project_id": "family-cashflow-4de1e", ...}',
            key="cfg_sa_json_direct_input",
            help="Ide is beillesztheted közvetlenül a letöltött Google Cloud JSON kulcsodat, majd kattints a Mentés gombra."
        )

        with st.expander("🔍 Secrets & Google Drive Diagnosztika"):
            secrets_keys = []
            try:
                secrets_keys = list(st.secrets.keys())
            except Exception:
                pass
            st.write("📋 Streamlit Secrets-ben látott kulcsok:", secrets_keys)
            if sa_info_fresh:
                st.success(f"✅ Aktív robot e-mail: `{sa_info_fresh.get('client_email')}`")
                st.info(f"📂 Projekt ID: `{sa_info_fresh.get('project_id')}`")
            else:
                st.warning("⚠️ Még nem sikerült felismerni a Service Account kulcsot a Secrets-ből.")

        if st.button("📁 Mappastruktúra Ellenőrzése / Létrehozása", key="btn_init_drive_dirs_v2", use_container_width=True):
            try:
                os.makedirs(drive_path_val, exist_ok=True)
                for f_key in ["kdp_interiors", "kdp_covers", "etsy_wallart", "etsy_clipart", "gumroad", "marketing"]:
                    f_path = resolve_drive_folder(f_key, custom_root=drive_path_val)
                    os.makedirs(f_path, exist_ok=True)
                st.success("✅ Minden almappa létrejött a meghajtón (03_📚_AMAZON_KDP, 04_🖼️_ETSY, 05_📖_GUMROAD, 06_📌_MARKETING_ES_SEO)!")
            except Exception as de:
                st.error(f"Hiba a mappák létrehozásakor: {de}")

        st.markdown("---")
        st.markdown("#### 🛍️ 4. Gumroad Integráció")
        gumroad_token_input = st.text_input(
            "🔑 Gumroad API Access Token:",
            value=cfg.get("gumroad_access_token", ""),
            type="password",
            key="cfg_gumroad_token_v2",
            help="Gumroad fiókod személyes Access Tokenje."
        )

        st.markdown("---")
        if st.button("💾 ÖSSZES BEÁLLÍTÁS MENTÉSE (config.json)", key="btn_save_all_config_v2", use_container_width=True):
            clean_paid_key = new_paid_key.strip()
            chosen_text = st.session_state.get("cfg_model_sel_v2", sel_model)

            cfg_save_dict = {
                "groq_api_key": new_groq_key.strip(),
                "openrouter_api_key": new_openrouter_key.strip(),
                "paid_key": clean_paid_key,
                "api_key": "",
                "free_keys": [],
                "api_keys": [],
                "selected_model": chosen_text,
                "temperature": sel_temp,
                "drive_root_path": drive_path_val,
                "gumroad_access_token": gumroad_token_input.strip()
            }
            if sa_json_val.strip():
                cfg_save_dict["google_service_account_json"] = sa_json_val.strip()

            save_config(cfg_save_dict)
            km.update_keys(
                paid_key=clean_paid_key,
                groq_key=new_groq_key.strip(),
                openrouter_key=new_openrouter_key.strip()
            )

            st.success("🎉 **Beállítások sikeresen elmentve!**")
            time.sleep(0.6)
            st.rerun()


            st.success("🎉 **Beállítások sikeresen elmentve!**")
            time.sleep(0.6)
            st.rerun()

    st.markdown("---")
    st.markdown("#### 📊 Rendszer Kvóták & AI Fallback Hierarchia")
    col_q1, col_q2, col_q3 = st.columns(3)
    with col_q1:
        st.markdown("""
        **🚀 1. Groq Cloud (Elsődleges)**
        - **Modell:** Llama 3.3 70B (300 szó/mp)
        - **Költség:** 100% Ingyenes & Villámgyors
        - **Használat:** Minden prompt, vázlat és marketing szöveg azonnali generálása
        """)
    with col_q2:
        st.markdown("""
        **🌐 2. OpenRouter (Másodlagos)**
        - **Modellek:** `:free` nyílt forráskódú modellek
        - **Költség:** 100% Ingyenes biztonsági háló
        - **Használat:** Automatikusan bekapcsol, ha a Groq kvóta kimerül
        """)
    with col_q3:
        st.markdown("""
        **💎 3. Fizetős Gemini (Tartalék)**
        - **Modellek:** Gemini 2.5 Flash / 3.7 Flash / Pro
        - **Költség:** Csak kimerüléskor fogyaszt
        - **Használat:** 0 Ft felesleges kiadás, folyamatos munka
        """)



# ==========================================================
# WORKSPACE: FFC MARKETING, COPYWRITING & GOOGLE SITES STÚDIÓ
# ==========================================================

elif "FFC Marketing" in menu_choice or "Google Sites" in menu_choice or "8." in menu_choice:
    st.markdown("<div class='path-badge'>🚀 FFC Launchpad, Reels Gyár & Google Ökoszisztéma</div>", unsafe_allow_html=True)
    st.markdown("### 🚀 FFC Faceless Funnel, Reels Tartalomgyár & Google Automatizáció")
    st.caption("Teljes Keresztény Értékesítési Ökoszisztéma: Russell Brunson 10-részes Sales Pack, 10 db Arc Nélküli Reels + FLUX.1 B-roll képek, 0 Ft-os Stripe ➔ Sheets ➔ Gmail webhook, Google Sites landing page és 30 napos e-mail tölcsér.")

    active_art_style_prompt = render_style_selector("ws_ffc")
    render_niche_status_bar("ffc")

    tab_ffc1, tab_ffc2, tab_ffc3, tab_ffc4, tab_ffc5, tab_ffc6 = st.tabs([
        "🎯 1. FFC 10-Részes Sales Pack & 3-Tagú Bulletek",
        "🎬 2. Arc Nélküli Reels & ManyChat Tartalomgyár (10 db)",
        "⚡ 3. Google Apps Script & Stripe 0 Ft-os Automatizáció",
        "🌐 4. Google Sites 0 Ft-os Landing Page",
        "📧 5. Automata E-mail Tölcsér (3 & 30 Napos)",
        "📅 6. 30 Napos Social SEO Naptár"
    ])

    # Get active niche info
    curr_niche_key = st.session_state.get("active_niche_choice", "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)")
    curr_niche_data = get_niche_prompt_context(curr_niche_key)
    niche_default_aud = curr_niche_data.get("default_audience", "Keresztény édesanyák, alkotók és hívők")

    # ─────────────────────────────────────────────────────
    # AL-FÜL 1: FFC 10-RÉSZES SALES PACK & 3-TAGÚ BULLETEK
    # ─────────────────────────────────────────────────────
    with tab_ffc1:
        st.markdown("#### 🎯 FFC 10-Részes Értékesítési Csomag & 3-Tagú Bulletek")
        st.caption("Russell Brunson és Alex Hormozi stílusú, mély lélektani konverziós gépezet: Big Domino, 5 Brunson Főcím, 3-Perces VSL, 3-Tagú Bulletek, Kifogáskezelés és Value Stack.")

        col_f1_in, col_f1_out = st.columns([1, 1.15], gap="large")

        with col_f1_in:
            n_slug_ffc = get_niche_slug(curr_niche_key)
            
            st.markdown("<div class='step-label'>1. Lépés: Nyelv & Terméktípus Kiválasztása</div>", unsafe_allow_html=True)
            col_sub1, col_sub2 = st.columns(2)
            with col_sub1:
                ffc_lang = st.selectbox("🌐 Nyelv (Language):", ["Magyar", "Angol (English)"], index=0, key="ffc_lang_v2")
            with col_sub2:
                ffc_prod_type = st.selectbox(
                    "📦 Terméktípus:",
                    ["30 Napos Áhítat & Napló", "Keresztény Színezőkönyv (KDP)", "Igés Falikép / Clipart Csomag", "Keresztény Tanfolyam / Masterclass"],
                    index=0,
                    key="ffc_prod_type_v2"
                )

            # Get tailored preset data for this exact Niche + Product Type + Language
            ffc_preset = get_ffc_preset(curr_niche_key, ffc_prod_type, ffc_lang)
            pt_slug = f"{sanitize_filename(ffc_prod_type)[:10]}_{'en' if 'Angol' in ffc_lang else 'hu'}"

            st.markdown("<div class='step-label'>2. Lépés: Termék & Transzformáció Paraméterek</div>", unsafe_allow_html=True)
            ffc_prod_name = st.text_input(
                "Termék / Ajánlat Neve:",
                value=st.session_state.get(f"ffc_prod_{n_slug_ffc}_{pt_slug}", ffc_preset["prod_name"]),
                key=f"ffc_prod_{n_slug_ffc}_{pt_slug}"
            )
            ffc_target_aud = st.text_input(
                "Célközönség / Avatár:",
                value=st.session_state.get(f"ffc_aud_{n_slug_ffc}_{pt_slug}", ffc_preset["target_aud"]),
                key=f"ffc_aud_{n_slug_ffc}_{pt_slug}"
            )
            ffc_main_trans = st.text_area(
                "Fő Transzformáció (Végső Érzelmi/Szellemi Ígéret):",
                value=st.session_state.get(f"ffc_trans_{n_slug_ffc}_{pt_slug}", ffc_preset["main_trans"]),
                height=75,
                key=f"ffc_trans_{n_slug_ffc}_{pt_slug}"
            )
            ffc_vehicle = st.text_input(
                "Az Új Módszer / Kulcs (The Vehicle):",
                value=st.session_state.get(f"ffc_veh_{n_slug_ffc}_{pt_slug}", ffc_preset["vehicle"]),
                key=f"ffc_veh_{n_slug_ffc}_{pt_slug}"
            )
            ffc_extra_notes = st.text_area(
                "➕ Extra preferenciák / Különleges részletek (opcionális):",
                value=st.session_state.get(f"ffc_notes_{n_slug_ffc}_{pt_slug}", ffc_preset["extra_notes"]),
                height=60,
                key=f"ffc_notes_{n_slug_ffc}_{pt_slug}"
            )

            btn_gen_sales_pack = st.button("🚀 Teljes FFC Értékesítési Csomag Generálása (AI)", key="btn_gen_sales_pack", use_container_width=True)

        with col_f1_out:
            st.markdown("<div class='step-label'>Generált 10-Részes FFC Értékesítési Csomag</div>", unsafe_allow_html=True)

            if btn_gen_sales_pack:
                with st.spinner("AI írja a teljes 10-részes FFC értékesítési csomagot (Big Domino, VSL, 3-tagú bulletek, Value Stack)..."):
                    if FFC_MODULES_AVAILABLE:
                        res_pack = generate_ffc_sales_pack(
                            topic=ffc_prod_name,
                            target_audience=ffc_target_aud,
                            core_transformation=ffc_main_trans,
                            language=ffc_lang,
                            vehicle=ffc_vehicle,
                            extra_notes=ffc_extra_notes,
                            product_type=ffc_prod_type
                        )
                    else:
                        p_av = build_ffc_avatar_research_prompt(ffc_prod_name, ffc_target_aud, ffc_main_trans, ffc_extra_notes, niche_name=curr_niche_key)
                        p_hk = build_ffc_big_domino_hooks_prompt(ffc_prod_name, ffc_target_aud, ffc_main_trans, ffc_vehicle, language=ffc_lang, niche_name=curr_niche_key)
                        p_sl = build_ffc_sales_letter_prompt(ffc_prod_name, ffc_target_aud, ffc_main_trans, ffc_extra_notes, ffc_vehicle, "", "", ffc_lang, curr_niche_key)
                        ok_call, res_text = km.generate_text_with_fallback(f"{p_av}\n\n{p_hk}\n\n{p_sl}")
                        res_pack = {"sales_letter_full": res_text, "big_domino": "Big Domino Generálva", "headlines": [], "three_part_bullets": []}
                    
                    st.session_state["active_ffc_sales_pack"] = res_pack
                    st.session_state["active_ffc_prod_name"] = ffc_prod_name

            if st.session_state.get("active_ffc_sales_pack"):
                pack = st.session_state["active_ffc_sales_pack"]
                p_name = st.session_state.get("active_ffc_prod_name", ffc_prod_name)

                # 1. Big Domino Card
                if "big_domino" in pack and pack["big_domino"]:
                    st.markdown(f"""
                    <div style='background: linear-gradient(135deg, rgba(16, 185, 129, 0.15), rgba(5, 150, 105, 0.08)); border-left: 4px solid #10b981; border-radius: 10px; padding: 14px 18px; margin-bottom: 15px;'>
                        <div style='font-size: 0.85rem; font-weight: 700; color: #34d399; text-transform: uppercase;'>🎯 A Big Domino Állítás:</div>
                        <div style='font-size: 1.02rem; color: #f8fafc; margin-top: 5px; font-weight: 600;'>{pack["big_domino"]}</div>
                    </div>
                    """, unsafe_allow_html=True)

                # 2. Headlines
                if "headlines" in pack and isinstance(pack["headlines"], list) and len(pack["headlines"]) > 0:
                    with st.expander("📢 5 db Russell Brunson Főcím (Headlines)", expanded=True):
                        for hl in pack["headlines"]:
                            st.markdown(f"- **{hl}**")

                # 3. 3-Minute VSL Script
                if "vsl_script" in pack and pack["vsl_script"]:
                    with st.expander("🎬 3-Perces VSL (Video Sales Letter) Forgatókönyv", expanded=False):
                        st.markdown(pack["vsl_script"])

                # 4. 3-Part Bullets
                if "three_part_bullets" in pack and isinstance(pack["three_part_bullets"], list) and len(pack["three_part_bullets"]) > 0:
                    with st.expander("💎 5 db 3-Tagú Termék Bullet Pont ('Mit kap' + 'Még akkor is ha' + 'Ami azt jelenti')", expanded=True):
                        for idx, b in enumerate(pack["three_part_bullets"], 1):
                            if isinstance(b, dict):
                                f_val = b.get('mit_kap') or b.get('feature') or b.get('what_you_get') or b.get('component') or b.get('title') or ''
                                e_val = b.get('meg_akkor_is_ha') or b.get('even_if') or b.get('objection') or ''
                                m_val = b.get('ami_azt_jelenti') or b.get('which_means') or b.get('meaning') or b.get('outcome') or ''
                                st.markdown(f"""
                                <div style='background: #1e293b; border-radius: 8px; padding: 10px 14px; margin-bottom: 8px; border: 1px solid #334155;'>
                                    <div style='color: #60a5fa; font-weight: 700;'>📦 #{idx}: {f_val}</div>
                                    <div style='color: #f87171; font-size: 0.92rem; margin-top: 3px;'>🛡️ <em>{e_val}</em></div>
                                    <div style='color: #34d399; font-size: 0.95rem; font-weight: 600; margin-top: 3px;'>✨ <strong>{m_val}</strong></div>
                                </div>
                                """, unsafe_allow_html=True)
                            else:
                                st.markdown(f"- {b}")

                # 5. Objection Handling & Value Stack
                if "objection_handling" in pack and isinstance(pack["objection_handling"], dict):
                    with st.expander("🛡️ 3 Fő Kifogáskezelés (Vehicle, Belső Önbizalom, Külső Időhiány)", expanded=False):
                        for k, v in pack["objection_handling"].items():
                            st.markdown(f"- **{k.replace('_', ' ').title()}:** {v}")

                if "value_stack" in pack and isinstance(pack["value_stack"], dict):
                    vs = pack["value_stack"]
                    with st.expander("💰 Value Stack Értékkosár & Garancia", expanded=True):
                        st.markdown(f"""
                        - 🎁 **Fő Termék:** {vs.get('main_product_value', '')}
                        - 🎁 **Bónusz #1:** {vs.get('bonus_1', '')}
                        - 🎁 **Bónusz #2:** {vs.get('bonus_2', '')}
                        - 🎁 **Bónusz #3:** {vs.get('bonus_3', '')}
                        - 📊 **Teljes Valós Érték:** `{vs.get('total_value', '')}` ➔ **Ajánlati Ár:** <span style='color:#34d399; font-weight:bold; font-size:1.1rem;'>{vs.get('offer_price', '')}</span>
                        - 🛡️ **Garancia:** {vs.get('guarantee_text', '')}
                        """, unsafe_allow_html=True)

                # 6. Full Sales Letter Text
                full_sl = pack.get("sales_letter_full", "")
                if not full_sl:
                    full_sl = f"# {p_name}\n\n## Big Domino:\n{pack.get('big_domino', '')}\n\n## VSL:\n{pack.get('vsl_script', '')}"

                with st.expander("📜 Teljes Formázott Értékesítési Levél (Sales Letter Másolása)", expanded=False):
                    st.markdown(full_sl)

                st.text_area(
                    "📋 Nyers Sales Pack Szöveg Másolása (Ctrl+A → Ctrl+C):",
                    value=json.dumps(pack, ensure_ascii=False, indent=2) if isinstance(pack, dict) else str(pack),
                    height=180,
                    key="ffc_pack_copy_area"
                )

                col_dl_p1, col_dl_p2 = st.columns(2)
                with col_dl_p1:
                    st.download_button(
                        label="⬇️ Letöltés (.txt)",
                        data=full_sl.encode("utf-8"),
                        file_name=f"Marketing_{sanitize_filename(p_name)}_Sales_Pack.txt",
                        mime="text/plain",
                        key="dl_ffc_pack_txt",
                        use_container_width=True
                    )
                with col_dl_p2:
                    if DOCX_AVAILABLE:
                        docx_bio = create_marketing_docx(p_name, full_sl, header_info=f"FFC 10-Részes Sales Pack - {p_name}")
                        st.download_button(
                            label="⬇️ Letöltés (.docx)",
                            data=docx_bio.getvalue(),
                            file_name=f"Marketing_{sanitize_filename(p_name)}_Sales_Pack.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_ffc_pack_docx",
                            use_container_width=True
                        )

                st.markdown("---")
                if st.button("💾 Mentés a 06_📌_MARKETING_ES_SEO Mappába (és Drive-ra)", key="btn_save_ffc_pack_drive", use_container_width=True):
                    ok_s, details, _ = save_marketing_file_to_drive(
                        product_name=p_name,
                        text_content=full_sl,
                        header_info=f"FFC 10-Részes Sales Pack - {p_name}",
                        content_type_tag="Sales_Pack"
                    )
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés!**\n\n`{details}`")
                    else:
                        st.error(f"Hiba a mentéskor: {details}")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🎯</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg a termék adatait a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">🚀 Teljes FFC Értékesítési Csomag Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────
    # AL-FÜL 2: ARC NÉLKÜLI REELS & MANYCHAT TARTALOMGYÁR
    # ─────────────────────────────────────────────────────
    with tab_ffc2:
        st.markdown("#### 🎬 Arc Nélküli (Faceless) Reels & ManyChat Tartalomgyár (10 db Batch)")
        st.caption("1 kattintással 10 db 5-7 másodperces virális keresztény Reels forgatókönyv, ManyChat CTA kulcsszó, posztszöveg és 9:16 vertikális FLUX.1 háttérkép prompt CapCut-hoz.")

        col_r_in, col_r_out = st.columns([1, 1.2], gap="large")

        with col_r_in:
            reels_lang = st.selectbox("🌐 Nyelv (Language):", ["Magyar", "Angol (English)"], index=0, key="reels_lang_select_v2")
            is_r_en = "Angol" in reels_lang or "English" in reels_lang
            r_lang_slug = "en" if is_r_en else "hu"

            default_r_prod = f"30-Day {curr_niche_data.get('name_en', 'Christian')} Guided Journal" if is_r_en else get_niche_field("reels_prod", curr_niche_key)
            default_r_aud = "Parents, creators, and believers seeking peace and focus" if is_r_en else get_niche_field("ffc_aud", curr_niche_key)

            st.markdown("<div class='step-label'>Reels & ManyChat Paraméterek</div>", unsafe_allow_html=True)
            reels_prod = st.text_input(
                "Termék / Téma:",
                value=st.session_state.get(f"reels_prod_{n_slug_ffc}_{r_lang_slug}", default_r_prod),
                key=f"reels_prod_{n_slug_ffc}_{r_lang_slug}"
            )
            reels_target = st.text_input(
                "Célközönség:",
                value=st.session_state.get(f"reels_aud_{n_slug_ffc}_{r_lang_slug}", default_r_aud),
                key=f"reels_aud_{n_slug_ffc}_{r_lang_slug}"
            )
            
            if is_r_en:
                reels_cta_kw_opts = ["PEACE", "FAITH", "MIRACLE", "GRACE", "FOCUS", "FREEDOM", "BLESSING", "SUCCESS", "GROWTH"]
            else:
                reels_cta_kw_opts = ["BÉKESSÉG", "CSODA", "ÁLDÁS", "HIT", "SIKER", "REMÉNY", "SZABADSÁG", "FÓKUSZ", "IRÁNYTŰ", "VAGYON", "EGÉSZSÉG"]
            
            reels_cta_kw = st.selectbox(
                "ManyChat CTA Kulcsszó (Ezt kell kommentelniük):",
                reels_cta_kw_opts,
                index=0,
                key=f"reels_cta_kw_{n_slug_ffc}_{r_lang_slug}"
            )

            btn_gen_reels = st.button("🎬 10 db Virális Reels & B-roll Prompt Generálása (AI)", key="btn_gen_reels_batch", use_container_width=True)

        with col_r_out:
            st.markdown("<div class='step-label'>10 db Kész Reels Forgatókönyv & Képgenerátor</div>", unsafe_allow_html=True)

            if btn_gen_reels:
                with st.spinner("AI generálja a 10 db virális Reels forgatókönyvet, ManyChat CTA-kat és 9:16 FLUX.1 képpromptokat..."):
                    if FFC_MODULES_AVAILABLE:
                        reels_data = generate_faceless_reels_batch(
                            topic=reels_prod,
                            cta_keyword=reels_cta_kw,
                            target_audience=reels_target,
                            count=10,
                            language=reels_lang
                        )
                    else:
                        reels_data = []
                    st.session_state["active_reels_batch"] = reels_data
                    st.session_state["active_reels_prod"] = reels_prod

            if st.session_state.get("active_reels_batch"):
                r_batch = st.session_state["active_reels_batch"]
                p_r_name = st.session_state.get("active_reels_prod", reels_prod)

                st.success(f"✅ **10 db Reels Sikeresen Legenerálva!** (ManyChat Kulcsszó: `{reels_cta_kw}`)")

                for item in r_batch:
                    r_id = item.get("id", 1)
                    r_title = item.get("title", f"Reels #{r_id}")
                    r_hook = item.get("hook_text", "")
                    r_body = item.get("body_text", "")
                    r_cta = item.get("cta_text", "")
                    r_caption = item.get("caption_text", "")
                    r_img_prompt = item.get("image_prompt", "")

                    with st.expander(f"🎬 #{r_id} · {r_hook[:45]}...", expanded=(r_id == 1)):
                        st.markdown(f"""
                        <div style='background: #1e293b; border-radius: 8px; padding: 12px; border: 1px solid #334155; margin-bottom: 8px;'>
                            <div style='color: #f87171; font-weight: 700;'>🎣 0:00-0:03 Horog (Hook): <span style='color:#f8fafc;'>"{r_hook}"</span></div>
                            <div style='color: #60a5fa; font-weight: 600; margin-top: 6px;'>💡 0:03-0:07 Megoldás: <span style='color:#cbd5e1;'>{r_body}</span></div>
                            <div style='color: #34d399; font-weight: 700; margin-top: 6px;'>📲 ManyChat CTA: <span style='color:#a7f3d0;'>"{r_cta}"</span></div>
                        </div>
                        """, unsafe_allow_html=True)

                        st.markdown("**📝 Instagram Posztszöveg & Hashtagek:**")
                        st.code(r_caption, language="text")

                        st.markdown("**🎨 9:16 FLUX.1 Képprompt (B-roll Háttér):**")
                        st.code(r_img_prompt, language="text")

                        col_gen_img1, col_gen_img2 = st.columns([1.2, 1])
                        with col_gen_img1:
                            btn_img_single = st.button(f"🎨 B-roll Kép Generálása (#{r_id})", key=f"btn_gen_img_r_{r_id}", use_container_width=True)
                        with col_gen_img2:
                            pass

                        if btn_img_single:
                            with st.spinner(f"FLUX.1 9:16 vertikális kép generálása (#{r_id})..."):
                                if FFC_MODULES_AVAILABLE:
                                    ok_img, img_bytes, img_msg = generate_reels_broll_image(r_img_prompt)
                                    if ok_img and img_bytes:
                                        st.session_state[f"reels_img_{r_id}"] = img_bytes
                                        st.success("✅ Kép sikeresen elkészült!")
                                    else:
                                        st.error(f"Hiba a képgeneráláskor: {img_msg}")

                        if st.session_state.get(f"reels_img_{r_id}"):
                            st.image(st.session_state[f"reels_img_{r_id}"], caption=f"Reels #{r_id} 9:16 FLUX.1 B-roll", use_container_width=True)
                            st.download_button(
                                label=f"📥 Kép Letöltése (Reels_{r_id}.png)",
                                data=st.session_state[f"reels_img_{r_id}"],
                                file_name=f"Reels_{sanitize_filename(p_r_name)}_{r_id}.png",
                                mime="image/png",
                                key=f"dl_r_img_{r_id}",
                                use_container_width=True
                            )

                st.markdown("---")
                all_reels_text = "\n\n" + "="*50 + "\n\n".join([
                    f"=== REELS #{it.get('id', 1)} ===\nHOOK: {it.get('hook_text','')}\nBODY: {it.get('body_text','')}\nCTA: {it.get('cta_text','')}\n\nCAPTION:\n{it.get('caption_text','')}\n\nB-ROLL PROMPT (FLUX.1 9:16):\n{it.get('image_prompt','')}"
                    for it in r_batch
                ])

                col_dl_r1, col_dl_r2 = st.columns(2)
                with col_dl_r1:
                    st.download_button(
                        label="⬇️ Mind a 10 db Reels Szöveg Letöltése (.txt)",
                        data=all_reels_text.encode("utf-8"),
                        file_name=f"Marketing_{sanitize_filename(p_r_name)}_10_Reels_Batch.txt",
                        mime="text/plain",
                        key="dl_reels_batch_txt",
                        use_container_width=True
                    )
                with col_dl_r2:
                    if st.button("💾 Összes Reels Mentése Drive-ra", key="btn_save_all_reels_drive", use_container_width=True):
                        ok_s, details, _ = save_marketing_file_to_drive(
                            product_name=p_r_name,
                            text_content=all_reels_text,
                            header_info=f"10 db Arc Nélküli Reels Forgatókönyv - {p_r_name}",
                            content_type_tag="10_Reels_Batch"
                        )
                        if ok_s:
                            st.success(f"💾 **Sikeres mentés!**\n\n`{details}`")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🎬</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg a témát és a ManyChat kulcsszót a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">🎬 10 db Virális Reels Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────
    # AL-FÜL 3: GOOGLE APPS SCRIPT & STRIPE 0 FT-OS AUTOMATIZÁCIÓ
    # ─────────────────────────────────────────────────────
    with tab_ffc3:
        st.markdown("#### ⚡ Google Apps Script & Stripe 0 Ft-os Értékesítési Automatizáció")
        st.caption("0 Ft-os, szervermentes digitális értékesítési tölcsér: Stripe fizetés ➔ Google Sheets rendelésnaplózás ➔ Azonnali digitális kézbesítő levél a Gmail fiókodból a privát Google Drive linkkel.")

        col_g_in, col_g_out = st.columns([1, 1.2], gap="large")

        with col_g_in:
            st.markdown("<div class='step-label'>Automatizáció Paraméterek</div>", unsafe_allow_html=True)
            wh_prod_name = st.text_input(
                "Értékesített Termék Neve (ami az e-mailben megjelenik):",
                value="30 Napos Békesség & Fókusz Vezetett Keresztény Csomag",
                key="wh_prod_name_input"
            )
            wh_drive_link = st.text_input(
                "Privát Google Drive Mappa Linkje (ezt kapja meg a vásárló a Gmailből):",
                value=st.session_state.get("last_uploaded_drive_link", "https://drive.google.com/drive/folders/00_VALLALKOZAS_AUDHD_DIGITALIS_BIRODALOM"),
                key="wh_drive_link_input",
                help="A vásárló ezt a titkos linket kapja meg közvetlenül a sikeres Stripe fizetés után a saját Gmail fiókodból."
            )
            wh_sheet_name = st.text_input("Google Sheets Munkalap Neve:", value="Rendelések", key="wh_sheet_name_input")
            wh_stripe_url = st.text_input("Stripe Fizetési Link (Buy Button linkje):", value="https://buy.stripe.com/pelda_fizetes", key="wh_stripe_url_input")

        with col_g_out:
            st.markdown("<div class='step-label'>Kész Google Apps Script Kód & Beállítás</div>", unsafe_allow_html=True)

            if FFC_MODULES_AVAILABLE:
                apps_script_code = get_apps_script_webhook_template(
                    drive_folder_url=wh_drive_link,
                    sheet_name=wh_sheet_name,
                    product_name=wh_prod_name
                )
                setup_guide = get_stripe_setup_guide()
                embed_btn_html = get_google_sites_embed_button(checkout_url=wh_stripe_url)
            else:
                apps_script_code = "// Google Hub modul nem érhető el"
                setup_guide = ""
                embed_btn_html = ""

            st.markdown("##### 📋 1. Másold be ezt a kódot a Google Sheets Apps Script felületére:")
            st.code(apps_script_code, language="javascript")

            with st.expander("🛠️ 2. Lépésről-lépésre Beállítási Útmutató (Google Sheets & Stripe)", expanded=True):
                st.markdown(setup_guide)

            with st.expander("🌐 3. Google Sites Beágyazható Vásárlás Gomb (HTML Kód)", expanded=False):
                st.caption("Másold ki ezt a HTML kódot, és a Google Sites felületén válaszd a **Beágyazás (Embed) ➔ Kód beágyazása** lehetőséget!")
                st.code(embed_btn_html, language="html")

    # ─────────────────────────────────────────────────────
    # AL-FÜL 4: GOOGLE SITES 0 FT-OS LANDING PAGE GENERÁLÓ
    # ─────────────────────────────────────────────────────
    with tab_ffc4:
        st.markdown("#### 🌐 Google Sites 0 Ft-os Landing Page & Értékesítési Szöveg Generáló")
        st.caption("Előállítja a 100%-ban ingyenes Google Sites (sites.google.com) keretrendszerhez szükséges teljes, blokkokra bontott struktúrát és szövegeket (Hero, Ingyenes Csalitermék, Kiemelt Termékkártyák, CTA gombok, és Színtéma Útmutató).")

        col_f3_in, col_f3_out = st.columns([1, 1.15], gap="large")

        with col_f3_in:
            st.markdown("<div class='step-label'>1. Lépés: Nyelv & Stílus Kiválasztása</div>", unsafe_allow_html=True)
            col_sub_p1, col_sub_p2 = st.columns(2)
            with col_sub_p1:
                gs_lang = st.selectbox("🌐 Nyelv (Language):", ["Magyar", "English"], index=0, key="gs_lang_select_v2")
            with col_sub_p2:
                gs_price = st.text_input("Ajánlati Ár (Gumroad):", value="$19 (Value: $67)", key="gs_price_v2")

            is_gs_en = "English" in gs_lang
            gs_lang_slug = "en" if is_gs_en else "hu"

            if is_gs_en:
                def_gs_prod = f"30-Day {curr_niche_data.get('name_en', 'Christian')} Guided Journal"
                def_gs_aud = "Parents, creators, and believers seeking daily quiet time and focus"
                def_gs_head = "Find Daily Peace, Spiritual Clarity, and Joy in Your Everyday Life"
                def_gs_tag = "A beautiful 30-day printable guided workbook, scripture collection, and mindfulness companion for calm, balanced days."
                def_gs_lm = "3-page free printable coloring sheet and mini workbook sample pack with instant download."
                def_gs_ft = "30-day guided workbook, 30 high-resolution printable templates, 5 color palette guides, and instant digital PDF download."
            else:
                def_gs_prod = get_niche_field("gs_prod", curr_niche_key)
                def_gs_aud = get_niche_field("ffc_aud", curr_niche_key)
                def_gs_head = get_niche_field("gs_headline", curr_niche_key)
                def_gs_tag = get_niche_field("gs_tagline", curr_niche_key)
                def_gs_lm = get_niche_field("gs_lead_magnet", curr_niche_key)
                def_gs_ft = get_niche_field("gs_features", curr_niche_key)

            st.markdown("<div class='step-label'>2. Lépés: Google Sites Landing Page Paraméterek</div>", unsafe_allow_html=True)
            gs_prod_name = st.text_input(
                "Termék Neve:",
                value=st.session_state.get(f"gs_prod_{n_slug_ffc}_{gs_lang_slug}", def_gs_prod),
                key=f"gs_prod_{n_slug_ffc}_{gs_lang_slug}"
            )
            gs_target_aud = st.text_input(
                "Célközönség:",
                value=st.session_state.get(f"gs_aud_{n_slug_ffc}_{gs_lang_slug}", def_gs_aud),
                key=f"gs_aud_{n_slug_ffc}_{gs_lang_slug}"
            )
            gs_headline = st.text_input(
                "Hero Főcímsor (Main Headline):",
                value=st.session_state.get(f"gs_head_{n_slug_ffc}_{gs_lang_slug}", def_gs_head),
                key=f"gs_head_{n_slug_ffc}_{gs_lang_slug}"
            )
            gs_tagline = st.text_area(
                "Alcím / Életérzés (Tagline):",
                value=st.session_state.get(f"gs_tag_{n_slug_ffc}_{gs_lang_slug}", def_gs_tag),
                height=70,
                key=f"gs_tag_{n_slug_ffc}_{gs_lang_slug}"
            )
            gs_lead_magnet = st.text_area(
                "🎁 Ingyenes Csalitermék (Lead Magnet) Leírása:",
                value=st.session_state.get(f"gs_lm_{n_slug_ffc}_{gs_lang_slug}", def_gs_lm),
                height=70,
                key=f"gs_lm_{n_slug_ffc}_{gs_lang_slug}"
            )
            gs_features = st.text_area(
                "Csomag Tartalma & Főbb Előnyök:",
                value=st.session_state.get(f"gs_ft_{n_slug_ffc}_{gs_lang_slug}", def_gs_ft),
                height=75,
                key=f"gs_ft_{n_slug_ffc}_{gs_lang_slug}"
            )

            gs_style = st.selectbox(
                "🎨 Vizuális Design & Színvilág (Theme):",
                [
                    "🌿 Keresztény Pasztell Minimalista (Zsályazöld #8A9A86, Bézs #F9F6F0 & Arany #D4AF37)",
                    "🌸 Pasztell Akvarell & Virágos Elegancia (Soft Rose, Lilac & Cream)",
                    "💼 Modern Sötét Palatábla & Smaragd (Deep Slate #1E293B & Emerald #34D399)",
                    "🌾 Földszínek & Boho Minimalista (Boho Earth Tones, Terracotta & Linen)",
                    "🏛️ Klasszikus Királykék & Arany (Royal Navy #1E3A8A & Gold #D4AF37)"
                ],
                index=0,
                key="gs_style"
            )

            st.markdown("##### 🔗 Közvetlen Értékesítési Gombok Linkjei (CTA Links):")
            col_l1, col_l2, col_l3 = st.columns(3)
            with col_l1:
                gs_amazon = st.text_input("Amazon KDP:", value="https://amazon.com/dp/B0EXAMPLE", key="gs_amazon")
            with col_l2:
                gs_etsy = st.text_input("Etsy Shop:", value="https://etsy.com/listing/EXAMPLE", key="gs_etsy")
            with col_l3:
                gs_gumroad = st.text_input("Gumroad / Stripe:", value="https://gumroad.com/l/EXAMPLE", key="gs_gumroad")

            btn_gen_gsites = st.button("🌐 Google Sites Landing Page Szöveg Generálása", key="btn_gen_gsites", use_container_width=True)

        with col_f3_out:
            st.markdown("<div class='step-label'>Kész Google Sites Struktúra & Másolható Tartalom</div>", unsafe_allow_html=True)

            if btn_gen_gsites:
                with st.spinner("0 Ft-os Google Sites landing page struktúra és szövegek generálása..."):
                    p_gs = build_google_sites_landing_page_prompt(
                        product_name=gs_prod_name,
                        target_audience=gs_target_aud,
                        headline=gs_headline,
                        tagline=gs_tagline,
                        lead_magnet_desc=gs_lead_magnet,
                        features=gs_features,
                        offer_price=gs_price,
                        amazon_url=gs_amazon,
                        etsy_url=gs_etsy,
                        gumroad_url=gs_gumroad,
                        style_theme=gs_style,
                        language=gs_lang.lower(),
                        niche_name=curr_niche_key
                    )
                    is_gs_en_call = "english" in gs_lang.lower() or "angol" in gs_lang.lower()
                    sys_inst_gs = (
                        "You are a world-class Conversion Rate Optimization and Landing Page copywriter for Google Sites. Output 100% fluent, natural English without mixing foreign languages."
                        if is_gs_en_call else
                        "Te egy világszínvonalú CRO és Landing Page szövegíró szakértő vagy a Google Sites keretrendszerhez. Válaszolj 100%-ban tiszta magyar nyelven."
                    )
                    ok_call, res_gs = km.generate_text_with_fallback(
                        prompt=p_gs,
                        model_name=text_model,
                        system_instruction=sys_inst_gs
                    )
                    if not ok_call:
                        res_gs = f"Hiba: {res_gs}"

                    st.session_state["gsites_landing_page_res"] = res_gs
                    st.session_state["gsites_prod_name"] = gs_prod_name

            if st.session_state.get("gsites_landing_page_res"):
                res_gs_val = st.session_state["gsites_landing_page_res"]
                p_gs_name = st.session_state.get("gsites_prod_name", gs_prod_name)

                st.markdown("""
                <div class='info-banner'>
                    <strong>💡 Hogyan hozd létre a 0 Ft-os Google Sites oldaladat 5 perc alatt?</strong><br>
                    1. Nyisd meg a <a href="https://sites.google.com" target="_blank" style="color:#34d399; font-weight:700;">sites.google.com</a> oldalt a Google fiókoddal.<br>
                    2. Kattints az <strong>Üres sablon (+)</strong> gombra az új oldal indításához.<br>
                    3. Másold be az alábbi blokkokat (Hero banner, 2-oszlopos Ingyenes Minta, 3-oszlopos Termékkártyák).<br>
                    4. Illeszd be az Amazon, Etsy és Gumroad/Stripe gombokat és publikáld 100%-ban ingyen saját vagy egyedi domain alatt!
                </div>
                """, unsafe_allow_html=True)

                st.markdown(res_gs_val)

                st.markdown("---")
                st.text_area(
                    "📋 Nyers Google Sites Szövegek Másolása (Ctrl+A → Ctrl+C):",
                    value=res_gs_val,
                    height=200,
                    key="gsites_copy_area"
                )

                col_dl_g1, col_dl_g2 = st.columns(2)
                with col_dl_g1:
                    st.download_button(
                        label="⬇️ Letöltés (.txt)",
                        data=res_gs_val.encode("utf-8"),
                        file_name=f"Marketing_{sanitize_filename(p_gs_name)}_Google_Sites_Landing_Page.txt",
                        mime="text/plain",
                        key="dl_gsites_txt",
                        use_container_width=True
                    )
                with col_dl_g2:
                    if DOCX_AVAILABLE:
                        docx_gs_bio = create_marketing_docx(p_gs_name, res_gs_val, header_info=f"Google Sites 0 Ft-os Landing Page - {p_gs_name}")
                        st.download_button(
                            label="⬇️ Letöltés (.docx)",
                            data=docx_gs_bio.getvalue(),
                            file_name=f"Marketing_{sanitize_filename(p_gs_name)}_Google_Sites_Landing_Page.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_gsites_docx",
                            use_container_width=True
                        )

                st.markdown("---")
                if st.button("💾 Mentés a 06_📌_MARKETING_ES_SEO mappába", key="btn_save_gsites_drive", use_container_width=True):
                    ok_s, details, _ = save_marketing_file_to_drive(
                        product_name=p_gs_name,
                        text_content=res_gs_val,
                        header_info=f"Google Sites Landing Page - {p_gs_name}",
                        content_type_tag="Google_Sites_Landing_Page"
                    )
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{details}`")
                    else:
                        st.error(f"Hiba a mentéskor: {details}")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>🌐</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Töltsd ki az adatokat a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">🌐 Google Sites Landing Page Szöveg Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────
    # AL-FÜL 5: AUTOMATA E-MAIL TÖLCSÉR (3-DAY & 30-DAY FUNNEL)
    # ─────────────────────────────────────────────────────
    with tab_ffc5:
        st.markdown("#### 📧 Automata E-mail Tölcsér Stúdió (3 Napos Indítás & 30 Napos Életút Csomag)")
        st.caption("Azonnal alkalmazható szekvenciák a '30 Email Marketing Bundle' alapján: Üdvözlés, Értékadás, Ajánlat, Bizonyítékok és Zárás.")

        em_mode_choice = st.radio(
            "Válassz E-mail Tölcsér Formátumot:",
            [
                "⚡ 3 Napos Gyors Indítás (Quick Launch Funnel)",
                "📬 30 Napos Teljes E-mail Rendszer (30 Email Marketing Bundle)"
            ],
            horizontal=True,
            key="em_mode_choice"
        )
        is_30day_mode = "30 Napos" in em_mode_choice

        col_f4_in, col_f4_out = st.columns([1, 1.1], gap="large")

        with col_f4_in:
            em_lang = st.selectbox("🌐 Nyelv (Language):", ["Magyar", "Angol (English)"], index=0, key="em_lang_select_v2")
            is_em_en = "Angol" in em_lang or "English" in em_lang
            em_lang_slug = "en" if is_em_en else "hu"

            if is_em_en:
                def_em_lm = f"Free Printable {curr_niche_data.get('name_en', 'Christian')} Sample Workbook"
                def_em_paid = f"Complete 30-Day {curr_niche_data.get('name_en', 'Christian')} Master Bundle"
                def_em_aud = "Parents, creators, and believers seeking daily quiet time and focus"
                def_em_disc = "25% exclusive welcome discount with coupon code PEACE25 (limited time)"
                def_em_offer = "30-day guided workbook, 30 printable high-resolution templates, daily gratitude tracker, and instant digital PDF access."
                def_em_story = "How I discovered peace, clarity, and daily focus in 10 minutes of quiet reflection during a chaotic season."
            else:
                def_em_lm = get_niche_field("em_lead_magnet", curr_niche_key)
                def_em_paid = get_niche_field("em_paid_prod", curr_niche_key)
                def_em_aud = get_niche_field("ffc_aud", curr_niche_key)
                def_em_disc = f"25% exkluzív üdvözlő kedvezmény a {get_niche_field('reels_cta_kw', curr_niche_key)}25 kuponkóddal (korlátozott határidő)"
                def_em_offer = "30 napos vezetett napló, 30 db művészi színező/fókusz kártya, napi hálaadás tracker, azonnali PDF hozzáférés."
                def_em_story = "Hogyan találtam meg a reggeli 10 perces csendességben a lelki békét és fókuszt egy kimerítő életszakaszomban."

            st.markdown(f"<div class='step-label'>{'30 Napos E-mail Csomag' if is_30day_mode else '3 Napos Tölcsér'} Paraméterek</div>", unsafe_allow_html=True)
            em_lead_magnet = st.text_input(
                "🎁 Ingyenes Csalitermék (Lead Magnet) Neve:",
                value=st.session_state.get(f"em_lm_{n_slug_ffc}_{em_lang_slug}", def_em_lm),
                key=f"em_lm_{n_slug_ffc}_{em_lang_slug}"
            )
            em_paid_prod = st.text_input(
                "💎 Értékesítendő Fizetős Termék / Ajánlat Neve:",
                value=st.session_state.get(f"em_paid_{n_slug_ffc}_{em_lang_slug}", def_em_paid),
                key=f"em_paid_{n_slug_ffc}_{em_lang_slug}"
            )
            em_target_aud = st.text_input(
                "Célközönség:",
                value=st.session_state.get(f"em_aud_{n_slug_ffc}_{em_lang_slug}", def_em_aud),
                key=f"em_aud_{n_slug_ffc}_{em_lang_slug}"
            )
            em_discount = st.text_input(
                "Exkluzív Kedvezmény / Ajánlat Kupon:",
                value=st.session_state.get(f"em_disc_{n_slug_ffc}_{em_lang_slug}", def_em_disc),
                key=f"em_disc_{n_slug_ffc}_{em_lang_slug}"
            )

            if is_30day_mode:
                em_core_offer_desc = st.text_area(
                    "Fő Ajánlat Részletes Leírása (Value Stack):",
                    value=st.session_state.get(f"em_offer_{n_slug_ffc}_{em_lang_slug}", def_em_offer),
                    height=70,
                    key=f"em_offer_{n_slug_ffc}_{em_lang_slug}"
                )
            else:
                em_story = st.text_area(
                    "Személyes Történet / Kapcsolódási Pont (2. naphoz):",
                    value=st.session_state.get(f"em_story_{n_slug_ffc}_{em_lang_slug}", def_em_story),
                    height=70,
                    key=f"em_story_{n_slug_ffc}_{em_lang_slug}"
                )

            btn_gen_email_funnel = st.button(
                f"📧 {'30 Napos E-mail Csomag' if is_30day_mode else '3 Napos E-mail Szekvencia'} Generálása (AI)",
                key="btn_gen_email_funnel",
                use_container_width=True
            )

        with col_f4_out:
            st.markdown(f"<div class='step-label'>Generált {'30 Napos' if is_30day_mode else '3 Napos'} E-mail Sorozat</div>", unsafe_allow_html=True)

            if btn_gen_email_funnel:
                with st.spinner(f"AI írja a {'30 napos teljes életút e-mail rendszert' if is_30day_mode else '3 napos automata tölcsért'}..."):
                    if is_30day_mode:
                        p_em_call = build_email_funnel_30day_prompt(
                            product_name=em_paid_prod,
                            target_audience=em_target_aud,
                            core_offer=em_core_offer_desc,
                            lead_magnet=em_lead_magnet,
                            discount_info=em_discount,
                            language=em_lang,
                            niche_name=curr_niche_key
                        )
                    else:
                        p_em_call = build_email_funnel_3day_prompt(
                            lead_magnet_name=em_lead_magnet,
                            paid_product_name=em_paid_prod,
                            target_audience=em_target_aud,
                            discount_offer=em_discount,
                            main_story=em_story,
                            language=em_lang,
                            niche_name=curr_niche_key
                        )

                    is_em_en_call = "english" in em_lang.lower() or "angol" in em_lang.lower()
                    sys_inst_em = (
                        "You are an elite lifecycle email marketing expert who writes warm, authentic, high-converting email sequences in 100% fluent English without mixing foreign languages."
                        if is_em_en_call else
                        "Te egy mester e-mail marketing specialista vagy, aki meleg, hiteles, emberi és nagy konverziójú e-mail szekvenciákat ír 100%-ban tiszta magyar nyelven."
                    )
                    ok_em, res_em = km.generate_text_with_fallback(
                        prompt=p_em_call,
                        system_instruction=sys_inst_em
                    )
                    st.session_state["ffc_em_res"] = res_em
                    st.session_state["ffc_em_prod"] = em_paid_prod
                    st.session_state["ffc_em_is_30day"] = is_30day_mode

            if st.session_state.get("ffc_em_res"):
                res_em_content = st.session_state["ffc_em_res"]
                p_em_name = st.session_state.get("ffc_em_prod", em_paid_prod)
                em_tag_type = "30Day_Email_Bundle" if st.session_state.get("ffc_em_is_30day") else "3Day_Email_Funnel"

                with st.container():
                    st.markdown(f"<div class='prompt-output'>{res_em_content}</div>", unsafe_allow_html=True)

                st.text_area(
                    "📋 Másold innen a teljes e-mail sorozatot (Ctrl+A → Ctrl+C):",
                    value=res_em_content,
                    height=240,
                    key="ffc_em_copy_area"
                )

                col_dl_em1, col_dl_em2 = st.columns(2)
                with col_dl_em1:
                    st.download_button(
                        label="⬇️ Letöltés (.txt)",
                        data=res_em_content.encode("utf-8"),
                        file_name=f"Marketing_{sanitize_filename(p_em_name)}_{em_tag_type}.txt",
                        mime="text/plain",
                        key="dl_ffc_em_txt",
                        use_container_width=True
                    )
                with col_dl_em2:
                    if DOCX_AVAILABLE:
                        docx_em_bio = create_marketing_docx(p_em_name, res_em_content, header_info=f"Automata E-mail Sorozat ({em_tag_type}) - {p_em_name}")
                        st.download_button(
                            label="⬇️ Letöltés (.docx)",
                            data=docx_em_bio.getvalue(),
                            file_name=f"Marketing_{sanitize_filename(p_em_name)}_{em_tag_type}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_ffc_em_docx",
                            use_container_width=True
                        )

                st.markdown("---")
                if st.button("💾 Mentés a 06_📌_MARKETING_ES_SEO mappába", key="btn_save_ffc_em_drive", use_container_width=True):
                    ok_s, details, _ = save_marketing_file_to_drive(
                        product_name=p_em_name,
                        text_content=res_em_content,
                        header_info=f"Automata E-mail Tölcsér ({em_tag_type}) - {p_em_name}",
                        content_type_tag=em_tag_type
                    )
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{details}`")
                    else:
                        st.error(f"Hiba a mentéskor: {details}")
            else:
                st.markdown(f"""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>📧</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg az adatokat a bal oldalon, majd<br>kattints az <strong style="color:#34d399;">📧 E-mail Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────
    # AL-FÜL 6: 30 NAPOS MARKETING & SOCIAL SEO NAPTÁR
    # ─────────────────────────────────────────────────────
    with tab_ffc6:
        st.markdown("#### 📅 30 Napos Marketing, Pinterest SEO & Social Media Naptár")
        st.caption("Komplett 30 napos organikus forgalomgeneráló stratégia: Pinterest SEO kulcsszavak, Instagram/TikTok horgok és Blog témák.")

        col_f5_in, col_f5_out = st.columns([1, 1.1], gap="large")

        with col_f5_in:
            cal_lang = st.selectbox(
                "🌐 Naptár Nyelve (Language):",
                ["Magyar", "Angol (English)"],
                index=0,
                key="cal_lang_select_v2"
            )
            is_cal_en = "angol" in cal_lang.lower() or "english" in cal_lang.lower()
            cal_lang_slug = "en" if is_cal_en else "hu"

            if is_cal_en:
                def_cal_prod = f"30-Day {curr_niche_data.get('name_en', 'Christian')} Mindful Journal"
                def_cal_aud = "Parents, creators, and believers seeking daily quiet time and focus"
                def_cal_top = "Week 1: Overcoming daily anxiety and finding spiritual peace\nWeek 2: Practical morning quiet time routines and focus habits\nWeek 3: Real transformations, gratitude journaling and mindfulness\nWeek 4: The complete bundle walkthrough, exclusive bonuses and limited offer"
            else:
                def_cal_prod = get_niche_field("cal_prod_name", curr_niche_key)
                def_cal_aud = get_niche_field("ffc_aud", curr_niche_key)
                def_cal_top = f"1. Hét: A(z) {get_niche_field('ffc_prod', curr_niche_key)} fő kihívásai és az akadályok leküzdése\n2. Hét: Gyakorlati lépések, fókusz és {get_niche_field('ffc_vehicle', curr_niche_key)}\n3. Hét: Esettanulmányok, sikerélmények és vizualizáció\n4. Hét: Záró ajánlat, bónuszok és sürgősség"

            st.markdown("<div class='step-label'>Tartalomnaptár Paraméterek</div>", unsafe_allow_html=True)
            cal_prod_name = st.text_input(
                "Termék / Csomag Neve:",
                value=st.session_state.get(f"cal_prod_{n_slug_ffc}_{cal_lang_slug}", def_cal_prod),
                key=f"cal_prod_{n_slug_ffc}_{cal_lang_slug}"
            )
            cal_target_aud = st.text_input(
                "Célközönség:",
                value=st.session_state.get(f"cal_aud_{n_slug_ffc}_{cal_lang_slug}", def_cal_aud),
                key=f"cal_aud_{n_slug_ffc}_{cal_lang_slug}"
            )
            cal_topics = st.text_area(
                "Fő Tartalmi Pillérek / Témák (opcionális):",
                value=st.session_state.get(f"cal_top_{n_slug_ffc}_{cal_lang_slug}", def_cal_top),
                height=85,
                key=f"cal_top_{n_slug_ffc}_{cal_lang_slug}"
            )
            cal_platforms = st.multiselect(
                "Célplatformok:",
                ["📌 Pinterest SEO", "📸 Instagram / Reels", "📱 TikTok", "📝 Blog / SEO Cikkek", "💌 Hírlevél"],
                default=["📌 Pinterest SEO", "📸 Instagram / Reels", "📝 Blog / SEO Cikkek"],
                key="cal_platforms"
            )

            btn_gen_social_cal = st.button("🚀 30 Napos Social SEO Naptár Generálása (AI)", key="btn_gen_social_cal", use_container_width=True)

        with col_f5_out:
            st.markdown("<div class='step-label'>Generált 30 Napos Tartalomterv</div>", unsafe_allow_html=True)

            if btn_gen_social_cal:
                with st.spinner("AI tervezi a 30 napos Pinterest SEO és Social Media naptárat..."):
                    p_cal_call = build_social_seo_calendar_30day_prompt(
                        product_name=cal_prod_name,
                        target_audience=cal_target_aud,
                        main_topics=cal_topics,
                        platforms=", ".join(cal_platforms),
                        language=cal_lang,
                        niche_name=curr_niche_key
                    )
                    sys_inst_cal = (
                        "You are a master Social Media Marketing and Pinterest SEO expert. Output a structured, clear, highly actionable 30-day content calendar in 100% fluent English without mixing foreign languages."
                        if is_cal_en else
                        "Te egy mester Social Media Marketing és Pinterest SEO szakértő vagy. Adj strukturált, átlátható és azonnal posztolható 30 napos tartalmi naptárat 100%-ban tiszta magyar nyelven."
                    )
                    ok_cal, res_cal = km.generate_text_with_fallback(
                        prompt=p_cal_call,
                        system_instruction=sys_inst_cal
                    )
                    st.session_state["ffc_cal_res"] = res_cal
                    st.session_state["ffc_cal_prod"] = cal_prod_name

            if st.session_state.get("ffc_cal_res"):
                res_cal_content = st.session_state["ffc_cal_res"]
                p_cal_name = st.session_state.get("ffc_cal_prod", cal_prod_name)

                with st.container():
                    st.markdown(f"<div class='prompt-output'>{res_cal_content}</div>", unsafe_allow_html=True)

                st.text_area(
                    "📋 Másold innen a teljes 30 napos naptárat (Ctrl+A → Ctrl+C):",
                    value=res_cal_content,
                    height=240,
                    key="ffc_cal_copy_area"
                )

                col_dl_c1, col_dl_c2 = st.columns(2)
                with col_dl_c1:
                    st.download_button(
                        label="⬇️ Letöltés (.txt)",
                        data=res_cal_content.encode("utf-8"),
                        file_name=f"Marketing_{sanitize_filename(p_cal_name)}_30Day_Social_Calendar.txt",
                        mime="text/plain",
                        key="dl_ffc_cal_txt",
                        use_container_width=True
                    )
                with col_dl_c2:
                    if DOCX_AVAILABLE:
                        docx_cal_bio = create_marketing_docx(p_cal_name, res_cal_content, header_info=f"30 Napos Social SEO Naptár - {p_cal_name}")
                        st.download_button(
                            label="⬇️ Letöltés (.docx)",
                            data=docx_cal_bio.getvalue(),
                            file_name=f"Marketing_{sanitize_filename(p_cal_name)}_30Day_Social_Calendar.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_ffc_cal_docx",
                            use_container_width=True
                        )

                st.markdown("---")
                if st.button("💾 Mentés a 06_📌_MARKETING_ES_SEO mappába", key="btn_save_ffc_cal_drive", use_container_width=True):
                    ok_s, details, _ = save_marketing_file_to_drive(
                        product_name=p_cal_name,
                        text_content=res_cal_content,
                        header_info=f"30 Napos Social Media & Pinterest SEO Naptár - {p_cal_name}",
                        content_type_tag="30Day_Social_Calendar"
                    )
                    if ok_s:
                        st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{details}`")
                    else:
                        st.error(f"Hiba a mentéskor: {details}")
            else:
                st.markdown("""
                <div style='text-align:center; padding: 50px 20px; border: 2px dashed #333f56; border-radius: 16px; background: #1e2536; color: #94a3b8;'>
                    <div style='font-size: 2.8rem;'>📅</div>
                    <div style='margin-top: 12px; font-size: 0.95rem; font-weight: 600; color:#e2e8f0;'>Add meg a témákat a bal oldalon, majd<br>kattints a <strong style="color:#34d399;">🚀 30 Napos Naptár Generálása</strong> gombra</div>
                </div>
                """, unsafe_allow_html=True)



# ══════════════════════════════════════════════════════════
# WORKSPACE: DFY CANVA SABLON & ERŐFORRÁS TÁR (RESOURCE HUB)
# ══════════════════════════════════════════════════════════

elif "Canva Sablon" in menu_choice or "DFY" in menu_choice or "9." in menu_choice:
    st.markdown("<div class='path-badge'>📦 DFY Erőforrás Tár & Sablonok</div>", unsafe_allow_html=True)
    st.markdown("### 📦 DFY Canva Sablonok & Digitális Erőforrás-Központ")
    st.caption("Azonnal használható, 1-kattintásos Canva sablonok, lead magnet vázlatok, 7-napos indítási stratégiák, 30 e-mail sablonok és Etsy SEO segédletek.")

    curr_niche_key = st.session_state.get("active_niche_choice", "✝️ Keresztény & Bibliai Rétegpiac (Alapértelmezett)")
    curr_niche_data = get_niche_prompt_context(curr_niche_key)

    st.markdown(f"""
    <div class='info-banner'>
        <strong>🎯 Aktív Niche Fókusz:</strong> {curr_niche_key}<br>
        <em>Minden sablon és struktúra azonnal átültethető a kiválasztott piaci kategóriába és a meglévő Canva szerkesztődbe!</em>
    </div>
    """, unsafe_allow_html=True)

    # 6 Rich Cards in 2 columns
    col_c1, col_c2 = st.columns(2, gap="large")

    with col_c1:
        # Card 1: Journal & Devotional Template
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>📖</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>1. Bibliai & Áhítat Napló Sablon (Canva DFY)</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    30 napos nyomtatható napló (8.5×11 inch & A4). Napi igehely, hálaadás blokk, ima kérések és heti önreflexiós lapok.
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 Sablon Struktúra & Canva Prompt Másolása", expanded=False):
                journal_prompt = (
                    "CANVA DESIGN SPECIFICATION: 30-Day Devotional & Prayer Journal (8.5x11 inches)\n\n"
                    "PAGE BREAKDOWN:\n"
                    "- Page 1: Elegant Minimalist Cover (Clean typography header, soft botanical frame)\n"
                    "- Page 2: 'This Journal Belongs To' Ownership Page with scripture verse\n"
                    "- Page 3: 30-Day Spiritual Goal & Daily Habit Tracker\n"
                    "- Pages 4-33: 30 Daily Devotional Journal Pages:\n"
                    "   * Top: Scripture Reference & Verse Quote Box\n"
                    "   * Middle: 'Key Takeaway & Reflection' lined writing area\n"
                    "   * Bottom-Left: 'My Daily Prayer' box with soft border\n"
                    "   * Bottom-Right: '3 Things I Am Grateful For Today' (3 checkbox lines)\n"
                    "- Page 34: Weekly Sabbath Reflection & Spiritual Check-in\n"
                    "- Page 35: Answered Prayers Celebration Log\n"
                    "- Page 36: Back Cover with blessing quote"
                )
                st.code(journal_prompt, language="text")
                st.link_button("🚀 Canva Megnyitása Új Tervezéshez", "https://www.canva.com/create/planners/", use_container_width=True)

        # Card 2: 7-Day Launch & 30-Day Marketing Strategy
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>🚀</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>2. 7-Napos Indítási Útmutató & Marketing Stratégia</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    Lépésről lépésre indítási ellenőrzőlista és 30 napos forgalomnövelési akcióterv a 0-tól az első 100 eladásig.
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 7-Napos Indítási Ellenőrzőlista Megtekintése", expanded=False):
                st.markdown("""
                - **0. Nap:** Termék PDF exportálása, Gumroad/Etsy feltöltés, árképzés (pl. $17-$27).
                - **1. Nap:** Ingyenes Lead Magnet (Csali) közzététele Pinteresten és Instagram Bioban.
                - **2. Nap:** Automata üdvözlő e-mail bekapcsolása (1. napos gyors győzelem).
                - **3. Nap:** Google Sites 0 Ft-os Landing Page élesítése az Amazon/Etsy/Gumroad gombokkal.
                - **4. Nap:** Első értékadó sztori e-mail kiküldése (A káoszból a megoldásig).
                - **5. Nap:** Exkluzív 24-48 órás bevezető kedvezmény kupon megnyitása.
                - **6. Nap:** Utolsó esély emlékeztető e-mail + 3 Pinterest Pin kiemelése.
                - **7. Nap:** Eredmények értékelése, átváltás a 30 napos organikus tartalomnaptárra.
                """)

        # Card 3: DFY Banners & Store Assets
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>🖼️</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>3. DFY Weboldal, Etsy és Social Banner Méretek</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    Pontos pixelméretek, elrendezési formulák és Canva promptok Google Sites-hoz, Etsy-hez, Facebookhoz és Pinteresthez.
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 Banner Méretek & Canva Specifikáció", expanded=False):
                st.markdown("""
                - **🌐 Google Sites / Shopify Hero Banner:** `1920 × 1080 px` (16:9 arány, bal oldalon cím, jobb oldalon termék mockup).
                - **🎨 Etsy Shop Big Banner:** `3360 × 840 px` (Elegáns logó, 3 fő előny, 5 csillagos értékelés badge).
                - **📱 Facebook Csoport / Oldal Fejléc:** `1640 × 924 px` (Középre zárt tartalom a mobil megjelenítéshez).
                - **📌 Pinterest Click-Magnet Pin:** `1000 × 1500 px` (2:3 arány, kontrasztos főcím és letöltési badge).
                - **📸 Instagram Carousel Diák:** `1080 × 1350 px` (4:5 arány, maximális képernyőkitöltés).
                """)

    with col_c2:
        # Card 4: Lead Magnet Template
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>🎁</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>4. Ingyenes Csali (Your Ultimate Lead Magnet)</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    5-8 oldalas magas konverziójú csalitermék sablon (Mini eBook, Gyors Útmutató vagy 3 Napos Kóstoló).
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 Lead Magnet Belső Szerkezet", expanded=False):
                st.markdown("""
                1. **Borító:** Kifejező cím + 'Azonnali Ingyenes Letöltés' badge.
                2. **Üdvözlés & Gyors Siker:** 1 bekezdés arról, miért ez a leghatékonyabb első lépés.
                3. **A 3 Kulcslépés:** Lényegretörő, azonnal végrehajtható gyakorlati tanácsok.
                4. **Munkafüzet / Kérdéssor:** 3 db kitölthető önreflexiós sor.
                5. **A Következő Lépés (The Bridge):** Kedvezményes kupon és közvetlen hivatkozás a teljes 30 napos csomagra!
                """)
                st.link_button("🚀 Canva Lead Magnet Tervező", "https://www.canva.com/create/ebooks/", use_container_width=True)

        # Card 5: 30 Email Marketing Bundle Templates
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>📧</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>5. 30 Email Marketing Bundle Sablontár</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    30 db előre megírt, kitöltős (fill-in-the-blanks) e-mail sablon 5 fázisban: Üdvözlés, Értékadás, Ajánlat, Bizonyíték, Sürgősség.
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 5 E-mail Fázis & Kitöltős Vázlat", expanded=False):
                st.markdown("""
                - **1. Fázis (1-5. nap):** *'Itt a(z) [CSALITERMÉK] letöltési linked!'* — Cél: Üdvözlés és azonnali sikerélmény.
                - **2. Fázis (6-12. nap):** *'A titok, amit bárcsak tudtam volna...'* — Cél: Mély kapcsolódás, történetmesélés és tiszta értékadás.
                - **3. Fázis (13-18. nap):** *'Bemutatom a(z) [FŐ AJÁNLAT]-ot'* — Cél: Value Stack és exkluzív bónuszok bemutatása.
                - **4. Fázis (19-24. nap):** *'Nem tudom, ez neked való-e, de...'* — Cél: Kifogások lebontása és vásárlói tapasztalatok.
                - **5. Fázis (25-30. nap):** *'Utolsó 24 óra: A(z) [KEDVEZMÉNY] ma éjfélkor zárul'* — Cél: Határidő és sürgősség.
                """)

        # Card 6: Etsy SEO, Pricing & Ads Guides
        with st.container():
            st.markdown("""
            <div class='sanctuary-card'>
                <div style='font-size: 1.8rem; margin-bottom: 6px;'>🛍️</div>
                <h4 style='margin: 0 0 6px 0; color: #34d399;'>6. Etsy SEO, Árazási & Ads Mester Stratégia</h4>
                <p style='color: #cbd5e1; font-size: 0.9rem; line-height: 1.6;'>
                    13 pontos kulcsszó-formula, profitabilitási kalkulátor és $1-$5/napos Etsy hirdetés-beállítási útmutató.
                </p>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📋 Etsy SEO & Árazási Útmutató", expanded=False):
                st.markdown("""
                - **13 SEO Tag Formula:** Használj több szavas kifejezéseket (long-tail keywords), pl. `Christian Journal PDF`, `Daily Devotional Printable`, `Prayer Tracker A4`.
                - **Árazási Arany Középút:** Egyedi printable termékek: `$7 - $12`; Komplett 30 napos mestercsomagok bónuszokkal: `$17 - $27`.
                - **Etsy Ads Tipp:** Indíts napi $1 - $3 költségkerettel, és kapcsold ki azokat a keresőszavakat 7 nap után, amik kattintást hoznak eladás nélkül.
                """)

    st.markdown("---")
    st.markdown("#### 💾 DFY Teljes Erőforrás Csomag Mentése")
    st.caption("Egyetlen kattintással elmentheted a fenti összes DFY sablonleírást, méretet és indítási stratégiát a Google Drive-ra.")

    all_dfy_content = f"""=== DFY ERŐFORRÁS TÁR & CANVA SABLONOK ===
Dátum: {time.strftime('%Y-%m-%d %H:%M:%S')}
Aktív Niche: {curr_niche_key}
Kategória: {curr_niche_data.get('group', 'General')}

1. BIBLIAI & ÁHÍTAT NAPLÓ CANVA SABLON SPECIFIKÁCIÓ:
- 8.5x11 hüvelyk és A4 méret, 36 oldalas struktúra.
- Napi igehely, önreflexió, ima kérések és hálaadás blokkok.

2. 7-NAPOS INDÍTÁSI ÚTMUTATÓ:
- 0. nap: Termék előkészítés, Gumroad/Etsy feltöltés.
- 1. nap: Lead magnet közzététel.
- 2. nap: Üdvözlő e-mail.
- 3. nap: Google Sites 0 Ft-os landing page élesítés.
- 4. nap: Sztori e-mail.
- 5. nap: Kedvezmény kupon megnyitása.
- 6-7. nap: Zárás és értékelés.

3. BANNER MÉRETEK:
- Google Sites / Shopify Hero: 1920x1080 px
- Etsy Shop Banner: 3360x840 px
- Facebook Header: 1640x924 px
- Pinterest Pin: 1000x1500 px (2:3)
- Etsy Shop Banner: 3360x840 px
- Facebook Header: 1640x924 px
- Pinterest Pin: 1000x1500 px (2:3)

4. LEAD MAGNET STRUKTÚRA:
- 5-8 oldalas mini eBook és munkafüzet.

5. 30 EMAIL MARKETING BUNDLE:
- 5 fázis: Üdvözlés, Értékadás, Ajánlat, Bizonyíték, Sürgősség.

6. ETSY SEO & ADS ÚTMUTATÓ:
- 13 long-tail tag, $17-$27 árképzés, $1-$3/napos mikro hirdetési keret.
"""

    col_dfy_b1, col_dfy_b2 = st.columns(2)
    with col_dfy_b1:
        st.download_button(
            label="⬇️ DFY Csomag Letöltése (.txt)",
            data=all_dfy_content.encode("utf-8"),
            file_name=f"DFY_Erőforrás_Csomag_{sanitize_filename(curr_niche_key)}.txt",
            mime="text/plain",
            key="dl_dfy_all_txt",
            use_container_width=True
        )
    with col_dfy_b2:
        if DOCX_AVAILABLE:
            docx_dfy_bio = create_marketing_docx(curr_niche_key, all_dfy_content, header_info=f"DFY Erőforrás Tár & Canva Sablonok - {curr_niche_key}")
            st.download_button(
                label="⬇️ DFY Csomag Letöltése (.docx)",
                data=docx_dfy_bio.getvalue(),
                file_name=f"DFY_Erőforrás_Csomag_{sanitize_filename(curr_niche_key)}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_dfy_all_docx",
                use_container_width=True
            )

    if st.button("💾 Mentés a 06_📌_MARKETING_ES_SEO mappába", key="btn_save_dfy_drive", use_container_width=True):
        ok_s, details, _ = save_marketing_file_to_drive(
            product_name=f"DFY_Erőforrás_Tár_{curr_niche_key}",
            text_content=all_dfy_content,
            header_info=f"DFY Canva Sablon & Erőforrás Tár - {curr_niche_key}",
            content_type_tag="DFY_Canva_Resource_Hub"
        )
        if ok_s:
            st.success(f"💾 **Sikeres mentés Google Drive-ra!**\n\n`{details}`")
        else:
            st.error(f"Hiba a mentéskor: {details}")




# ── MINDEN OLDALON ELÉRHETŐ ALSÓ GEMINI & NOTEBOOKLM GYORS-HÍD & CSEVEGŐ ──
if render_sidecar_dock:
    st.markdown("---")
    render_sidecar_dock()

# ─────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────

st.markdown("---")
col_f1, col_f2, col_f3 = st.columns(3)
with col_f1:
    st.markdown("<div style='text-align:center; color:#94a3b8; font-size:0.85rem; font-weight:500;'>📖 <strong>Amazon KDP</strong><br><a href='https://kdp.amazon.com' target='_blank' style='color:#34d399; text-decoration:none;'>kdp.amazon.com</a></div>", unsafe_allow_html=True)
with col_f2:
    st.markdown("<div style='text-align:center; color:#94a3b8; font-size:0.85rem; font-weight:500;'>🎨 <strong>Etsy Shop</strong><br><a href='https://etsy.com' target='_blank' style='color:#34d399; text-decoration:none;'>etsy.com</a></div>", unsafe_allow_html=True)
with col_f3:
    st.markdown("<div style='text-align:center; color:#94a3b8; font-size:0.85rem; font-weight:500;'>✝️ <strong>Gumroad PLR</strong><br><a href='https://gumroad.com' target='_blank' style='color:#34d399; text-decoration:none;'>gumroad.com</a></div>", unsafe_allow_html=True)



